"""
ITC Initiation Report — Section 1
Cover Page + Investment Summary + Table of Contents
Output: ~/ITC_report/ITC_Initiation_Report_2026-03-18.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── PATHS ─────────────────────────────────────────────────────────────────
BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

# ─── COLOURS ───────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x2B, 0x55)
GREEN  = RGBColor(0x1A, 0x75, 0x3F)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GOLD   = RGBColor(0xB7, 0x95, 0x0B)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY  = RGBColor(0x5D, 0x6D, 0x7E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
MIDBLUE= RGBColor(0xAA, 0xBE, 0xCF)

# ─── HELPERS ───────────────────────────────────────────────────────────────
def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def shade_cell(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic,
          color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing chart: {fname}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def sub_head(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sfont(r, size=size, bold=True, color=NAVY)

def body(doc, text, size=10, justify=True, indent=0, sb=0, sa=4):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    sfont(r, size=size)

# ─── CREATE DOC ────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(0.90)
    section.right_margin  = Inches(0.90)

# Header / Footer
for sec in doc.sections:
    hdr = sec.header
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("ITC LIMITED (ITC.NS)  |  INITIATING COVERAGE  |  BUY  |  ₹380 TARGET")
    sfont(r, size=8, color=DGRAY)

    ftr = sec.footer
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run("Institutional Equity Research  ·  March 18, 2026  ·  For Important Disclosures see Appendix")
    sfont(r2, size=8, italic=True, color=DGRAY)

# ══════════════════════════════════════════════════════════════════════════
# PAGE 1 — INVESTMENT SUMMARY
# ══════════════════════════════════════════════════════════════════════════

# Firm / product line
p0 = doc.add_paragraph()
p0.paragraph_format.space_before = Pt(4)
p0.paragraph_format.space_after  = Pt(2)
r0 = p0.add_run("EQUITY RESEARCH  ·  INITIATING COVERAGE  ·  INDIA CONSUMER & FMCG")
sfont(r0, size=8, color=DGRAY)

# Company name
p1 = doc.add_paragraph()
p1.paragraph_format.space_after = Pt(2)
r1a = p1.add_run("ITC LIMITED")
sfont(r1a, size=26, bold=True, color=NAVY)

# Sub-line
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run("NSE: ITC  ·  BSE: 500875  ·  Market Cap: ₹3,87,153 Cr (~USD 45B)  ·  "
                 "Sector: FMCG / Diversified Conglomerate")
sfont(r2, size=9.5, color=DGRAY)

# Rating + details table (left) | Stock chart placeholder (right)
meta = doc.add_table(rows=1, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
lc = meta.rows[0].cells[0]
rc = meta.rows[0].cells[1]
shade_cell(lc, NAVY)
shade_cell(rc, RGBColor(0xF5, 0xF8, 0xFA))

# Left cell — rating details
lp = lc.paragraphs[0]
lp.paragraph_format.space_before = Pt(5)
lp.paragraph_format.space_after  = Pt(4)

def lr(p, label, value, vc=WHITE, vb=False):
    rl = p.add_run(f"  {label:<28}")
    sfont(rl, size=9, color=MIDBLUE)
    rv = p.add_run(f"{value}\n")
    sfont(rv, size=9, bold=vb, color=vc)

lr(lp, "Rating",                    "BUY",                   GREEN, True)
lr(lp, "Current Price (18-Mar-26)", "₹306.80",               WHITE)
lr(lp, "12-Month Price Target",     "₹380",                  GOLD,  True)
lr(lp, "Upside",                    "+23.8%",                GREEN, True)
lr(lp, "Total Return (w/ dividend)","~27.3%",                GREEN)
lr(lp, "52-Week High / Low",        "₹444.20 / ₹296.45",    WHITE)
lr(lp, "Market Cap",                "₹3,87,153 Cr (USD 45B)",WHITE)
lr(lp, "Enterprise Value",          "₹3,59,153 Crore",       WHITE)
lr(lp, "Shares Outstanding",        "1,257 Crore",           WHITE)
lr(lp, "Net Cash",                  "₹28,000 Crore",         WHITE)
lr(lp, "FY26E EV/EBITDA",          "13.5x",                 WHITE)
lr(lp, "FY26E P/E",                "17.1x",                 WHITE)
lr(lp, "FY26E Dividend Yield",      "3.1%",                  WHITE)

# Right cell — analyst + headline
rp = rc.paragraphs[0]
rp.paragraph_format.space_before = Pt(5)

r_label = rp.add_run("RESEARCH TEAM\n")
sfont(r_label, size=8, bold=True, color=DGRAY)
r_name = rp.add_run("Institutional Equity Research\nMarch 18, 2026\n\n")
sfont(r_name, size=9, color=NAVY)

headline = (
    "ITC's stock has fallen ~31% from its 52-week high, overwhelmingly reflecting the "
    "January 2026 excise duty shock on cigarettes. We believe this correction is excessive. "
    "ITC has navigated 30+ years of successive excise hikes with volumes recovering each time. "
    "At 13.5x FY26E EV/EBITDA — a 55% discount to Indian FMCG peers — with ₹28,000 Crore "
    "of net cash, an FMCG business approaching profitability inflection, and the Hotels demerger "
    "providing a structural re-rating catalyst, the risk-reward is asymmetric. "
    "We initiate coverage with BUY and a ₹380 twelve-month price target."
)
r_hl = rp.add_run(headline)
sfont(r_hl, size=9.5, italic=True, color=NAVY)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Figure 1: Stock price chart
add_chart(doc, "chart_01_stock_price_performance.png", width=7.0, fig_num=1,
          caption="ITC Limited — Stock Price Performance vs. Nifty 50 & Nifty FMCG Index (12 Months to March 2026)")

# Gray header bar
tblh = doc.add_table(rows=1, cols=1)
tblh.style = "Table Grid"
hcell = tblh.rows[0].cells[0]
shade_cell(hcell, NAVY)
hp3 = hcell.paragraphs[0]
hp3.paragraph_format.space_before = Pt(2)
hp3.paragraph_format.space_after  = Pt(2)
rh = hp3.add_run(
    "  BUY RECOMMENDATION  ·  INITIATION OF COVERAGE  ·  12-MONTH PRICE TARGET: ₹380 (+23.8%)"
)
sfont(rh, size=10, bold=True, color=WHITE)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Four key bullets
bullets = [
    (
        "Market has materially over-penalised ITC for the January 2026 excise duty hike.",
        "The Central Excise (Amendment) Act 2025 imposed a 10–15x increase in specific cigarette "
        "duties effective January 1, 2026, and ITC's stock fell ~10% on announcement, extending a "
        "31% correction from its 52-week high of ₹444.20. This is an overreaction. ITC has navigated "
        "30+ years of successive excise hikes — including significant FY10 and FY17 increases — "
        "with volumes recovering within 2–3 quarters each time. Our base case assumes only a 2–3% "
        "annual volume decline, more than offset by 7–10% price increases preserving cigarettes EBIT "
        "at ₹15,900–16,000 Crore in FY26E. The first full-quarter data point (Q1 FY27 results, "
        "July 2026) will be the pivotal confirmation catalyst."
    ),
    (
        "FMCG-Others profitability inflection ascribed near-zero value by the market.",
        "The FMCG-Others segment (₹21,982 Crore revenue, FY25A) crossed EBIT breakeven in FY23 "
        "after 15+ years of brand-building investment. EBIT has grown from ₹120 Crore (FY23A) to "
        "₹1,050 Crore (FY25A) and we project ₹7,210 Crore by FY30E (7.0% margin). Our SOTP DCF "
        "assigns ₹62,351 Crore of EV to FMCG-Others vs. the market's implied ₹15,000–20,000 Crore "
        "— a 3–4x discount to intrinsic value. Aashirvaad (#1 branded atta in India), Sunfeast "
        "(top-2 biscuits), Bingo! (fast-growing snacks), and Yippee! (25–30% of instant noodles) "
        "carry durable distribution advantages through ITC's proprietary 2-million-outlet network. "
        "Each 1% improvement in FMCG-Others EBIT margin adds ~₹450 Crore (~2%) to consolidated EBIT."
    ),
    (
        "ITC Hotels demerger unlocks structural re-rating; ₹28,000 Crore net cash signals "
        "imminent capital return.",
        "ITC Hotels Limited listed on January 29, 2026, with ITC retaining a 40% strategic stake "
        "(~₹8,500 Crore). As ITC Hotels targets 220 properties by 2030 from 140+ today, the retained "
        "stake will compound. Separately, ITC's net cash has grown to ₹28,000 Crore (7.2% of market "
        "cap) with FY26E FCF of ₹18,880 Crore. With Hotels capex removed from the parent's books, "
        "we model a special dividend or buyback in the next 12–18 months — a ₹10,000 Crore buyback "
        "at CMP would retire ~3% of shares and add ₹0.50 to FY27E EPS. BAT's stake reduction "
        "(from ~29% to 22.9%) has also improved ITC's corporate governance trajectory."
    ),
    (
        "Valuation at multi-year lows; risk-reward is asymmetric at current levels.",
        "At ₹306.80, ITC trades at 13.5x FY26E EV/EBITDA vs. Indian FMCG median ~30x — the "
        "widest discount in our analysis. Our probability-weighted price target of ₹380 (12 months) "
        "reflects: ₹313 weighted intrinsic value today + ~20% forward re-rating as FMCG profitability "
        "crystallises. Bull case (25% probability): ₹480, assuming excise absorbed and FMCG margin "
        "reaches 10%+. Base case (55%): ₹380. Bear case (20%): ₹240 — only 22% downside. FCF "
        "yield of 4.9% (FY26E) vs. 10-year GoI bond at 6.8% is compelling for a business with "
        "30+ consecutive years of dividend growth and 85–90% payout ratio."
    ),
]

for bold_t, body_t in bullets:
    pb = doc.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pb.paragraph_format.space_before = Pt(3)
    pb.paragraph_format.space_after  = Pt(3)
    pb.paragraph_format.left_indent  = Inches(0.1)
    r_sq = pb.add_run("■ ")
    sfont(r_sq, size=10, bold=True, color=NAVY)
    r_bd = pb.add_run(bold_t + "  ")
    sfont(r_bd, size=10, bold=True, color=BLACK)
    r_by = pb.add_run(body_t)
    sfont(r_by, size=10, color=BLACK)

# ─── FINANCIAL SUMMARY TABLE ──────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(3)
cp2 = doc.add_paragraph()
r_cp = cp2.add_run(
    "Figure 2 – ITC Limited: Summary Financials & Valuation  (₹ Crore, except per-share; "
    "gross revenue basis for margin calc.)"
)
sfont(r_cp, size=8.5, bold=True, color=NAVY)

cols = ["Metric", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A",
        "FY26E", "FY27E", "FY28E", "FY29E", "FY30E"]
data = [
    ["Gross Revenue", "47,902", "59,214", "69,446", "71,985", "73,465",
     "77,640", "83,580", "90,020", "97,100", "1,05,000"],
    ["Net Revenue",   "27,448", "34,588", "39,460", "41,870", "43,011",
     "47,310", "52,140", "57,360", "62,100", "67,200"],
    ["  YoY Growth",  "—",      "26.1%",  "14.1%",  "6.1%",  "2.7%",
     "10.0%",  "10.2%",  "10.0%",  "8.3%",  "8.2%"],
    ["EBITDA",        "14,356", "17,458", "20,141", "21,857", "24,025",
     "26,670", "29,420", "32,500", "35,940", "39,750"],
    ["  EBITDA Margin","34.7%", "35.4%",  "35.5%",  "36.0%", "34.7%",
     "37.3%",  "37.6%",  "37.8%",  "38.3%", "38.7%"],
    ["EBIT",          "12,821", "15,732", "18,234", "19,780", "21,982",
     "24,650", "27,200", "30,200", "33,500", "37,160"],
    ["PAT (Shareholders)","12,946","15,150","18,638","19,766","22,592",
     "22,955", "25,120", "27,614", "30,396", "33,488"],
    ["  Net Margin",  "47.2%",  "43.8%",  "47.2%",  "47.2%", "52.5%",
     "48.5%",  "48.2%",  "48.1%",  "48.9%", "49.8%"],
    ["EPS (₹)",       "10.32",  "12.08",  "14.85",  "15.74", "17.97",
     "18.26",  "19.98",  "21.97",  "24.18", "26.64"],
    ["DPS (₹)",       "5.75",   "6.25",   "6.75",   "7.50",  "14.35",
     "15.50",  "17.00",  "18.65",  "20.55", "22.65"],
    ["Free Cash Flow","12,040", "13,290", "15,630", "16,100", "16,900",
     "18,880", "20,882", "23,124", "25,866", "28,668"],
    ["  FCF Margin",  "43.9%",  "38.4%",  "39.6%",  "38.5%", "39.3%",
     "39.9%",  "40.1%",  "40.3%",  "41.7%", "42.7%"],
    ["EV / EBITDA (x)","—",    "—",      "—",      "—",     "15.0x",
     "13.5x",  "12.2x",  "11.1x",  "10.0x",  "9.1x"],
    ["P / E (x)",     "—",      "—",      "—",      "—",     "17.1x",
     "16.8x",  "15.4x",  "14.0x",  "12.7x", "11.5x"],
    ["Dividend Yield","—",      "—",      "—",      "—",     "4.7%",
     "5.1%",   "5.5%",   "6.1%",   "6.7%",  "7.4%"],
]

ft = doc.add_table(rows=len(data)+1, cols=len(cols))
ft.style = "Table Grid"
ft.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, h in enumerate(cols):
    c = ft.rows[0].cells[ci]
    shade_cell(c, NAVY)
    cell_text(c, h, size=8, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

for ri, row in enumerate(data):
    is_main = not row[0].startswith("  ")
    bg = LGRAY if ri % 2 == 0 else WHITE
    for ci, val in enumerate(row):
        c = ft.rows[ri+1].cells[ci]
        shade_cell(c, bg)
        cell_text(c, val, size=8, bold=is_main,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(1)
sr = sp.add_run(
    "Source: ITC Annual Reports, SEBI filings, Institutional Equity Research estimates.  "
    "A = Actual; E = Estimate.  Margins on gross revenue basis for all years."
)
sfont(sr, size=7.5, italic=True, color=DGRAY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PAGE 2 — TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "TABLE OF CONTENTS")

toc = [
    ("Investment Summary",                                "1",  True),
    ("Table of Contents",                                 "2",  True),
    ("Investment Thesis",                                 "3",  True),
    ("    Pillar 1: Excise Duty Overreaction",            "3",  False),
    ("    Pillar 2: FMCG-Others Profitability Inflection","4",  False),
    ("    Pillar 3: Hotels Demerger Re-rating",           "4",  False),
    ("    Pillar 4: Net Cash & Capital Return",           "5",  False),
    ("    Pillar 5: Conglomerate Discount Closing",       "5",  False),
    ("Key Investment Risks",                              "6",  True),
    ("    Risk 1: Cigarette Volume Decline",              "6",  False),
    ("    Risk 2: FMCG-Others Competitive Intensity",    "6",  False),
    ("    Risk 3: Further Regulatory Tightening",         "7",  False),
    ("    Risk 4: BAT Stake Overhang",                    "7",  False),
    ("    Risk 5: Conglomerate Discount Persistence",     "7",  False),
    ("Company 101",                                       "8",  True),
    ("    Company Overview",                              "8",  False),
    ("    Company History (1910–2026)",                   "10", False),
    ("    Management Team",                               "12", False),
    ("    Governance & Ownership Structure",              "14", False),
    ("Products & Services",                               "15", True),
    ("    FMCG-Cigarettes",                               "15", False),
    ("    FMCG-Others (Branded Consumer Goods)",          "16", False),
    ("    Paperboards & Packaging",                       "17", False),
    ("    Agribusiness & e-Choupal",                      "18", False),
    ("    Hotels (ITC Hotels Limited — Demerged)",        "18", False),
    ("Customers & Go-to-Market",                          "19", True),
    ("Industry Overview",                                 "21", True),
    ("    Indian FMCG Market",                            "21", False),
    ("    Indian Cigarettes Industry",                    "22", False),
    ("    Hospitality & Paper Industries",                "23", False),
    ("Competitive Landscape",                             "24", True),
    ("Market Opportunity & TAM",                          "26", True),
    ("Financial Analysis (FY21A–FY25A)",                  "27", True),
    ("    Income Statement & Margin Analysis",            "27", False),
    ("    Cash Flow Analysis",                            "30", False),
    ("    Balance Sheet Review",                          "31", False),
    ("    Key Operating Metrics",                         "32", False),
    ("Financial Projections (FY26E–FY30E)",               "33", True),
    ("    Segment Revenue Assumptions",                   "33", False),
    ("    Scenario Analysis (Bull / Base / Bear)",        "35", False),
    ("Valuation Analysis",                                "36", True),
    ("    SOTP DCF Valuation",                            "36", False),
    ("    Consolidated DCF",                              "37", False),
    ("    Comparable Companies Analysis",                 "38", False),
    ("    Dividend Discount Model",                       "39", False),
    ("    Valuation Summary & Football Field",            "40", False),
    ("    Price Target & Recommendation",                 "40", False),
    ("Key Investment Catalysts",                          "41", True),
    ("Appendix A: Extended Risk Assessment",              "42", True),
    ("Appendix B: Modelling Assumptions",                 "44", True),
    ("Appendix C: Data Sources & Citations",              "45", True),
    ("Important Disclosures & Analyst Certification",     "47", True),
]

toc_t = doc.add_table(rows=len(toc), cols=2)
toc_t.style = "Table Grid"
for ri, (item, pg, bold) in enumerate(toc):
    bg = LGRAY if bold else WHITE
    shade_cell(toc_t.rows[ri].cells[0], bg)
    shade_cell(toc_t.rows[ri].cells[1], bg)
    cell_text(toc_t.rows[ri].cells[0], item, size=9,
              bold=bold, color=NAVY if bold else BLACK)
    cell_text(toc_t.rows[ri].cells[1], pg, size=9,
              bold=bold, color=NAVY if bold else BLACK,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_chart(doc, "chart_02_revenue_growth_trajectory.png", width=7.0,
          fig_num=3,
          caption="ITC Limited — Net Revenue & EBITDA Growth Trajectory (FY21A–FY30E)")

doc.add_page_break()

# ─── SAVE ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✅ Section 1 complete → {OUTPUT}")
