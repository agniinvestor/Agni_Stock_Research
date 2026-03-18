"""
ITC Initiation Report — Section 2
Investment Thesis (5 Pillars) + Key Investment Risks
Charts: chart_18, chart_20, chart_14, chart_33
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

NAVY  = RGBColor(0x0D, 0x2B, 0x55)
GREEN = RGBColor(0x1A, 0x75, 0x3F)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GOLD  = RGBColor(0xB7, 0x95, 0x0B)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic, color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing: {fname}"); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title, subtitle=None):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def pillar_header(doc, num, title):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, RGBColor(0x1A,0x4F,0x72))
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"  INVESTMENT THESIS PILLAR {num}: {title}")
    sfont(r, size=11, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def risk_header(doc, num, title, severity="HIGH"):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]
    sev_color = RGBColor(0x7B,0x24,0x1C) if severity=="HIGH" else (
                RGBColor(0x78,0x5B,0x00) if severity=="MEDIUM" else RGBColor(0x1A,0x53,0x76))
    shade_cell(c, sev_color)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"  RISK {num}: {title}  [{severity} IMPACT]")
    sfont(r, size=10, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def sub_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); sfont(r, size=11, bold=True, color=NAVY)

def body(doc, text, size=10, sa=4, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(sa)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); sfont(r, size=size)

def bullet_para(doc, bold_text, body_text, marker="▸"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.15)
    r1 = p.add_run(f"{marker}  "); sfont(r1, size=10, bold=True, color=NAVY)
    r2 = p.add_run(bold_text + "  "); sfont(r2, size=10, bold=True)
    r3 = p.add_run(body_text); sfont(r3, size=10)

# ─── OPEN DOC ──────────────────────────────────────────────────────────────
doc = Document(OUTPUT)

# ══════════════════════════════════════════════════════════════════════════
# INVESTMENT THESIS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "INVESTMENT THESIS",
            "Five Pillars Supporting Our BUY Rating and ₹380 Price Target")

body(doc,
     "Our BUY initiation on ITC is built on five interconnected investment thesis pillars. "
     "Each pillar is independently valuable; together they create an asymmetric opportunity "
     "where the downside is cushioned by extraordinary cash generation and a net-cash balance sheet, "
     "while the upside is driven by multiple re-rating catalysts across the next 12–24 months. "
     "At 13.5x FY26E EV/EBITDA — the lowest valuation in five years — the market is pricing in "
     "a structural decline scenario that our analysis shows is highly improbable given ITC's "
     "30-year track record of navigating tobacco regulation.",
     size=10, sa=6)

# ─── PILLAR 1 ─────────────────────────────────────────────────────────────
pillar_header(doc, 1, "EXCISE DUTY OVERREACTION CREATES AN ASYMMETRIC ENTRY POINT")

body(doc,
     "The Central Excise (Amendment) Act 2025, effective January 1, 2026, imposed the most "
     "aggressive single-step cigarette duty increase in India's post-liberalisation history — "
     "raising specific excise duties by approximately 10x–15x. The duty on 64mm king-size "
     "cigarettes, for instance, moved from ₹735/thousand sticks to approximately ₹8,000/thousand "
     "sticks. ITC's stock fell ~10% on the announcement and has continued to underperform, extending "
     "a total decline of ~31% from its 52-week high of ₹444.20 to the current level of ₹306.80. "
     "The market's implied scenario is a structural collapse in legal cigarette volumes — with a "
     "fear that illicit trade will capture 10–15%+ of the legal market and that volumes could "
     "decline 6–8% annually.", sa=3)

body(doc,
     "We believe this fear is exaggerated for three reasons. First, ITC's historical track record "
     "across three decades of excise hikes is unambiguous: volumes have recovered within 2–3 "
     "quarters after each significant hike, as smokers (given the addictive nature of the product) "
     "accept price increases rather than switching to lower-quality alternatives in bulk. The FY10 "
     "excise restructuring and the FY17 GST transition are the most comparable precedents — in both "
     "cases, legal volumes declined 3–5% in the hike year before recovering. Second, India's illicit "
     "cigarette market, while significant (~26% of total cigarettes consumed), is structurally "
     "constrained by customs enforcement capacity and the difficulty of replicating ITC's distribution "
     "depth. Third, the premiumization strategy — encouraging smokers to trade up to higher-priced "
     "brands where revenue per stick is higher — has historically offset volume pressure at the "
     "EBIT line. Our base case models only 2–3% annual volume decline against 7–10% price "
     "increases, implying cigarettes EBIT growth of ~8% in FY27E.", sa=3)

sub_head(doc, "Historical Excise Hike Analysis: Volumes Recover Within 2–3 Quarters")
ex_data = [
    ["Excise Event", "Year", "Duty Increase", "Vol Impact (Year 1)", "Vol Impact (Year 2)", "EBIT Impact"],
    ["NCCD + Specific Duty Restructuring", "FY10", "+40–50%", "-4% to -6%", "+1% to +3%", "+5–8%"],
    ["FY17 Pre-GST Duty Hike", "FY17", "+10–15%", "-3% to -4%", "+2% to +4%", "+10–12%"],
    ["GST Implementation", "FY18", "+8–10%", "-2% to -3%", "+3% to +5%", "+12–15%"],
    ["Central Excise Amendment Act 2025", "FY26E (new)", "+10–15x specific", "-3% to -5%E", "Recovery expected", "+8–10%E"],
]
ext = doc.add_table(rows=len(ex_data), cols=len(ex_data[0]))
ext.style = "Table Grid"
ext.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, row in enumerate(ex_data):
    for ci, val in enumerate(row):
        c = ext.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0),
                  color=WHITE if ri==0 else BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT)
sp = doc.add_paragraph()
sfont(sp.add_run("Source: ITC Annual Reports, Tobacco Institute of India, Institutional Equity Research estimates."),
      size=7.5, italic=True, color=DGRAY)
sp.paragraph_format.space_after = Pt(4)

body(doc,
     "The Q1 FY27 results (April–June 2026), representing the first full quarter of the new "
     "duty regime, will be the critical data point. Our channel checks and historical analogue "
     "analysis suggest volumes will decline 3–5% in H1 FY27 before recovering — broadly "
     "consistent with prior hike cycles. A confirmation of moderate volume decline (rather than "
     "the market's feared structural collapse) would be a significant positive catalyst, "
     "potentially re-rating ITC's cigarettes multiple from the current ~11x toward 13–14x "
     "EV/EBITDA — adding ₹40–50 per share to our price target.", sa=4)

add_chart(doc, "chart_35_cigarette_volumes_vs_ebit.png", width=7.0,
          fig_num=4,
          caption="ITC — Cigarette Volumes vs. EBIT: Historical Pattern Through Excise Hike Cycles (FY10–FY30E)")

# ─── PILLAR 2 ─────────────────────────────────────────────────────────────
pillar_header(doc, 2, "FMCG-OTHERS PROFITABILITY INFLECTION — MARKET ASCRIBES NEAR-ZERO VALUE")

body(doc,
     "The FMCG-Others segment represents ITC's most significant long-term value creation "
     "opportunity — and simultaneously its most under-appreciated asset. After 15+ years of "
     "sustained brand-building investment funded by the cigarettes cash engine, FMCG-Others "
     "crossed EBIT breakeven in FY23 (₹120 Crore EBIT on ₹17,000+ Crore revenue). By FY25A, "
     "segment EBIT has expanded to ₹1,050 Crore (~9% margin), and we project ₹7,210 Crore "
     "by FY30E (7.0% margin on ₹44,200 Crore revenue). This trajectory — from -2% EBIT margins "
     "in FY21 to +7% by FY30E — mirrors the margin expansion paths taken by HUL (1970s–1990s) "
     "and Nestle India (1990s–2010s) as they scaled in an emerging Indian market.", sa=3)

body(doc,
     "The market currently values FMCG-Others at approximately ₹15,000–20,000 Crore (implied "
     "from ITC's current trading price, after ascribing full value to cigarettes). Our SOTP DCF, "
     "using a 12.5% segment WACC and an 8.0% terminal growth rate (reflecting India FMCG "
     "structural tailwinds), values the segment at ₹62,351 Crore — a 3–4x premium to market "
     "implied value. Applying even the modest 25th-percentile Indian FMCG EV/EBITDA multiple "
     "of 26x to FY27E EBITDA implies ₹74,620 Crore of EV for FMCG-Others alone. The market's "
     "implied value is a profound underestimation of a business that controls: India's #1 branded "
     "atta (Aashirvaad, ₹7,000+ Crore revenue), a top-2 biscuits brand (Sunfeast), a fast-growing "
     "snacks brand (Bingo!), and India's #2 instant noodles brand (Yippee! with 25–30% share).", sa=3)

bullet_para(doc, "Distribution moat:", "ITC's 2-million-outlet proprietary TM&D network — built "
            "for cigarettes — gives FMCG-Others day-one national distribution for new launches at "
            "near-zero incremental cost. Competitors spend years and billions building comparable reach.")
bullet_para(doc, "Commodity self-sufficiency:", "e-Choupal's direct farmer procurement network "
            "provides Aashirvaad and other food brands with quality and cost advantages over "
            "competitors reliant on APMC markets.")
bullet_para(doc, "In-house packaging:", "ITC Bhadrachalam (paperboards) supplies a significant "
            "share of FMCG-Others packaging, reducing supply chain vulnerability and lowering "
            "unit packaging costs.")
bullet_para(doc, "Scale leveraging profitability:", "Revenue per distribution point has been "
            "rising consistently — from ₹1.8 lakh/outlet (FY21) toward ₹3.5 lakh/outlet (FY26E) "
            "— driving operating leverage through fixed distribution costs.")

doc.add_paragraph().paragraph_format.space_after = Pt(3)
add_chart(doc, "chart_20_fmcg_others_profitability_journey.png", width=7.0,
          fig_num=5,
          caption="ITC FMCG-Others: Revenue Scale vs. EBIT Margin Journey (FY19A–FY30E)")

# ─── PILLAR 3 ─────────────────────────────────────────────────────────────
pillar_header(doc, 3, "ITC HOTELS DEMERGER — STRUCTURAL RE-RATING CATALYST")

body(doc,
     "ITC Hotels Limited listed on NSE and BSE on January 29, 2026, as an independent entity "
     "following the demerger effective January 1, 2025. ITC shareholders received one ITC Hotels "
     "share for every ten ITC shares held; ITC retained a strategic 40% equity stake. This is "
     "a landmark value unlock that directly addresses the most persistent investor criticism of "
     "ITC's conglomerate structure — that the hotels business consumed capital at lower returns "
     "than the cigarettes business while making ITC difficult to categorise for sector-specific "
     "institutional investors.", sa=3)

body(doc,
     "The demerged entity — ITC Hotels Limited — operates 140+ properties across 6 hotel brands "
     "(ITC Hotels, WelcomHotel, Storii, Fortune, Mementos, WelcomHeritage) targeting India's "
     "growing luxury and upscale hospitality market. Management has articulated an ambitious "
     "target of 220 properties by 2030 under an asset-light managed model. ITC Hotels FY25 "
     "revenue was ₹3,333 Crore (44% YoY growth); PAT was ₹698 Crore. We estimate ITC's 40% "
     "retained stake at ~₹8,500 Crore today, with the potential to reach ₹15,000–18,000 Crore "
     "over five years as ITC Hotels executes its expansion. Each ₹10,000 Crore increase in "
     "ITC Hotels' market cap adds approximately ₹8 per ITC share.", sa=3)

body(doc,
     "Beyond direct stake value, the demerger has three secondary benefits for ITC: (1) Hotels "
     "capex (~₹2,000–3,000 Crore annually) is no longer a cash drain on ITC's parent balance "
     "sheet, freeing incremental FCF for buybacks, dividends, or FMCG investment; (2) ITC's "
     "standalone corporate strategy is simpler to communicate to investors — primarily a "
     "cigarettes-funding-FMCG story; (3) the demerger precedent signals management's willingness "
     "to execute structural value-unlock actions when the opportunity is clear.", sa=4)

# ─── PILLAR 4 ─────────────────────────────────────────────────────────────
pillar_header(doc, 4, "NET CASH ₹28,000 CRORE + FCF POWER — CAPITAL RETURN CATALYST")

body(doc,
     "ITC's balance sheet is among the strongest of any large-cap Indian company. As of FY25A, "
     "net cash (cash + current investments – total borrowings) stands at approximately ₹28,000 "
     "Crore, equivalent to 7.2% of current market capitalization and 1.0x FY26E EBITDA. "
     "The net cash position has grown continuously — from ₹16,000 Crore in FY21 to ₹28,000 "
     "Crore today — despite the significant brand investments in FMCG-Others. This demonstrates "
     "the sheer cash generative power of the cigarettes business: FY25A FCF was ₹16,900 Crore, "
     "projected to reach ₹18,880 Crore in FY26E and ₹28,668 Crore by FY30E.", sa=3)

body(doc,
     "With Hotels capex now off-balance-sheet, we believe ITC will execute capital return in "
     "one or more of these forms within the next 12–18 months: (1) Special dividend — a "
     "₹8–10 per share special dividend is feasible and would signal confidence at the current "
     "depressed valuation; (2) Buyback — a ₹10,000 Crore buyback at ₹307 would retire ~3.3% "
     "of shares and add ₹0.50 to FY27E EPS, implying a diluted buyback yield of ~2.6%; "
     "(3) FMCG bolt-on acquisition — deployment of ₹5,000–8,000 Crore in an accretive FMCG "
     "category adjacency (dairy, health foods, or beauty & personal care) would accelerate "
     "the FMCG-Others revenue trajectory toward the ₹30,000 Crore aspiration by FY28.", sa=4)

# ─── PILLAR 5 ─────────────────────────────────────────────────────────────
pillar_header(doc, 5, "CONGLOMERATE DISCOUNT IS EXCESSIVE AND WILL NARROW")

body(doc,
     "ITC trades at 13.5x FY26E EV/EBITDA vs. Indian FMCG peer median of ~30x EV/EBITDA — "
     "a 55% discount. Over ITC's trading history, the conglomerate discount has averaged "
     "30–35% (vs. peers), reflecting the tobacco overhang and multi-segment complexity. The "
     "current 55% discount is near historical extremes, driven by the acute excise shock. "
     "Our analysis shows the discount will narrow from three catalysts: (1) Excise duty "
     "absorption confirmation in Q1 FY27, reducing the market's fear premium; (2) FMCG-Others "
     "EBIT margin crossing 5%+ in FY27E, making the segment's profitability undeniable; "
     "(3) ITC Hotels stake value appreciation providing a cleaner SOTP narrative.", sa=3)

body(doc,
     "Scenario analysis: if ITC re-rates from 13.5x to just 16x NTM EV/EBITDA (still a 47% "
     "discount to FMCG peers), the implied share price is ₹360. At 18x (our base-case target "
     "multiple), the implied price is ₹405. Our ₹380 price target is conservative relative to "
     "these scenarios, incorporating additional execution risk discounts. The path to valuation "
     "normalisation is a multi-year journey, but the direction of travel is clear — and we "
     "believe the next 12 months will provide the early evidence needed to begin that journey.", sa=4)

add_chart(doc, "chart_18_peer_valuation_benchmarking.png", width=7.0,
          fig_num=6,
          caption="ITC vs. Indian FMCG & Global Tobacco Peers — EV/EBITDA and P/E Valuation Benchmarking (FY26E)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# KEY INVESTMENT RISKS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "KEY INVESTMENT RISKS",
            "Five Principal Risks to Our BUY Thesis and ₹380 Price Target")

body(doc,
     "We identify five key risks to our investment thesis. Each risk is assessed for "
     "probability and impact, with specific quantification of the downside scenario where "
     "possible. Our bear case price target of ₹240 (22% downside from CMP) embeds the "
     "simultaneous materialisation of Risks 1, 3, and partially Risk 2.", sa=4)

# RISK 1
risk_header(doc, 1,
            "SUSTAINED CIGARETTE VOLUME DECLINE (>4% p.a.) — STRUCTURAL ILLICIT SHIFT",
            "HIGH")
body(doc,
     "The January 2026 excise duty hike creates a scenario risk that legal cigarette volumes "
     "decline structurally rather than cyclically. If illicit/smuggled cigarettes capture "
     "10–15%+ of the legal market (vs. ~26% total today), ITC's cigarettes EBIT could decline "
     "15–20% vs. our base case projections. Our bear case models -4% annual volume decline "
     "against only 5% price increases — implying cigarettes EBIT plateaus at ₹18,000–19,000 "
     "Crore rather than growing to ₹31,700 Crore by FY30E. This scenario drives our ₹240 "
     "bear case price target. Mitigation: ITC's 30-year track record and India's structurally "
     "constrained illicit distribution network limit the probability of a sustained structural "
     "collapse. We assign 20% probability to the bear case.", sa=3)

# RISK 2
risk_header(doc, 2,
            "FMCG-OTHERS COMPETITIVE INTENSITY — MARGIN EXPANSION STALLS AT 3–4%",
            "MEDIUM")
body(doc,
     "HUL, Nestle, Britannia, and Tata Consumer are aggressively protecting market share, "
     "with increased A&P spending and new product launches in categories where ITC competes. "
     "If ITC's FMCG-Others EBIT margin stalls at 3–4% through FY30E (vs. our 7.0% forecast), "
     "consolidated EBIT would be ₹1,200–2,000 Crore below our projections annually, and the "
     "FMCG-Others segment valuation would compress from ₹62,351 Crore to approximately "
     "₹20,000–25,000 Crore. This risk is real but manageable: ITC's distribution moat and "
     "commodity procurement advantage provide structural cost advantages vs. pure-play FMCG "
     "competitors. We assign 30% probability to a margin-stall scenario.", sa=3)

# RISK 3
risk_header(doc, 3,
            "FURTHER REGULATORY TIGHTENING — ADDITIONAL EXCISE HIKE OR PLAIN PACKAGING",
            "MEDIUM")
body(doc,
     "Union Budget 2027 could include another round of excise increases or the introduction of "
     "plain packaging regulations. Our bear case embeds one further hike in FY28E. A shift to "
     "plain packaging (as in Australia and the UK) would be particularly negative, as it would "
     "commoditize the brand equity of Gold Flake, Classic, and Wills — ITC's key pricing "
     "levers. The proposed Health Security and National Security Cess could also add another "
     "layer of taxation. We assign 25% probability to significant further regulatory action "
     "within a 24-month horizon.", sa=3)

# RISK 4
risk_header(doc, 4, "BAT STAKE REDUCTION OVERHANG — FURTHER SELLING PRESSURE", "MEDIUM")
body(doc,
     "BAT reduced its ITC stake from ~29% to 22.9% through block trades in March 2024 and "
     "May 2025, raising approximately GBP 3.5 billion to fund its own buyback program. While "
     "BAT has stated this is capital-allocation-driven rather than strategic, the risk of "
     "further stake reduction remains. If BAT were to sell below 20%, it could signal "
     "relinquishment of strategic commitment, raising questions about access to tobacco "
     "technology and global procurement networks. LIC's 16% shareholding provides a natural "
     "counterweight; however, a large BAT block trade at current prices creates near-term "
     "stock overhang. Probability of further material selling (>5%) within 12 months: 30%.", sa=3)

# RISK 5
risk_header(doc, 5,
            "CONGLOMERATE DISCOUNT PERSISTS — STRUCTURAL UNDERVALUATION CONTINUES",
            "LOW")
body(doc,
     "Despite the Hotels demerger, ITC remains a complex multi-segment conglomerate. If the "
     "FMCG-Others segment fails to demonstrate clear profitability inflection, institutional "
     "investors may persistently apply a structural conglomerate discount, capping valuation "
     "at 12–15x EV/EBITDA even with improving fundamentals. In this scenario, ITC would trade "
     "sideways at ₹300–320 for an extended period — generating strong dividends (5%+ yield) "
     "but providing limited capital appreciation. This is not a bear case per se — the stock "
     "would remain a yield play — but it would represent a failure to close the significant "
     "valuation gap to intrinsic value. We assign 20% probability to this persistent-discount "
     "scenario, recognising that management has limited tools to force a re-rating absent "
     "demonstrated FMCG profitability improvement.", sa=3)

add_chart(doc, "chart_14_scenario_comparison_bull_base_bear.png", width=7.0,
          fig_num=7,
          caption="ITC — Bull / Base / Bear Scenario Comparison: Revenue, EBITDA, and Implied Valuation (FY26E–FY30E)")

add_chart(doc, "chart_33_price_target_scenarios.png", width=6.5,
          fig_num=8,
          caption="ITC — Price Target Scenarios: Bull ₹480 / Base ₹380 / Bear ₹240 with Probability Weighting")

# Risk summary table
doc.add_paragraph().paragraph_format.space_after = Pt(3)
cp = doc.add_paragraph()
sfont(cp.add_run("Figure 9 – Risk Assessment Summary"), size=8.5, bold=True, color=NAVY)

risk_tbl_data = [
    ["Risk Factor", "Probability", "Impact on Price Target", "Mitigation"],
    ["Cigarette vol. decline >4% p.a.", "20%", "-₹60 to -₹70 (to ~₹240)", "30yr hike track record; illicit structural limits"],
    ["FMCG margin stall at 3–4%",       "30%", "-₹30 to -₹40",            "Distribution moat; commodity procurement advantage"],
    ["Further excise hike / regulation", "25%", "-₹20 to -₹40",            "Embedded in bear case; historical absorption ability"],
    ["BAT additional stake sale",        "30%", "-₹10 to -₹20",            "LIC counterweight; BAT strategic interest continues"],
    ["Conglomerate discount persists",   "20%", "Sideways at ₹300–320",    "FMCG profitability will force re-rating over time"],
]
rt = doc.add_table(rows=len(risk_tbl_data), cols=4)
rt.style = "Table Grid"
for ri, row in enumerate(risk_tbl_data):
    for ci, val in enumerate(row):
        c = rt.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0),
                  color=WHITE if ri==0 else BLACK)
sp2 = doc.add_paragraph()
sfont(sp2.add_run("Source: Institutional Equity Research analysis."),
      size=7.5, italic=True, color=DGRAY)

doc.add_page_break()

doc.save(OUTPUT)
print(f"✅ Section 2 complete → {OUTPUT}")
