"""
ITC Initiation Report — Section 3
Company 101: Overview + History + Management + Governance
Charts: chart_05, chart_06, chart_07, chart_24, chart_27
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

NAVY  = RGBColor(0x0D, 0x2B, 0x55)
GREEN = RGBColor(0x1A, 0x75, 0x3F)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
TEAL  = RGBColor(0x0E, 0x6B, 0x6B)

def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic, color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing: {fname}"); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title, subtitle=None):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sub_head(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    sfont(r, size=11, bold=True, color=color or NAVY)

def sub2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); sfont(r, size=10, bold=True, color=TEAL)

def body(doc, text, size=10, sa=4, sb=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); sfont(r, size=size)

def bullet(doc, text, indent=0.2, marker="•"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent  = Inches(indent)
    r = p.add_run(f"{marker}  {text}"); sfont(r, size=10)

def mgmt_bio(doc, name, title, text):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, RGBColor(0x1A,0x4F,0x72))
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"  {name}  —  ")
    sfont(r1, size=11, bold=True, color=WHITE)
    r2 = p.add_run(title)
    sfont(r2, size=10, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    body(doc, text, sa=6)

# ─── OPEN DOC ──────────────────────────────────────────────────────────────
doc = Document(OUTPUT)

# ══════════════════════════════════════════════════════════════════════════
# COMPANY 101 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "COMPANY 101 — COMPANY OVERVIEW",
            "From Imperial Tobacco (1910) to India's Unique FMCG Conglomerate (2026)")

body(doc,
     "ITC Limited is India's most structurally unique conglomerate: a business that generates "
     "exceptional, near-monopoly returns from one of the world's most heavily regulated industries "
     "and systematically deploys that capital to build the consumer goods company of the next "
     "generation. Incorporated on August 24, 1910, as the Imperial Tobacco Company of India "
     "Limited — a subsidiary of British-American Tobacco — ITC has evolved over 115 years from "
     "a pure tobacco manufacturer into a diversified Indian enterprise with principal operations "
     "across five verticals: FMCG-Cigarettes, FMCG-Others (packaged foods and consumer goods), "
     "Hotels (now partially demerged as ITC Hotels Limited), Paperboards/Paper/Packaging, and "
     "Agribusiness.", sa=3)

body(doc,
     "ITC Limited (NSE: ITC, BSE: 500875) is headquartered in Kolkata's Virginia House, a "
     "landmark colonial-era building that has been the company's registered office since 1928. "
     "The company employs approximately 26,000–30,000 people across its businesses and had "
     "consolidated gross revenues of ₹73,465 Crore in FY25 — making it one of India's largest "
     "consumer-facing corporates. Market capitalisation as of March 18, 2026 stands at "
     "approximately ₹3,87,153 Crore (USD ~45 billion), the twelfth largest company on the "
     "National Stock Exchange by market capitalisation.", sa=3)

add_chart(doc, "chart_05_company_overview_snapshot.png", width=7.0,
          fig_num=10,
          caption="ITC Limited — Business Snapshot: Revenue Mix, EBIT Contribution & Strategic Positioning (FY25A)")

sub_head(doc, "Revenue Architecture and Profit Asymmetry")
body(doc,
     "The most distinctive feature of ITC's financial architecture is its profound profit "
     "asymmetry: cigarettes contribute approximately 44% of gross revenues but generate "
     "approximately 78% of segment EBIT. This is not a structural weakness but a deliberate "
     "strategic construct — the extraordinary free cash flows from the oligopolistic cigarettes "
     "business fund the brand-building investments required to build next-generation FMCG "
     "franchises that would be prohibitively expensive to establish on standalone capital.", sa=2)

# Segment metrics table
seg_data = [
    ["Segment", "FY25A Revenue (₹ Cr)", "% of Total", "FY25A EBIT (₹ Cr)", "EBIT Margin", "Growth Outlook"],
    ["FMCG-Cigarettes",    "32,631", "44%", "21,091", "58–65%*", "5–7% EBIT CAGR (pricing-led)"],
    ["FMCG-Others",        "21,982", "30%", "1,050",  "~4.8%",   "15–18% revenue; margin expanding"],
    ["Agribusiness",       "19,753", "27%", "850",    "~4.3%",   "9–12% revenue CAGR"],
    ["Paperboards & Paper","8,200",  "11%", "890",    "~10.9%",  "Recovery from FY26E (anti-dumping)"],
    ["ITC Infotech",       "3,480",  "5%",  "490",    "~14.1%",  "12–15% CAGR (IT services growth)"],
    ["Hotels (discontinued)","3,333","5%",  "210",    "~6.3%",   "Demerged Jan 2026; ITC retains 40%"],
    ["Unallocated / Elim.", "—",     "—",   "(2,599)","—",       "Corporate costs"],
    ["TOTAL (Consolidated)","73,465","100%","24,025 EBITDA","34.7%","9–10% net rev CAGR (base)"],
]
st = doc.add_table(rows=len(seg_data), cols=6)
st.style = "Table Grid"
for ri, row in enumerate(seg_data):
    for ci, val in enumerate(row):
        c = st.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0 or ri==len(seg_data)-1),
                  color=WHITE if ri==0 else BLACK,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci in [1,2,3,4] else WD_ALIGN_PARAGRAPH.LEFT)
sp = doc.add_paragraph()
sfont(sp.add_run("Source: ITC Annual Report FY25, SEBI filings.  * Cigarettes EBIT margin on net segment revenues."),
      size=7.5, italic=True, color=DGRAY)

sub_head(doc, "Why ITC Is Structurally Unique")
body(doc,
     "No other Indian company has assembled a comparable combination of: (1) a domestic oligopoly "
     "in a heavily regulated, sin-taxed consumer product generating 60%+ EBIT margins; "
     "(2) a nascent but growing FMCG portfolio (now EBIT-positive) funded entirely from internal "
     "accruals; (3) one of India's largest rural internet and e-commerce infrastructure networks "
     "(e-Choupal, 35,000+ villages); (4) an integrated paperboards-to-packaging value chain "
     "servicing its own FMCG and cigarettes businesses; and (5) a luxury hospitality brand with "
     "pan-India presence now pursuing independent growth as a listed entity. This combination "
     "creates complexity that institutional investors struggle to price, but also creates resilience: "
     "when one segment faces headwinds (as paperboards did in FY25 from Chinese paper dumping), "
     "the cigarettes cash engine absorbs the impact without compromising investment in growth "
     "businesses.", sa=3)

body(doc,
     "The 'ITC Next' strategy articulated by Chairman & MD Sanjiv Puri encompasses three pillars: "
     "(1) Accelerated FMCG growth — 100+ new product launches annually, targeting ₹30,000+ Crore "
     "in FMCG-Others by FY28; (2) Digital and agritech investment — e-Choupal 4.0, targeting "
     "70,000 villages with integrated agri-services aggregation; (3) Structural simplification — "
     "the Hotels demerger being the landmark example, with further portfolio reviews possible for "
     "Paperboards and Infotech over the medium term. The ₹20,000 Crore medium-term manufacturing "
     "expansion plan announced at the FY25 AGM reflects management's conviction in India's "
     "consumption growth story.", sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# COMPANY HISTORY
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "COMPANY HISTORY", "115 Years: From Imperial Tobacco to India's FMCG Challenger")

body(doc,
     "ITC's 115-year journey from a British colonial tobacco company to an Indian conglomerate "
     "aspiring to FMCG leadership is one of the most remarkable corporate transformations in "
     "Indian business history — a story of Indianisation, diversification, and the disciplined "
     "deployment of tobacco cash flows into next-generation consumer businesses.",
     sa=3)

sub_head(doc, "Founding & Colonial Origins (1910–1947)")
body(doc,
     "ITC Limited was established on August 24, 1910, as the Imperial Tobacco Company of India "
     "Limited, registered in Kolkata as a subsidiary of the British-American Tobacco Company (BAT). "
     "The company's genesis was rooted in the colonial tobacco trade, succeeding W.D. & H.O. Wills's "
     "Indian operations to manufacture and distribute cigarettes to a growing Indian market. In 1911, "
     "the company entered partnerships with farmers in southern India to source leaf tobacco. The "
     "Indian Leaf Tobacco Development Company was formed in the Guntur district of Andhra Pradesh "
     "in 1912, establishing what would become ITC's deep agribusiness roots. The first cigarette "
     "factory was established in Bangalore in 1913. By 1926, the company had acquired its landmark "
     "Kolkata property, and by 1928, construction had commenced on 'Virginia House' — the Kolkata "
     "headquarters that remains in use today.", sa=3)

sub_head(doc, "Indianisation & Post-Independence (1947–1970)")
body(doc,
     "Post-Independence India necessitated progressive Indianisation of British-owned companies. "
     "ITC was converted into a Public Limited Company on October 27, 1954, with an initial Indian "
     "shareholding of 6%. The pivotal moment came in 1969, when Ajit Narain Haskar became the "
     "first Indian Chairman of ITC, marking the formal transition from British to Indian executive "
     "leadership. The company's name was changed from 'Imperial Tobacco Company' to 'India Tobacco "
     "Company' in 1970, then to 'I.T.C. Limited' in 1974 — deliberately stripping the "
     "tobacco-centric nomenclature to signal diversification intent.", sa=3)

sub_head(doc, "Strategic Diversification: Hotels, Paper & e-Choupal (1970–2000)")
body(doc,
     "The most consequential strategic decisions of ITC's post-Independence history were the entries "
     "into hotels (mid-1970s), paperboards (1979), and agribusiness-digital innovation (2000). "
     "ITC Hotels Limited was incorporated and opened its first luxury hotel in Chennai in 1975, "
     "with ITC Maurya (New Delhi) and ITC Windsor (Bangalore) following as flagship properties. "
     "The agribusiness deepening accelerated through the 1990s as ITC expanded its leaf tobacco "
     "procurement network. The transformative e-Choupal initiative launched in June 2000 — "
     "deploying internet kiosks in 35,000+ rural villages to provide farmers with real-time price "
     "information, weather data, and agronomic guidance, bypassing exploitative middlemen. "
     "e-Choupal became one of the most celebrated private-sector rural development interventions "
     "globally, the subject of Harvard Business School and INSEAD case studies.", sa=3)

sub_head(doc, "The FMCG Pivot (2000–2019)")
body(doc,
     "The early 2000s saw ITC make the most ambitious bet in its history: systematic entry into "
     "consumer goods categories. Aashirvaad atta launched in 2002, Sunfeast biscuits in 2003, "
     "and Bingo! snack foods in 2007. These were not opportunistic forays but a coordinated "
     "strategy leveraging ITC's cigarettes distribution infrastructure, e-Choupal procurement "
     "network, and Bhadrachalam paperboards division. For over a decade, FMCG-Others was "
     "EBIT-negative, drawing investor criticism that capital should be returned rather than "
     "deployed in loss-making consumer goods. This debate was the core of the long-running "
     "conglomerate discount — a discount that is now beginning to unwind as the FMCG thesis "
     "is vindicated.", sa=3)

sub_head(doc, "ITC Next & Hotels Demerger (2019–Present)")
body(doc,
     "Sanjiv Puri, who became Chairman & MD in May 2019, articulated the 'ITC Next' strategy — "
     "a comprehensive reset toward a consumer-centric, technology-enabled, climate-positive "
     "enterprise. The Hotels demerger, announced September 2023, received NCLT Kolkata approval "
     "on October 4, 2024, became effective January 1, 2025, and saw ITC Hotels Limited list on "
     "NSE/BSE on January 29, 2025. BAT reduced its ITC stake from ~29% to 22.9% through block "
     "trades in 2024–2025, partly funding its own GBP 2 billion share buyback. The January 2026 "
     "Central Excise Amendment Act imposed the largest single-step cigarette duty increase in "
     "India's history, creating the current investment opportunity as the market has over-reacted "
     "to cyclical regulatory pressure.", sa=3)

add_chart(doc, "chart_06_key_milestones_timeline.png", width=7.2,
          fig_num=11,
          caption="ITC Limited — Key Corporate Milestones Timeline (1910–2026)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# MANAGEMENT TEAM
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "MANAGEMENT TEAM",
            "ITC's Cadre of Career Executives: Deep Institutional Knowledge Across Five Segments")

body(doc,
     "ITC's management depth is exceptional by Indian corporate standards. Unlike the typical "
     "large Indian conglomerate where leadership is family-driven or heavily reliant on lateral "
     "external hires, ITC has developed a cadre of career executives who have spent 30–40 years "
     "within the ITC system — building institutional knowledge across businesses, understanding "
     "the nuances of regulated-industry management, and maintaining the long-term strategic "
     "discipline that the cigarettes-funding-FMCG model requires. This 'ITC lifer' culture is "
     "both a strength (deep knowledge, continuity) and a risk (potential insularity).", sa=3)

mgmt_bio(doc,
    "Sanjiv Puri", "Chairman & Managing Director",
    "Sanjiv Puri is among India's most accomplished corporate leaders, recognised with the Best CEO "
    "Award in FMCG at Fortune India's Best CEOs 2025 and the AIMA-JRD Tata Corporate Leadership Award. "
    "Born in 1963, Puri is a graduate of IIT Kanpur (India's most selective engineering institution) "
    "and holds a postgraduate degree from Wharton School of Business, University of Pennsylvania. He "
    "joined ITC in January 1986, beginning a four-decade career within a single company — a rarity "
    "in the modern era of executive mobility. His career trajectory has been notable for its breadth: "
    "early operational roles across ITC's businesses; MD of ITC Infotech (2006–2010); MD of Surya "
    "Nepal Private Limited, an international tobacco management role; Divisional Chief Executive of "
    "India Tobacco Division (from 2009); President of FMCG Businesses (2014); COO (2016–2017); "
    "CEO (2017); MD (2018); and Chairman & MD from May 13, 2019. As Chairman, Puri has presided "
    "over three landmark outcomes: FMCG-Others crossing EBIT profitability (FY23), the Hotels "
    "demerger (2025), and maintaining cigarettes resilience through unprecedented regulatory pressure. "
    "He serves as Immediate Past President of CII, Chairman of IIT Gandhinagar Board, member of "
    "BRICS Business Council, and holds an Honorary Doctorate from XIM University."
)

mgmt_bio(doc,
    "S. Sivakumar", "Group Head, Agri & IT Businesses",
    "S. Sivakumar is India's foremost practitioner of agribusiness innovation and rural development. "
    "A class topper of Institute of Rural Management, Anand (IRMA) in 1983, Sivakumar joined ITC in "
    "1989 after six years with a farmers' cooperative, grounding his career in smallholder agriculture. "
    "As Divisional Chief Executive of ITC's Agri Business Division from 1996 to 2017, he was the "
    "principal architect of e-Choupal (launched June 2000) — the world's largest internet-based "
    "rural intervention, covering 35,000+ villages across 10 states through ~6,100 kiosks, providing "
    "farmers with real-time price data, weather forecasts, and agronomic information. e-Choupal has "
    "been the subject of case studies at Harvard Business School and INSEAD and is cited globally as "
    "a template for private-sector rural development. In his current role as Group Head of Agri & IT "
    "Businesses, Sivakumar oversees both the Agri Business Division and ITC Infotech — an unusual "
    "portfolio reflecting his comfort with technology as a business enabler. He serves on the Advisory "
    "Council to the Ministry of Rural Development for the National Rural Livelihoods Mission and on "
    "SEBI's Commodity Derivatives Advisory Committee. The Agri Business segment under his stewardship "
    "delivered ~25% revenue growth in FY25, driven by strong leaf tobacco export demand."
)

mgmt_bio(doc,
    "Hemant Malik", "Wholetime Director, Foods Business Division",
    "Hemant Malik represents the commercial leadership backbone of ITC's FMCG transformation. An MBA "
    "graduate from IIM Kolkata — India's oldest and most prestigious management institution — Malik "
    "joined ITC in 1989, beginning a career spanning the full arc of ITC's consumer goods ambition. "
    "His progression through commercial functions took him through Tobacco, Lifestyle Retailing, and "
    "Packaged Foods before he was elevated to COO of Cigarette Brands and Supply Chain in August 2012 "
    "and Divisional Chief Executive of India Tobacco Division in August 2015. In October 2016, he "
    "transitioned to lead the Foods Business Division as its Divisional Chief Executive, presiding "
    "over its most critical phase: under his leadership, Aashirvaad atta crossed ₹7,000 Crore to "
    "become India's leading branded atta; Sunfeast consolidated a top-2 biscuit position; Bingo! "
    "became a major snacks player; and FMCG-Others crossed EBIT breakeven (FY23). Malik was "
    "appointed to ITC's Board as Wholetime Director effective August 12, 2023. His mandate is to "
    "reach ₹30,000+ Crore in FMCG-Others revenue by FY28 while improving EBIT margins from ~5% "
    "today to 7%+ — the most critical execution challenge in ITC's strategic plan."
)

mgmt_bio(doc,
    "Nakul Anand", "Formerly Executive Director, Hotels Division",
    "Nakul Anand's career at ITC spanned four-plus decades, nearly as long as the Hotels business "
    "itself. An Economics Honours graduate from Delhi University with an Advanced Management Programme "
    "from Bond University, Australia, Anand joined erstwhile ITC Hotels as a Management Trainee in "
    "1978. He rose to MD of ITC Hotels (2003–2005) and served as Executive Director on ITC's Board "
    "from January 3, 2011, overseeing Hospitality and Travel & Tourism. Over his tenure, he oversaw "
    "expansion across six hotel brands — ITC Hotels, WelcomHotel, Fortune, Storii, Mementos, and "
    "WelcomHeritage — and positioned ITC Hotels as one of India's most awarded luxury brands, "
    "renowned for LEED Platinum green-building certifications and culinary excellence. The Hotels "
    "demerger, which Anand's work helped structure, was completed in January 2025. Following "
    "demerger and his retirement from ITC, Anand joined Tricone Luxury Hotels as Chairman of its "
    "Advisory Board, taking his deep hospitality expertise into the broader industry."
)

add_chart(doc, "chart_07_organisational_structure.png", width=7.0,
          fig_num=12,
          caption="ITC Limited — Organisational Structure and Business Division Chart (March 2026)")

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────
# GOVERNANCE & OWNERSHIP
# ──────────────────────────────────────────────────────────────────────────
section_bar(doc, "GOVERNANCE & OWNERSHIP STRUCTURE",
            "Professional Management, Institutional Anchors, and BAT Stake Evolution")

body(doc,
     "ITC's governance structure is distinctive by Indian corporate standards. Unlike most large "
     "Indian companies where founding families or promoters hold controlling stakes and exercise "
     "operational influence, ITC has no founding family or controlling promoter group — it is a "
     "professionally managed, institutionally owned company. This structure provides management "
     "independence but has historically made rapid strategic pivots more challenging.", sa=3)

sub_head(doc, "Major Shareholders")
own_data = [
    ["Shareholder", "Category", "Stake (%)", "Notes"],
    ["British American Tobacco (BAT)", "Strategic Industrial Partner", "22.9%",
     "Reduced from ~29% via block trades in 2024–25; retains board representation"],
    ["Life Insurance Corporation of India (LIC)", "Government Financial Institution", "~16.1%",
     "Stable anchor; counterweight to BAT; long-term holding"],
    ["Domestic Mutual Funds", "Institutional", "~13–14%",
     "HDFC MF, Kotak MF, SBI MF leading holders"],
    ["Foreign Portfolio Investors (FPIs)", "Institutional", "~14–15%",
     "Gradual reduction on tobacco ESG concerns; some re-entry on valuation"],
    ["ITC Employee Trust / ESOPs", "Management / Employee", "~2%",
     "Alignment mechanism; modest"],
    ["Retail & Other Investors", "Public", "~26–28%",
     "Large retail base; significant holding via NPS / EPF / insurance funds"],
]
ot = doc.add_table(rows=len(own_data), cols=4)
ot.style = "Table Grid"
for ri, row in enumerate(own_data):
    for ci, val in enumerate(row):
        c = ot.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0),
                  color=WHITE if ri==0 else BLACK)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

body(doc,
     "BAT's role as ITC's strategic partner merits detailed analysis. BAT's historical ~29% "
     "stake provided ITC with access to global tobacco technology (manufacturing processes, "
     "blend development), global leaf procurement networks, and BAT's international cigarette "
     "brand expertise. Post the 2024–2025 stake reductions to 22.9%, BAT retains board "
     "representation but has relinquished its formal veto power under ITC's shareholder "
     "agreements — a significant governance improvement. BAT's continued presence (22.9% remains "
     "a strategic stake in any classification) means the strategic relationship continues, but "
     "the overhang risk of further selling persists. LIC's ~16% stake — the government's largest "
     "insurance and investment arm — provides a natural counterweight and represents an implicit "
     "institutional endorsement of ITC's long-term value.", sa=3)

body(doc,
     "Board composition (as of March 2026): 14 directors including Chairman & MD Sanjiv Puri, "
     "3 Executive Directors (incl. Hemant Malik, one additional Wholetime Director), 2 BAT nominee "
     "non-executive directors, 1 LIC nominee director, and 7 independent directors — broadly "
     "meeting Listing Obligations and Disclosure Requirements (LODR) governance standards. "
     "The absence of a family-promoter group is a distinctive governance feature; all major "
     "capital allocation decisions require Board approval and are subject to independent director "
     "scrutiny.", sa=4)

add_chart(doc, "chart_24_shareholding_pattern.png", width=6.5,
          fig_num=13,
          caption="ITC Limited — Shareholding Pattern (March 2026)")

add_chart(doc, "chart_27_ownership_structure_breakdown.png", width=6.5,
          fig_num=14,
          caption="ITC Limited — Ownership Structure: BAT Stake Evolution vs. Other Categories (FY20–FY26)")

doc.add_page_break()

doc.save(OUTPUT)
print(f"✅ Section 3 complete → {OUTPUT}")
