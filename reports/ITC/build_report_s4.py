"""
ITC Initiation Report — Section 4
Products & Services + Customers + Industry + Competitive Landscape + TAM
Charts: chart_08, chart_09, chart_15, chart_16, chart_17
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

NAVY  = RGBColor(0x0D, 0x2B, 0x55)
GREEN = RGBColor(0x1A, 0x75, 0x3F)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
TEAL  = RGBColor(0x0E, 0x6B, 0x6B)

def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic, color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing: {fname}"); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title, subtitle=None):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sub_head(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); sfont(r, size=11, bold=True, color=color or NAVY)

def sub2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); sfont(r, size=10, bold=True, color=TEAL)

def body(doc, text, size=10, sa=4, sb=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); sfont(r, size=size)

def bullet(doc, text, bold_prefix=None, indent=0.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent  = Inches(indent)
    if bold_prefix:
        r1 = p.add_run(f"▸  {bold_prefix}: ")
        sfont(r1, size=10, bold=True)
        r2 = p.add_run(text); sfont(r2, size=10)
    else:
        r = p.add_run(f"▸  {text}"); sfont(r, size=10)

# ─── OPEN DOC ──────────────────────────────────────────────────────────────
doc = Document(OUTPUT)

# ══════════════════════════════════════════════════════════════════════════
# PRODUCTS & SERVICES
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "PRODUCTS & SERVICES",
            "Five Business Verticals United by Distribution Scale and Commodity Self-Sufficiency")

body(doc,
     "ITC's product portfolio spans five verticals, united by the common threads of deep "
     "distribution scale, commodity procurement expertise via e-Choupal, and packaging "
     "self-sufficiency through ITC Bhadrachalam. Each vertical has distinct competitive dynamics "
     "and margin profiles, but the cigarettes business fundamentally cross-subsidises the brand "
     "investment required in FMCG-Others — the strategic rationale that defines ITC's conglomerate "
     "model.", sa=4)

add_chart(doc, "chart_08_product_portfolio_fmcg_others.png", width=7.0,
          fig_num=15,
          caption="ITC FMCG-Others — Product Portfolio: Category Leadership and Revenue Scale (FY25A)")

sub_head(doc, "1. FMCG-Cigarettes: The Cash Engine")
body(doc,
     "Cigarettes remain ITC's defining business — financially, strategically, and historically. "
     "ITC commands approximately 75–77% of India's legal cigarettes market by volume, a position "
     "that has been remarkably stable over three decades. The brand portfolio spans every price "
     "point in the legal market:", sa=2)

bullet(doc, "India's largest-selling cigarette brand by volume; spans premium to mid-premium "
       "segments; Gold Flake Kings and Gold Flake Premium are the volume leaders.",
       bold_prefix="Gold Flake (incl. Kings, Premium)")
bullet(doc, "Premium positioning; popular with urban consumers; Classic Milds has a strong "
       "loyal consumer base among light smokers.",
       bold_prefix="Classic (incl. Classic Milds, Regular)")
bullet(doc, "Heritage brand with strong urban loyalty; older smoker demographic with high "
       "repeat purchase rates.",
       bold_prefix="Wills (incl. Navy Cut)")
bullet(doc, "Value segment; maintains ITC's presence across all price points; critical for "
       "retaining consumers who might otherwise downtrade to bidis.",
       bold_prefix="Bristol")

body(doc,
     "\nThe economics of ITC's cigarettes business are extraordinary by any standard. Segment "
     "EBIT margins of 58–65% on net revenues (after excise) make this among the highest-margin "
     "legal consumer businesses globally — comparable only to luxury goods and branded spirits. "
     "The business requires minimal capital reinvestment (manufacturing assets are largely "
     "depreciated; the product is non-perishable), generates strong FCF, and has demonstrated "
     "consistent semi-annual price increase capability. The January 2026 excise hike represents "
     "an acute short-term test of this pricing power model; historical precedents strongly suggest "
     "ITC will absorb the hike with moderate volume impact.", sa=3)

sub_head(doc, "2. FMCG-Others: The Growth Engine")
body(doc,
     "The FMCG-Others business (₹21,982 Crore revenue, FY25A) is ITC's strategic growth engine "
     "and the most compelling long-term value creation opportunity. The portfolio spans multiple "
     "consumer categories, with staples and packaged foods constituting approximately 75% of "
     "revenues:", sa=2)

sub2(doc, "Staples & Packaged Foods (~75% of FMCG-Others)")
bullet(doc, "India's #1 branded atta with ₹7,000+ Crore annual sales; expanded into spices, "
       "salt, instant mixes, dairy, and poha. Distribution in 600,000+ outlets; leverages "
       "e-Choupal wheat procurement.",
       bold_prefix="Aashirvaad")
bullet(doc, "Top-2 biscuits brand; direct competitor to Britannia across Marie Light, Farmlite "
       "(healthy positioning), and cookie segments. Also pasta and noodles adjacent to Yippee!.",
       bold_prefix="Sunfeast")
bullet(doc, "India's #2 instant noodles; 25–30% market share, primarily gained from Maggi "
       "after the 2015 lead contamination controversy. Multi-format (round, roundwich) "
       "differentiates from Maggi's square format.",
       bold_prefix="Yippee!")
bullet(doc, "Fast-growing salted snacks and chips brand; genuine market-share gainer in a "
       "high-growth category competing with PepsiCo's Frito-Lay. Bingo! Mad Angles is a "
       "distinctive format with strong consumer recall.",
       bold_prefix="Bingo!")
bullet(doc, "Fruit juices and beverages; positioned in natural/no-concentrate segment; "
       "distribution leverage through ITC's network.",
       bold_prefix="B Natural")

sub2(doc, "Personal Care & Others (~25% of FMCG-Others)")
bullet(doc, "Deodorants and fragrances; competing in premium mass market vs. HUL's Axe and "
       "Marico's Set Wet.",
       bold_prefix="Engage")
bullet(doc, "Shower gels and soaps; mid-premium positioning; leverages ITC's distribution "
       "network for chemist and modern trade penetration.",
       bold_prefix="Fiama")
bullet(doc, "Value soaps and shampoos; broader geographic reach into Tier 2–4 cities and "
       "rural markets.",
       bold_prefix="Vivel")
bullet(doc, "India's largest stationery brand by volume; school and education segment; "
       "reliable cash-generating business with steady demand.",
       bold_prefix="Classmate")

sub_head(doc, "3. Paperboards, Paper & Packaging")
body(doc,
     "ITC Bhadrachalam (ITC's Paperboards and Specialty Papers Division, PSPD) is one of India's "
     "largest integrated paper mills, located on the banks of the Godavari River in Telangana. "
     "Capacity: ~677,000 tonnes per annum. Products include value-added paperboards for consumer "
     "goods packaging, specialty papers for printing/writing, and coated paperboards for food, "
     "beverage, and pharmaceutical packaging. Critically, PSPD supplies a significant portion of "
     "ITC's own FMCG and cigarettes packaging — an integrated supply chain advantage reducing "
     "exposure to third-party packaging price volatility. FY25 was difficult due to low-priced "
     "Chinese/Indonesian paper imports and elevated wood costs; recovery expected from FY26E as "
     "anti-dumping measures are implemented.", sa=3)

sub_head(doc, "4. Agribusiness & e-Choupal")
body(doc,
     "ITC's Agribusiness Division is organised around two complementary activities: procurement, "
     "processing, and export of agricultural commodities (leaf tobacco, wheat, spices, coffee, "
     "shrimp); and the e-Choupal rural platform. The Agri Business segment posted revenues of "
     "~₹19,753 Crore in FY25 (25% YoY growth), driven by strong international demand for Indian "
     "leaf tobacco. ITC is one of India's largest leaf tobacco exporters, participating in a "
     "global Indian-leaf export market estimated at USD 1+ billion annually. The e-Choupal "
     "network — covering 35,000+ villages through ~6,100 kiosks — is transitioning to version "
     "4.0, targeting 70,000 villages with integrated agri-services aggregation.", sa=3)

sub_head(doc, "5. Hotels (ITC Hotels Limited — Now Independently Listed)")
body(doc,
     "ITC Hotels Limited, demerged and listed on January 29, 2025, operates 140+ properties "
     "with ~13,300 rooms under six brands: ITC Hotels (flagship luxury, LEED Platinum certified); "
     "Mementos (distinctive luxury positioning); WelcomHotel (upper-upscale); Storii (boutique "
     "premium); Fortune (mid-market, largest by count); and WelcomHeritage (heritage properties). "
     "FY25 revenue: ₹3,333 Crore (44% growth); PAT: ₹698 Crore. Target: 220 properties by 2030 "
     "under an asset-light managed model. ITC retains a strategic 40% equity stake (~₹8,500 Crore "
     "by our estimate), sufficient to benefit from ITC Hotels' value creation without the capital "
     "intensity of the hotels business in ITC's books.", sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CUSTOMERS & GO-TO-MARKET
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "CUSTOMERS & GO-TO-MARKET",
            "A 2-Million-Outlet Proprietary Network: India's Most Powerful FMCG Distribution Asset")

body(doc,
     "ITC's most durable competitive advantage in consumer businesses is not brand equity alone — "
     "it is the proprietary Trade Marketing and Distribution (TM&D) network that reaches "
     "approximately two million retail points without reliance on third-party super-stockists or "
     "carrying-and-forwarding agents. This network, built over decades to serve the cigarettes "
     "business under COTPA advertising restrictions, is the single most important structural "
     "advantage ITC brings to its FMCG-Others launches.", sa=3)

add_chart(doc, "chart_09_customer_channel_segmentation.png", width=7.0,
          fig_num=16,
          caption="ITC — Customer & Channel Segmentation: Retail Reach, Trade Channel Mix, and FMCG Penetration (FY25A)")

sub_head(doc, "Cigarettes Distribution: The Direct-to-Retail Model")
body(doc,
     "The cigarettes business is served through one of India's most extensive direct distribution "
     "networks, reaching approximately 2 million retail outlets across urban, semi-urban, and rural "
     "India. Given the legal restrictions on tobacco advertising under COTPA (Cigarettes and Other "
     "Tobacco Products Act), distribution is the primary route to consumer — shelf availability "
     "and retailer relationship management are the critical commercial levers. ITC's TM&D vertical "
     "manages this network directly, with deep retailer-level data on sell-through, competitive "
     "shelf presence, and pricing execution. This direct relationship gives ITC unparalleled "
     "market intelligence and execution capability vs. competitors using traditional distribution "
     "hierarchies.", sa=3)

sub_head(doc, "FMCG-Others Distribution: Leveraging the Cigarettes Network")
body(doc,
     "ITC's most strategically intelligent decision in building the FMCG business was leveraging "
     "the cigarettes distribution network for packaged foods and personal care launches. Aashirvaad, "
     "Sunfeast, Bingo!, and Yippee! were distributed through the same TM&D network that had spent "
     "decades servicing cigarettes — dramatically reducing the distribution costs and timelines that "
     "typically constrain a new FMCG entrant. This meant ITC could launch new FMCG products with "
     "day-one distribution in 600,000–800,000+ retail outlets, a capability that typically takes a "
     "decade and billions of rupees for a startup brand to build. By comparison, new FMCG entrants "
     "typically spend 8–10% of revenues on distribution setup; ITC's incremental distribution cost "
     "for FMCG-Others is estimated at 2–3% of revenues — a structural cost advantage of 5–7 "
     "percentage points of margin vs. greenfield FMCG competitors.", sa=3)

body(doc,
     "ITC sells across all trade channels: General Trade (traditional kirana stores, pan shops, "
     "neighbourhood grocers — ~60% of FMCG-Others revenue); Modern Trade (supermarkets, "
     "hypermarkets, convenience stores — ~20%); and E-Commerce (Amazon, Flipkart, Blinkit, "
     "Swiggy Instamart, Zepto — ~8–10% and growing at 30%+ annually). The company is investing "
     "in direct-to-consumer digital capabilities through ITC Store and its own brand websites.", sa=3)

sub_head(doc, "Agribusiness Go-to-Market: B2B Direct Model")
body(doc,
     "ITC's Agribusiness operates a pure B2B model. The e-Choupal network enables direct sourcing "
     "from farmers, bypassing the APMC auction system for most commodities. On the sales side, "
     "leaf tobacco is exported directly to global tobacco companies (Universal Corporation, "
     "Standard Industries, Alliance One International) and international cigarette manufacturers. "
     "Wheat, spices, and coffee are sold to domestic food processors and international buyers. "
     "This direct procurement-to-trade model gives ITC significant quality control and cost "
     "advantages over commodity traders relying on mandis.", sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# INDUSTRY OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "INDUSTRY OVERVIEW",
            "Five Industries: FMCG, Cigarettes, Hospitality, Paper, and Agribusiness")

sub_head(doc, "Indian FMCG Market: ₹5.5–6 Lakh Crore and Growing at 9–13% p.a.")
body(doc,
     "The Indian Fast-Moving Consumer Goods market is one of the world's most attractive consumer "
     "arenas, with a current market size estimated at ₹5.5–6 lakh crore (USD 65–72 billion). "
     "The market is characterised by its bifurcation between an organised sector (branded products) "
     "and a large unorganised sector, with the former steadily gaining share through GST "
     "formalisation, digital payments, and growing consumer brand awareness. Growth expectations "
     "for the organised FMCG sector stand at 9–13% per annum over the medium term, with two "
     "distinct demand engines:", sa=2)

bullet(doc, "Urban India's premiumisation — consumers upgrading from commodity to branded, "
       "and from branded to premium branded; e-commerce accelerating in metro and Tier-1 cities")
bullet(doc, "Rural India's volume growth — driven by agricultural income, MGNREGS wages, "
       "PM Kisan transfers, expanding organised retail reach, and first-generation brand adoption")

body(doc,
     "\nIndia's per-capita FMCG consumption remains at approximately USD 50/year vs. USD 120–140 "
     "in China — a 2.5x–3x consumption gap that will progressively close over the next decade as "
     "income levels converge. India's median age of ~28 years, the expanding working-age "
     "population, accelerating urbanisation, and first-generation branded consumer creation "
     "provide powerful demographic tailwinds. The major bellwether companies — Hindustan Unilever "
     "(~₹60,000 Crore revenue), Nestle India, Britannia Industries, and Dabur India — provide "
     "the competitive reference framework for ITC's FMCG-Others ambitions.", sa=3)

sub_head(doc, "Indian Cigarettes Industry: Oligopoly Under Regulatory Pressure")
body(doc,
     "India's legal cigarettes market is among the world's most unique: a heavily regulated, "
     "high-taxation oligopoly where the dominant player (ITC, ~75–77% volume share) has "
     "maintained dominance over decades, yet where volumes have been broadly flat to declining "
     "due to aggressive excise taxation and health awareness. India's per-capita cigarette "
     "consumption is approximately 85–100 sticks per adult per year — among the lowest in Asia "
     "and globally — compared to 1,600+ in China and 600+ in Japan. This paradoxically low "
     "per-capita reflects decades of punitive excise policy targeting cigarettes while exempting "
     "bidis (hand-rolled tobacco at ~₹1/thousand sticks vs. ₹8,000+/thousand sticks for "
     "cigarettes post-January 2026). The illicit cigarettes market accounts for approximately "
     "26% of total cigarette consumption per the Tobacco Institute of India.", sa=3)

body(doc,
     "The competitive structure is a functional duopoly: ITC (~75–77% share), Godfrey Phillips "
     "India (~10% share, Philip Morris affiliate), and VST Industries (~9% share). This "
     "concentration means ITC effectively sets the reference price; competitor price hikes "
     "follow ITC's increases within weeks. The critical question for 2026–2027 is the degree "
     "to which the January 2026 excise hike accelerates substitution to illicit alternatives — "
     "our base case assumes modest and temporary volume impact, consistent with 30 years of "
     "historical precedent.", sa=3)

sub_head(doc, "Hospitality: Post-COVID Recovery and Structural Demand Growth")
body(doc,
     "India's hospitality industry has staged a strong post-COVID recovery, with RevPAR across "
     "premium and luxury segments recovering to and exceeding pre-pandemic levels by FY24–25. "
     "Structural demand drivers are strong: India's expanding business travel and MICE (meetings, "
     "incentives, conferences, exhibitions) market; rapid domestic leisure tourism growth; "
     "wedding and events demand; and inbound international tourism growing from a low base. "
     "Supply addition in the luxury segment has been slow due to high capital intensity, "
     "creating a sustained demand-supply imbalance that benefits existing premium inventory. "
     "GST at 18% on premium hotel rooms (above ₹7,500/night) remains a structural drag but "
     "has been broadly absorbed by the industry through pricing.", sa=3)

sub_head(doc, "Paper & Packaging: Near-Term Headwinds, Long-Term Tailwinds")
body(doc,
     "India's paper and packaging industry (₹1.5 lakh crore, growing at 6–8% p.a.) faces a "
     "near-term headwind from Chinese and Indonesian paper imports that have flooded domestic "
     "markets with below-cost pricing. ITC Bhadrachalam's FY25 EBIT declined ~34% as a result. "
     "Anti-dumping measures currently under review by the Ministry of Commerce would provide "
     "meaningful relief. Long-term, e-commerce-driven demand for corrugated packaging and "
     "sustainability-led shift from plastic to paper-based packaging are powerful secular "
     "tailwinds for ITC's value-added paperboards business.", sa=3)

add_chart(doc, "chart_15_market_size_tam_evolution.png", width=7.0,
          fig_num=17,
          caption="ITC's Addressable Markets: Indian FMCG TAM, Cigarettes Market, and Agribusiness Export TAM Evolution (₹ Lakh Crore)")

add_chart(doc, "chart_17_cigarette_market_share.png", width=6.5,
          fig_num=18,
          caption="India Legal Cigarettes Market — ITC vs. Godfrey Phillips vs. VST Industries: Volume Share Trend (FY15–FY26E)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "COMPETITIVE LANDSCAPE",
            "Dominant Oligopoly in Cigarettes; Challenger Position in FMCG")

body(doc,
     "ITC competes in distinct markets with entirely different competitive dynamics across "
     "its business segments. The conglomerate structure means that its cigarettes competitors "
     "(a two-player functional duopoly) are entirely different from its FMCG competitors "
     "(HUL-Britannia-Nestle oligopoly) or luxury hotels competitors (Taj, Oberoi). We analyse "
     "each competitive arena separately.", sa=4)

sub_head(doc, "Cigarettes: Near-Permanent Competitive Moat")
body(doc,
     "ITC's position in the Indian legal cigarettes market constitutes one of the most durable "
     "competitive moats in Indian consumer goods. With ~75–77% volume share maintained over "
     "three decades despite relentless regulatory pressure, ITC's cigarettes franchise has "
     "demonstrated extraordinary resilience. The moat has four components: (1) Distribution "
     "depth — 2 million direct retail relationships; (2) Brand portfolio depth — coverage of "
     "every price point from value to ultra-premium; (3) Manufacturing scale — ITC's four "
     "cigarette factories operate with scale efficiencies unavailable to smaller players; "
     "(4) Pricing leadership — in a duopoly, ITC sets prices and competitors follow.", sa=3)

body(doc,
     "Godfrey Phillips India (Philip Morris affiliate, ~10% share) is the only credible "
     "challenger. FY25 revenues were ₹5,611 Crore — roughly one-sixth of ITC's cigarettes "
     "revenue — with 39.6% growth. While Godfrey has been gaining marginal share in some "
     "sub-segments, it lacks the manufacturing scale and distribution depth to threaten ITC's "
     "dominance. VST Industries (~9% share, revenues ₹1,844 Crore FY25) is a distant third, "
     "primarily in South and East India. The structural competitor to legal cigarettes is the "
     "informal tobacco economy (bidis, illicit smuggled cigarettes) — not organised competitors.", sa=3)

sub_head(doc, "FMCG-Others: Gaining Share Against Entrenched Incumbents")
body(doc,
     "In FMCG-Others, ITC competes against India's most formidable consumer goods companies. "
     "The key competitors by category:", sa=2)

comp_data = [
    ["Competitor", "Revenues (₹ Cr, FY25E)", "Key Categories vs. ITC", "ITC's Position", "Assessment"],
    ["Hindustan Unilever (HUL)", "~60,000", "Personal care (Dove, Surf), foods (Knorr, Kissan)",
     "Challenger in PC, limited food overlap", "HUL's distribution & marketing scale is formidable; limited direct food overlap"],
    ["Nestle India", "~18,000", "Maggi noodles (vs. Yippee!), confectionery",
     "Gaining share — 25–30% noodles market", "ITC has structurally dented Maggi; Nestle rebuilding"],
    ["Britannia Industries", "~17,000", "Biscuits/cookies (vs. Sunfeast)",
     "Top-2 biscuits; gaining share", "Most direct competitor; both investing heavily in distribution"],
    ["Tata Consumer Products", "~16,000", "Salt, tea, spices (adjacent to staples)",
     "Competitor in adjacencies", "Fastest-growing; aggressive acquisitions; distribution expanding"],
    ["Dabur India", "~12,000", "Juices, health foods (vs. B Natural)",
     "Niche overlap in juices/health", "Dabur's Réal vs. B Natural; ITC has distribution edge"],
    ["ITC (FMCG-Others)", "~22,000", "All above categories",
     "Challenger gaining share", "Distribution moat; improving profitability; scale yet to be fully leveraged"],
]
ct = doc.add_table(rows=len(comp_data), cols=5)
ct.style = "Table Grid"
for ri, row in enumerate(comp_data):
    for ci, val in enumerate(row):
        c = ct.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0),
                  color=WHITE if ri==0 else BLACK)
sp2 = doc.add_paragraph()
sfont(sp2.add_run("Source: Company filings, Institutional Equity Research estimates."),
      size=7.5, italic=True, color=DGRAY)
sp2.paragraph_format.space_after = Pt(3)

body(doc,
     "The fundamental competitive assessment for FMCG-Others: ITC is a credible and improving "
     "challenger in most of its categories, with the distribution network as the key structural "
     "differentiator. HUL, Nestle, and Britannia are scale incumbents with deeper brand equity, "
     "but ITC's distribution cost advantage means it can sustain competitive brand-building at "
     "lower absolute A&P spend. Long-term success will depend on ITC building genuine independent "
     "brand equity — i.e., consumers choosing Sunfeast over Britannia based on taste preference "
     "and loyalty, not just distribution availability. The evidence from Aashirvaad's atta "
     "dominance and Yippee!'s structural share gain is encouraging.", sa=3)

add_chart(doc, "chart_16_competitive_positioning_matrix.png", width=7.0,
          fig_num=19,
          caption="ITC vs. FMCG Peers — Competitive Positioning Matrix: Revenue Scale vs. EBIT Margin (FY25A)")

sub_head(doc, "Hotels: Tier-1 Luxury Brand, Competition from Taj and Oberoi")
body(doc,
     "Indian Hotels Company (Taj Hotels, part of Tata Group, ~200+ properties) is the most "
     "directly comparable competitor to ITC Hotels in the luxury segment. Taj's heritage and "
     "international recognition (including New York and London flagships) give it a slight "
     "premium in international MICE and tourism. ITC Hotels competes effectively with Taj "
     "on sustainability (LEED Platinum certifications), culinary excellence, and prime "
     "locations of specific flagships (ITC Maurya in New Delhi, ITC Grand Chola in Chennai). "
     "EIH (Oberoi Hotels, ~30 ultra-luxury properties) commands the highest room rates in "
     "India but is not a direct volume competitor to ITC Hotels' broader portfolio.", sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# MARKET OPPORTUNITY / TAM
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "MARKET OPPORTUNITY & TOTAL ADDRESSABLE MARKET",
            "India's Fastest-Growing Large Economy Provides Multi-Decade Consumer TAM Tailwind")

body(doc,
     "The total addressable market for ITC's consumer businesses is among the largest in the "
     "world's fastest-growing large economy. ITC's strategic positioning across the value chain "
     "— from farmgate procurement (e-Choupal) through manufacturing (Bhadrachalam) to consumer "
     "distribution (2-million-outlet TM&D network) — creates durable structural advantages "
     "for capturing this opportunity.", sa=3)

tam_data = [
    ["Market Segment", "FY25 Market Size", "FY30E Market Size", "CAGR", "ITC's Addressable Share"],
    ["Organised FMCG (India)", "₹3.0–3.5 lakh crore", "₹5.0–5.5 lakh crore", "~10%",
     "~₹2.0–2.5 lakh crore (staples + snacks + personal care)"],
    ["Branded Atta / Staples", "₹25,000–30,000 crore", "₹45,000–55,000 crore", "~10%",
     "ITC #1 in branded atta (Aashirvaad); ~35% addressable share"],
    ["Biscuits & Cookies", "₹25,000–28,000 crore", "₹40,000–45,000 crore", "~8–9%",
     "Top-2 position (Sunfeast vs. Britannia); ~25–30% share"],
    ["Instant Noodles", "₹4,500–5,000 crore", "₹8,000–10,000 crore", "~12%",
     "~25–30% share (Yippee!); structural shift from Maggi"],
    ["Salted Snacks & Chips", "₹18,000–20,000 crore", "₹35,000–40,000 crore", "~13%",
     "~10–12% share (Bingo!); gaining in West/South India"],
    ["Legal Cigarettes (India)", "₹32,000–35,000 crore net", "₹38,000–42,000 crore", "~4–5%",
     "~75–77% volume share; value-led growth"],
    ["Leaf Tobacco Exports", "USD 1.0–1.2 billion", "USD 1.5–1.8 billion", "~8%",
     "ITC dominant buyer-exporter in India"],
    ["India Luxury Hospitality", "₹12,000–14,000 crore RevPAR", "₹22,000–26,000 crore", "~12–13%",
     "ITC Hotels: 40% stake in Tier-1 luxury brand; 220 props by FY30"],
]
tt = doc.add_table(rows=len(tam_data), cols=5)
tt.style = "Table Grid"
for ri, row in enumerate(tam_data):
    for ci, val in enumerate(row):
        c = tt.rows[ri].cells[ci]
        shade_cell(c, NAVY if ri==0 else (LGRAY if ri%2==0 else WHITE))
        cell_text(c, val, size=8, bold=(ri==0),
                  color=WHITE if ri==0 else BLACK)
sp3 = doc.add_paragraph()
sfont(sp3.add_run("Source: IMARC Group, ICICI Securities, company estimates, Institutional Equity Research."),
      size=7.5, italic=True, color=DGRAY)
sp3.paragraph_format.space_after = Pt(4)

body(doc,
     "ITC's current FMCG-Others revenue of ₹21,982 Crore represents approximately 9–11% "
     "penetration of its ~₹2.0–2.5 lakh crore addressable TAM. The path to ₹30,000 Crore "
     "by FY28 (management target) and ₹44,200 Crore by FY30E (our estimate) represents "
     "penetration of 12–18% of addressable TAM — still leaving extraordinary headroom for "
     "the following decade. India's per-capita FMCG consumption at USD 50/year vs. China's "
     "USD 120–140 implies a structural 2.5–3x consumption gap that will narrow progressively "
     "as India's per-capita GDP rises from ~USD 2,800 (FY26E) toward USD 4,000–5,000 by "
     "the early 2030s.", sa=4)

doc.add_page_break()

doc.save(OUTPUT)
print(f"✅ Section 4 complete → {OUTPUT}")
