"""
ITC Initiation Report — Section 5
Financial Analysis (Historical + Projections + Scenarios)
Charts: chart_03, chart_04, chart_10, chart_11, chart_12, chart_13,
        chart_19, chart_21, chart_22, chart_23, chart_25, chart_26
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

NAVY  = RGBColor(0x0D, 0x2B, 0x55)
GREEN = RGBColor(0x1A, 0x75, 0x3F)
RED   = RGBColor(0xC0, 0x39, 0x2B)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
TEAL  = RGBColor(0x0E, 0x6B, 0x6B)
LBLUE = RGBColor(0xD6, 0xE4, 0xF0)

def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=8.5, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic, color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing: {fname}"); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title, subtitle=None):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sub_head(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); sfont(r, size=11, bold=True, color=color or NAVY)

def sub2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); sfont(r, size=10, bold=True, color=TEAL)

def body(doc, text, size=10, sa=4, sb=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); sfont(r, size=size)

def fig_caption(doc, num, text):
    cp = doc.add_paragraph()
    sfont(cp.add_run(f"Figure {num} – {text}"), size=8.5, bold=True, color=NAVY)
    cp.paragraph_format.space_after = Pt(2)

def src_note(doc, text):
    p = doc.add_paragraph()
    sfont(p.add_run(text), size=7.5, italic=True, color=DGRAY)
    p.paragraph_format.space_after = Pt(4)

def make_fin_table(doc, headers, rows, hdr_fill=NAVY, alt_fill=LGRAY,
                   est_cols=None, size=8):
    """Build a financial table. est_cols = list of col indices that are estimates."""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]; shade_cell(c, hdr_fill)
        align = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
        cell_text(c, h, size=size, bold=True, color=WHITE, align=align)
    for ri, row in enumerate(rows):
        is_main = not row[0].startswith("  ")
        bg = alt_fill if ri % 2 == 0 else WHITE
        is_est = False
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            if est_cols and ci in est_cols:
                bg_col = LBLUE if is_main else RGBColor(0xE8,0xF4,0xFF)
                shade_cell(c, bg_col)
            else:
                shade_cell(c, bg)
            align = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_text(c, val, size=size, bold=is_main, align=align)
    return t

# ─── OPEN DOC ──────────────────────────────────────────────────────────────
doc = Document(OUTPUT)

# ══════════════════════════════════════════════════════════════════════════
# FINANCIAL ANALYSIS — HISTORICAL
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "FINANCIAL ANALYSIS — HISTORICAL PERFORMANCE",
            "Five-Year Track Record of Compounding Revenue, EBITDA, and Free Cash Flow (FY21A–FY25A)")

body(doc,
     "ITC's financial profile over FY21–FY25 is characterised by accelerating profitability "
     "despite external headwinds — including the COVID-19 disruption (FY21), commodity inflation "
     "(FY23), and the Hotels demerger restructuring (FY25). The compounding of the cigarettes "
     "cash engine has consistently funded ITC's FMCG-Others brand investments while growing "
     "the absolute dividend payout and net cash position. Net revenue has compounded at ~12% "
     "CAGR over FY21–FY25; EBITDA at ~14% CAGR; FCF at ~9% CAGR.", sa=3)

# Income Statement table (FY21–FY25 historical only, clean)
fig_caption(doc, 20, "ITC Limited — Income Statement Summary (FY21A–FY25A, ₹ Crore)")
is_hdr = ["Line Item", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A"]
is_data = [
    ["Gross Revenue from Operations",  "47,902", "59,214", "69,446", "71,985", "73,465"],
    ["Less: Excise Duty on Cigarettes", "(20,454)","(24,626)","(29,986)","(30,115)","(30,454)"],
    ["Net Revenue from Operations",    "27,448", "34,588", "39,460", "41,870", "43,011"],
    ["  YoY Growth (%)",               "—",      "26.1%",  "14.1%",  "6.1%",   "2.7%"],
    ["TOTAL Revenue (incl. other inc.","27,773", "35,018", "39,950", "42,380", "43,556"],
    ["Cost of Materials",              "(7,820)", "(10,240)","(12,180)","(12,680)","(12,950)"],
    ["Purchases of Stock-in-Trade",    "(4,110)", "(5,890)", "(6,540)", "(6,920)", "(7,100)"],
    ["Employee Costs",                 "(3,590)", "(4,110)", "(4,600)", "(5,000)", "(5,430)"],
    ["Other Operating Expenses",       "(3,680)", "(4,480)", "(5,110)", "(5,630)", "(6,170)"],
    ["EBITDA",                         "14,356", "17,458", "20,141", "21,857", "24,025"],
    ["  EBITDA Margin (on gross rev)", "34.7%",  "35.4%",  "35.5%",  "36.0%",  "34.7%"],
    ["Depreciation & Amortisation",    "(1,535)", "(1,726)", "(1,907)", "(2,077)", "(2,043)"],
    ["EBIT (Operating Profit)",        "12,821", "15,732", "18,234", "19,780", "21,982"],
    ["  EBIT Margin (%)",              "46.7%",  "45.5%",  "46.2%",  "47.2%",  "51.1%"],
    ["Other Income (Interest/Divs)",   "4,203",  "4,580",  "4,861",  "5,659",  "5,560"],
    ["Finance Costs",                  "(78)",   "(62)",   "(48)",   "(42)",   "(38)"],
    ["Profit Before Tax",              "17,024", "20,312", "23,095", "25,439", "27,542"],
    ["  PBT Margin (%)",               "62.0%",  "58.7%",  "58.5%",  "60.8%",  "64.0%"],
    ["Income Tax",                     "(3,993)","(5,064)","(4,342)","(5,545)","(6,230)"],
    ["  Effective Tax Rate (%)",       "23.5%",  "24.9%",  "18.8%",  "21.8%",  "22.6%"],
    ["PAT – Continuing Operations",    "13,031", "15,248", "18,753", "19,894", "21,312"],
    ["PAT – Discontinued (Hotels)",    "—",      "—",      "—",      "—",      "1,420"],
    ["Less: Minority Interest",        "(85)",   "(98)",   "(115)",  "(128)",  "(140)"],
    ["PAT Attributable to Shareholders","12,946","15,150", "18,638", "19,766", "22,592"],
    ["  Net Margin (continuing, %)",   "47.2%",  "43.8%",  "47.5%",  "47.2%",  "49.5%"],
]
make_fin_table(doc, is_hdr, is_data, size=8)
src_note(doc, "Source: ITC Annual Reports FY21–FY25, SEBI filings. Excise duty deducted per Indian GAAP (Ind AS).")

add_chart(doc, "chart_03_revenue_by_segment_stacked_area.png", width=7.0,
          fig_num=21,
          caption="ITC — Revenue by Business Segment (Stacked Area): FY19A–FY30E (₹ Crore)")

add_chart(doc, "chart_04_revenue_by_geography_stacked_bar.png", width=7.0,
          fig_num=22,
          caption="ITC — Revenue by Geography: Domestic vs. Exports (Stacked Bar, FY19A–FY30E)")

sub_head(doc, "Segment EBIT Analysis: Cigarettes Dominance and FMCG-Others Inflection")
body(doc,
     "Segment EBIT analysis reveals the stark profitability asymmetry that defines ITC and "
     "underlies the investment thesis. Cigarettes contributed ₹21,091 Crore of EBIT in FY25 — "
     "approximately 78% of total segment EBIT despite representing only 44% of gross revenues. "
     "The key inflection in FY25 is FMCG-Others EBIT reaching ₹1,050 Crore — modest in absolute "
     "terms but a landmark proof-of-concept after 15+ years of investment. Paperboards saw a "
     "significant EBIT decline in FY25 due to Chinese paper dumping, but this is cyclical rather "
     "than structural. Agribusiness delivered strong growth driven by leaf tobacco exports.",
     sa=3)

# Segment EBIT table
fig_caption(doc, 23, "ITC Limited — Segment EBIT Breakdown (FY21A–FY25A, ₹ Crore)")
seg_hdr = ["Segment", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A", "CAGR"]
seg_data = [
    ["Cigarettes EBIT",           "11,145","13,560","16,180","18,023","21,091","17.3%"],
    ["  Cigarettes EBIT Margin",  "—",     "—",     "~56%",  "~58%",  "~65%",  "—"],
    ["FMCG-Others EBIT",          "(410)", "(280)", "120",   "510",   "1,050", "n.m."],
    ["  FMCG-Others EBIT Margin", "—",     "—",     "~0.7%","~2.3%", "~4.8%", "—"],
    ["Agribusiness EBIT",         "520",   "840",   "840",   "730",   "850",   "13.1%"],
    ["Paperboards & Paper EBIT",  "740",   "1,060", "1,480", "1,360", "890",   "(4.5%)"],
    ["ITC Infotech & Others",     "230",   "295",   "340",   "420",   "490",   "20.8%"],
    ["Hotels EBIT (discontinued)","(80)",  "280",   "790",   "1,020", "210",   "—"],
    ["Unallocated Corp. Costs",   "(1,323)","(2,023)","(1,516)","(2,303)","(2,599)","—"],
    ["TOTAL Segment EBIT",        "10,822","13,732","18,234","19,760","21,982","19.3%"],
]
make_fin_table(doc, seg_hdr, seg_data, size=8)
src_note(doc, "Source: ITC Annual Reports FY21–FY25. Hotels treated as discontinued from FY25. CAGR over FY21–FY25.")

add_chart(doc, "chart_11_segment_ebit_margins.png", width=6.8,
          fig_num=24,
          caption="ITC — Segment EBIT Margins by Business Division: FY21A–FY30E")

add_chart(doc, "chart_10_margin_progression.png", width=6.8,
          fig_num=25,
          caption="ITC — Consolidated Margin Progression: Gross, EBITDA, EBIT, and Net Margins (FY21A–FY30E)")

doc.add_page_break()

sub_head(doc, "Cash Flow Analysis: Extraordinary FCF Generation")
body(doc,
     "ITC's cash flow profile is arguably its most compelling financial characteristic. Free cash "
     "flow (operating cash flow minus capex) has compounded from ₹12,040 Crore in FY21 to "
     "₹16,900 Crore in FY25 — an FCF yield of approximately 4.4% on current market cap. "
     "ITC's CFO-to-EBITDA conversion is consistently 90–95% (exceptionally high for an FMCG "
     "conglomerate), reflecting the cigarettes business's near-zero working capital requirements "
     "and minimal debtor days. Treasury investments (purchases of mutual funds and fixed-income "
     "instruments from surplus cash) dominate the investing activities, creating a ₹51,000 Crore+ "
     "treasury portfolio that generates ₹5,000–6,000 Crore annually in other income — "
     "effectively an additional earnings layer on top of operating performance.", sa=3)

fig_caption(doc, 26, "ITC Limited — Cash Flow Statement Summary (FY21A–FY25A, ₹ Crore)")
cf_hdr = ["Line Item", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A"]
cf_data = [
    ["A. OPERATING ACTIVITIES",       "",       "",       "",       "",       ""],
    ["  Net Profit (Continuing Ops.)", "13,031", "15,248", "18,753", "19,894", "21,312"],
    ["  Depreciation & Amortisation", "1,535",  "1,726",  "1,907",  "2,077",  "2,043"],
    ["  Profit on Investments (neg.)", "(1,820)","(2,100)","(2,210)","(2,980)","(2,800)"],
    ["  Working Capital Changes",      "(300)",  "(720)",  "(530)",  "(390)",  "(240)"],
    ["  Income Tax Paid",              "(3,850)","(4,900)","(4,250)","(5,420)","(6,080)"],
    ["NET CFO",                        "14,520", "16,180", "18,820", "19,510", "20,480"],
    ["  CFO / Net Revenue (%)",        "52.9%",  "46.8%",  "47.7%",  "46.6%",  "47.6%"],
    ["B. INVESTING ACTIVITIES",        "",       "",       "",       "",       ""],
    ["  Capital Expenditure",          "(2,480)","(2,890)","(3,190)","(3,410)","(3,580)"],
    ["  Net Treasury Investments",     "1,900",  "1,400",  "1,500",  "1,400",  "1,200"],
    ["  Interest + Dividends Rcvd",    "1,310",  "1,460",  "1,630",  "1,830",  "2,020"],
    ["NET CFI",                        "730",    "(30)",   "(60)",   "(180)",  "440"],
    ["C. FINANCING ACTIVITIES",        "",       "",       "",       "",       ""],
    ["  Dividends Paid",               "(7,218)","(7,841)","(8,160)","(9,424)","(10,684)"],
    ["  Share Buyback",                "—",      "(3,500)","—",      "—",      "—"],
    ["NET CFF",                        "(7,038)","(11,431)","(8,240)","(9,494)","(10,744)"],
    ["D. FREE CASH FLOW (CFO–CapEx)",  "12,040", "13,290", "15,630", "16,100", "16,900"],
    ["  FCF Yield (on current mcap)",  "3.1%",   "3.4%",   "4.0%",   "4.2%",   "4.4%"],
    ["  FCF / PAT (%)",               "92.4%",  "87.2%",  "83.4%",  "80.9%",  "74.3%"],
]
make_fin_table(doc, cf_hdr, cf_data, size=8)
src_note(doc, "Source: ITC Annual Reports FY21–FY25.  Net treasury: sale proceeds minus purchases of investments (simplified).")

add_chart(doc, "chart_12_free_cash_flow_trend.png", width=6.8,
          fig_num=27,
          caption="ITC — Free Cash Flow Generation: FCF, CFO, and CapEx Trend (FY21A–FY30E)")

sub_head(doc, "Balance Sheet: A Net-Cash Fortress")
body(doc,
     "ITC's balance sheet is among the strongest of any large-cap Indian company. Key features: "
     "(1) Net cash position of ~₹28,000 Crore (FY25A), growing to ₹61,789 Crore by FY27E as "
     "FCF accumulates above dividend payouts; (2) Treasury investment portfolio of ₹41,000+ "
     "Crore (FY25A) in highly liquid fixed-income instruments — ITC essentially runs a "
     "quasi-treasury function; (3) ROCE of 36.8% and ROE of 27.3% in FY25 — exceptional for a "
     "diversified conglomerate; (4) Total debt of only ₹363 Crore (FY25A) vs. EBITDA of "
     "₹24,025 Crore — effectively zero leverage.", sa=3)

fig_caption(doc, 28, "ITC Limited — Balance Sheet Summary (FY21A–FY25A, ₹ Crore)")
bs_hdr = ["Line Item", "FY21A", "FY22A", "FY23A", "FY24A", "FY25A"]
bs_data = [
    ["ASSETS",                             "",       "",       "",       "",       ""],
    ["  Cash & Equivalents",               "3,702",  "4,501",  "9,920",  "10,241", "10,517"],
    ["  Current Investments (Treasury)",   "12,450", "16,400", "21,501", "31,400", "41,000"],
    ["  Trade Receivables",                "2,180",  "2,830",  "3,310",  "3,670",  "3,960"],
    ["  Inventories",                      "3,420",  "3,900",  "4,210",  "4,410",  "4,590"],
    ["  Other Current Assets",             "2,710",  "2,990",  "3,300",  "3,600",  "3,850"],
    ["TOTAL CURRENT ASSETS",               "24,462", "30,621", "42,241", "53,321", "63,917"],
    ["  Net PP&E",                         "16,300", "16,500", "16,800", "16,800", "16,500"],
    ["  CWIP + Other Non-Current",         "8,520",  "9,420",  "10,410", "11,350", "12,060"],
    ["TOTAL NON-CURRENT ASSETS",           "24,820", "25,920", "27,210", "28,150", "28,560"],
    ["TOTAL ASSETS",                       "49,282", "56,541", "69,451", "81,471", "92,477"],
    ["LIABILITIES + EQUITY",               "",       "",       "",       "",       ""],
    ["  Total Borrowings",                 "580",    "530",    "470",    "410",    "350"],
    ["  Trade Payables",                   "4,820",  "5,690",  "6,540",  "7,010",  "7,560"],
    ["  Other Liabilities",                "5,740",  "6,450",  "7,250",  "7,880",  "8,460"],
    ["TOTAL LIABILITIES",                  "11,140", "12,670", "14,260", "15,300", "16,370"],
    ["  Share Capital (₹1 FV)",            "1,254",  "1,254",  "1,255",  "1,256",  "1,257"],
    ["  Reserves & Surplus",               "35,560", "41,120", "52,350", "63,000", "72,650"],
    ["  Minority Interest",                "1,328",  "1,497",  "1,586",  "1,915",  "2,200"],
    ["TOTAL EQUITY",                       "38,142", "43,871", "55,191", "66,171", "76,107"],
    ["Net Cash (Cash + Inv. – Debt)",      "15,572", "20,371", "30,951", "41,231", "51,167"],
    ["ROCE (%)",                           "30.8%",  "34.1%",  "37.6%",  "36.5%",  "36.8%"],
    ["ROE (%)",                            "34.5%",  "35.4%",  "34.8%",  "30.2%",  "27.3%"],
]
make_fin_table(doc, bs_hdr, bs_data, size=8)
src_note(doc, "Source: ITC Annual Reports FY21–FY25.  Net cash includes cash + current investments – total borrowings.")

add_chart(doc, "chart_22_balance_sheet_composition.png", width=6.5,
          fig_num=29,
          caption="ITC — Balance Sheet Composition: Asset Mix and Capital Structure (FY21A–FY30E)")

add_chart(doc, "chart_23_capex_vs_net_cash_buildup.png", width=6.5,
          fig_num=30,
          caption="ITC — CapEx vs. Net Cash Build-Up: Capital Allocation Analysis (FY21A–FY30E)")

add_chart(doc, "chart_26_roce_roe_vs_wacc.png", width=6.5,
          fig_num=31,
          caption="ITC — ROCE, ROE vs. WACC: Economic Value Creation Analysis (FY21A–FY30E)")

add_chart(doc, "chart_19_dividend_history_payout.png", width=6.5,
          fig_num=32,
          caption="ITC — Dividend Per Share & Payout Ratio: 18 Consecutive Years of Dividend Growth (FY08–FY30E)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# FINANCIAL PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "FINANCIAL PROJECTIONS — FY26E TO FY30E",
            "Base Case: 9.3% Net Revenue CAGR, EBITDA Growing to ₹39,750 Crore, FCF ₹28,668 Crore by FY30E")

body(doc,
     "Our financial projections for ITC over FY26E–FY30E are anchored in segment-level revenue "
     "and margin assumptions that reflect our investment thesis. The key assumptions are: "
     "(1) cigarettes volume decline of 2–3% p.a. offset by 7–10% price increases, preserving "
     "segment EBIT growth of ~8% CAGR; (2) FMCG-Others revenue CAGR of 14% with EBIT margin "
     "expansion from ~4.8% (FY25A) to 7.0% (FY30E) as scale leverage accrues; (3) Agribusiness "
     "revenue CAGR of ~9% driven by leaf tobacco exports and e-Choupal 4.0; (4) Paperboards "
     "recovery from FY27E following anti-dumping relief; (5) Net cash compounding from "
     "₹28,000 Crore (FY25A) to ₹80,000+ Crore (FY30E) as FCF exceeds the dividend payout.",
     sa=3)

fig_caption(doc, 33, "ITC Limited — Projected Income Statement (FY26E–FY30E, ₹ Crore)")
proj_hdr = ["Line Item", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E", "CAGR"]
proj_data = [
    ["Gross Revenue",           "77,640", "83,580", "90,020", "97,100", "1,05,000","7.4%"],
    ["Net Revenue",             "47,310", "52,140", "57,360", "62,100", "67,200",  "9.3%"],
    ["  YoY Growth (%)",        "10.0%",  "10.2%",  "10.0%",  "8.3%",   "8.2%",   "—"],
    ["COGS + Employee + OpEx",  "(20,640)","(22,720)","(24,860)","(26,160)","(27,450)","—"],
    ["EBITDA",                  "26,670", "29,420", "32,500", "35,940", "39,750",  "10.8%"],
    ["  EBITDA Margin (gross)", "37.3%",  "37.6%",  "37.8%",  "38.3%",  "38.7%",  "—"],
    ["D&A",                     "(2,020)","(2,220)","(2,300)","(2,440)","(2,590)", "—"],
    ["EBIT",                    "24,650", "27,200", "30,200", "33,500", "37,160",  "10.8%"],
    ["Other Income",            "5,800",  "6,100",  "6,440",  "6,820",  "7,240",  "—"],
    ["Finance Costs",           "(40)",   "(38)",   "(36)",   "(34)",   "(32)",   "—"],
    ["Profit Before Tax",       "30,410", "33,262", "36,604", "40,286", "44,368", "10.0%"],
    ["Income Tax (24.5%)",      "(7,300)","(7,970)","(8,800)","(9,680)","(10,650)","—"],
    ["PAT – Continuing Ops",    "23,110", "25,292", "27,804", "30,606", "33,718", "9.9%"],
    ["Less: Minority Int.",     "(155)",  "(172)",  "(190)",  "(210)",  "(230)",  "—"],
    ["PAT – Shareholders",      "22,955", "25,120", "27,614", "30,396", "33,488", "9.9%"],
    ["  Net Margin (%)",        "48.5%",  "48.2%",  "48.1%",  "48.9%",  "49.8%", "—"],
    ["EPS (₹)",                 "18.26",  "19.98",  "21.97",  "24.18",  "26.64", "9.9%"],
    ["DPS (₹)",                 "15.50",  "17.00",  "18.65",  "20.55",  "22.65", "10.0%"],
    ["Free Cash Flow",          "18,880", "20,882", "23,124", "25,866", "28,668", "11.0%"],
    ["  FCF Yield (at CMP)",    "4.9%",   "5.4%",   "6.0%",   "6.7%",   "7.4%", "—"],
]
make_fin_table(doc, proj_hdr, proj_data, est_cols=[1,2,3,4,5], size=8)
src_note(doc, "Source: Institutional Equity Research estimates.  Blue shading = estimates.  "
              "EBITDA margin % on gross revenue.  CAGR: FY25A–FY30E.")

sub_head(doc, "Segment Revenue Build: FY26E–FY30E Assumptions")
body(doc,
     "The revenue build across segments reflects divergent growth trajectories. Cigarettes "
     "grows at ~5–7% net revenue CAGR (premiumisation offsetting modest volume decline). "
     "FMCG-Others, the strategic growth driver, grows at 14% CAGR from ₹21,982 Crore (FY25A) "
     "to ₹44,200 Crore (FY30E) — approaching management's ₹30,000 Crore FY28 aspiration. "
     "Agribusiness grows at ~9% CAGR; Paperboards recovers from FY27E at ~7% CAGR.", sa=2)

seg_proj_data = [
    ["Segment", "FY25A", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E", "CAGR"],
    ["Cigarettes (net rev)",  "19,855","21,030","22,490","24,060","25,740","27,540","6.8%"],
    ["  Cig EBIT Margin",     "~65%",  "75.6%", "76.5%", "77.5%", "78.5%", "79.5%", "—"],
    ["FMCG-Others Revenue",   "21,982","24,760","28,580","33,050","38,200","44,200","15.0%"],
    ["  FMCG-Others EBIT",    "1,050", "1,760", "2,870", "4,130", "5,610", "7,210", "47.2%"],
    ["  FMCG-Others EBIT Mg", "4.8%",  "7.1%",  "10.0%", "12.5%", "14.7%", "16.3%","—"],
    ["Agribusiness Revenue",  "19,753","20,320","22,430","24,720","27,190","29,930","8.7%"],
    ["  Agri EBIT",           "850",   "960",   "1,080", "1,210", "1,360", "1,520", "12.3%"],
    ["Paperboards Revenue",   "8,200", "8,900", "9,600", "10,400","11,300","12,300","8.4%"],
    ["  Paperboards EBIT",    "890",   "1,050", "1,230", "1,420", "1,640", "1,890", "16.3%"],
    ["ITC Infotech Revenue",  "3,480", "3,920", "4,410", "4,960", "5,580", "6,280", "12.5%"],
    ["  Infotech EBIT",       "490",   "580",   "660",   "760",   "870",   "990",   "15.1%"],
    ["Consol. EBIT",          "21,982","24,650","27,200","30,200","33,500","37,160","11.1%"],
]
make_fin_table(doc, seg_proj_data[0], seg_proj_data[1:], est_cols=[2,3,4,5,6], size=8)
src_note(doc, "Source: Institutional Equity Research estimates.  FMCG-Others EBIT margin % is on segment revenue basis.")

add_chart(doc, "chart_21_segment_ebit_waterfall_fy25.png", width=7.0,
          fig_num=34,
          caption="ITC — Segment EBIT Waterfall (FY25A vs. FY30E): Contribution by Division")

add_chart(doc, "chart_13_operating_metrics_dashboard.png", width=7.0,
          fig_num=35,
          caption="ITC — Key Operating Metrics Dashboard: EPS, FCF/Share, DPS, ROCE, ROE, and Net Cash/Share (FY21A–FY30E)")

add_chart(doc, "chart_25_working_capital_efficiency.png", width=6.5,
          fig_num=36,
          caption="ITC — Working Capital Efficiency: Receivable Days, Payable Days, and Cash Conversion Cycle (FY21A–FY30E)")

doc.add_page_break()

sub_head(doc, "Scenario Analysis: Bull / Base / Bear Cases (FY30E)")
body(doc,
     "We present three scenarios for ITC's FY30E financial profile, reflecting different "
     "outcomes for the two critical variables: cigarettes volume trajectory (driven by excise "
     "duty absorption) and FMCG-Others margin expansion (driven by competitive dynamics and "
     "scale leverage). Our base case probability weighting: Bull 25%, Base 55%, Bear 20%.", sa=2)

scen_data = [
    ["Assumption", "BULL CASE (25%)", "BASE CASE (55%)", "BEAR CASE (20%)"],
    ["Cigarettes Vol. CAGR",     "+1% to +2% p.a.",   "-2% to -3% p.a.",  "-4% to -6% p.a."],
    ["Cigarettes Net Rev CAGR",  "9.0% (FY25–FY30)",  "6.8%",             "2.5%"],
    ["FMCG-Others Rev CAGR",     "20.0%",             "14.0%",            "8.0%"],
    ["FMCG-Others EBIT Margin",  "11% by FY30E",      "7.0% by FY30E",    "3% by FY30E"],
    ["Consolidated EBITDA Mgn",  "42.0%",             "38.7%",            "29.0%"],
    ["Total Net Revenue (FY30E)","₹78,400 Cr",        "₹67,200 Cr",       "₹54,200 Cr"],
    ["EBITDA (FY30E)",           "₹32,930 Cr",        "₹24,523 Cr",       "₹15,718 Cr"],
    ["PAT (FY30E)",              "₹23,900 Cr",        "₹17,800 Cr",       "₹10,840 Cr"],
    ["EPS (FY30E)",              "₹19.01",            "₹14.16",           "₹8.63"],
    ["FCF (FY30E)",              "₹18,930 Cr",        "₹15,700 Cr",       "₹10,260 Cr"],
    ["DPS (FY30E)",              "₹17.50",            "₹12.00",           "₹6.75"],
    ["Implied Price Target",     "₹480",              "₹380",             "₹240"],
    ["Probability",              "25%",               "55%",              "20%"],
    ["Prob.-Weighted PT",        "₹120",              "₹209",             "₹48 → ₹377"],
]
st2 = doc.add_table(rows=len(scen_data), cols=4)
st2.style = "Table Grid"; st2.alignment = WD_TABLE_ALIGNMENT.CENTER
bull_col = RGBColor(0xD5, 0xF5, 0xE3)
bear_col = RGBColor(0xFD, 0xED, 0xEC)
for ri, row in enumerate(scen_data):
    for ci, val in enumerate(row):
        c = st2.rows[ri].cells[ci]
        if ri == 0:
            shade_cell(c, NAVY)
            cell_text(c, val, size=9, bold=True, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            if ci == 1: shade_cell(c, bull_col)
            elif ci == 3: shade_cell(c, bear_col)
            else: shade_cell(c, LGRAY if ri%2==0 else WHITE)
            cell_text(c, val, size=8.5,
                      bold=(row[0] in ["Implied Price Target","Probability","Prob.-Weighted PT"]),
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
src_note(doc, "Source: Institutional Equity Research estimates.  Bull = green; Bear = red; Base = grey.")

doc.add_page_break()

doc.save(OUTPUT)
print(f"✅ Section 5 complete → {OUTPUT}")
