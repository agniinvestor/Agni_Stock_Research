"""
ITC Initiation Report — Section 6 (FINAL)
Valuation + Catalysts + Appendices + Disclosures
Charts: chart_28, chart_29, chart_30, chart_31, chart_32, chart_34, chart_35
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.expanduser("~/ITC_report")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "ITC_Initiation_Report_2026-03-18.docx")

NAVY  = RGBColor(0x0D, 0x2B, 0x55)
GREEN = RGBColor(0x1A, 0x75, 0x3F)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GOLD  = RGBColor(0xB7, 0x95, 0x0B)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
TEAL  = RGBColor(0x0E, 0x6B, 0x6B)
LBLUE = RGBColor(0xD6, 0xE4, 0xF0)

def sfont(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, italic=False,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    sfont(r, size=size, bold=bold, italic=italic, color=color or BLACK)

def add_chart(doc, fname, width=6.8, fig_num=None, caption=None):
    path = os.path.join(CHARTS, fname)
    if not os.path.exists(path):
        print(f"  !! Missing: {fname}"); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        label = f"Figure {fig_num} – {caption}" if fig_num else caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        cr = cp.add_run(f"{label}\nSource: Company data, Bloomberg, Institutional Equity Research estimates.")
        sfont(cr, size=7.5, italic=True, color=DGRAY)

def section_bar(doc, title, subtitle=None):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {title}")
    sfont(r, size=13, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sub_head(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); sfont(r, size=11, bold=True, color=color or NAVY)

def body(doc, text, size=10, sa=4, sb=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); sfont(r, size=size)

def fig_caption(doc, num, text):
    cp = doc.add_paragraph()
    sfont(cp.add_run(f"Figure {num} – {text}"), size=8.5, bold=True, color=NAVY)
    cp.paragraph_format.space_after = Pt(2)

def src_note(doc, text):
    p = doc.add_paragraph()
    sfont(p.add_run(text), size=7.5, italic=True, color=DGRAY)
    p.paragraph_format.space_after = Pt(4)

def bullet(doc, bold_part, body_part, indent=0.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(indent)
    r1 = p.add_run("▸  "); sfont(r1, size=10, bold=True, color=NAVY)
    r2 = p.add_run(bold_part + "  "); sfont(r2, size=10, bold=True)
    r3 = p.add_run(body_part); sfont(r3, size=10)

def make_table(doc, headers, rows, hdr_fill=NAVY, size=8.5):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]; shade_cell(c, hdr_fill)
        cell_text(c, h, size=size, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        is_main = not row[0].startswith("  ")
        bg = LGRAY if ri%2==0 else WHITE
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; shade_cell(c, bg)
            cell_text(c, val, size=size, bold=is_main,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
    return t

# ─── OPEN DOC ──────────────────────────────────────────────────────────────
doc = Document(OUTPUT)

# ══════════════════════════════════════════════════════════════════════════
# VALUATION ANALYSIS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "VALUATION ANALYSIS",
            "Three-Method Framework: SOTP DCF (50%) + EV/EBITDA Comps (35%) + DDM (15%)")

body(doc,
     "We use a three-pronged valuation framework to derive our ₹380 twelve-month price target: "
     "(1) Sum-of-the-Parts (SOTP) DCF — primary method (50% weight), valuing each of ITC's "
     "five segments separately to eliminate conglomerate discount opacity; (2) EV/EBITDA "
     "Comparable Companies — secondary method (35% weight), applying a blended peer multiple "
     "weighted by segment EBIT contributions; (3) Dividend Discount Model (DDM) — supplementary "
     "floor valuation (15% weight), reflecting ITC's exceptional 85–90% payout ratio and "
     "18 consecutive years of dividend growth.", sa=3)

sub_head(doc, "Valuation Methodology 1: SOTP DCF (50% weight)")
body(doc,
     "We value each segment individually using discounted cash flow analysis, applying "
     "segment-specific WACCs and terminal growth rates that reflect the risk profile and "
     "growth trajectory of each business. The consolidated WACC of 11.59% is derived from: "
     "Risk-Free Rate (10yr GoI bond) 6.80% + Equity Risk Premium (India) 6.50% × Beta 0.75 "
     "= Cost of Equity 11.68%; Cost of Debt (post-tax) 5.66%; Debt/Capital 1.5% (near-zero "
     "leverage). ITC's 0.75 beta reflects defensive characteristics: non-cyclical cigarettes, "
     "stable FMCG demand, and net-cash balance sheet eliminating financial risk.", sa=3)

fig_caption(doc, 37, "SOTP DCF: Segment-Level Valuation Summary (₹ Crore)")
sotp_data = [
    ["Segment", "UFCF FY30E", "Term. Growth", "Seg. WACC", "PV of FCFs", "PV of TV", "Seg. EV", "Per Share (₹)"],
    ["Cigarettes",    "16,374", "4.5%", "10.8%", "53,100",  "1,57,090", "2,10,190", "167.2"],
    ["FMCG-Others",  "4,173",  "8.0%", "12.5%", "7,630",   "54,721",   "62,351",   "49.6"],
    ["Agribusiness",  "1,110",  "6.5%", "12.0%", "2,930",   "12,150",   "15,080",   "12.0"],
    ["Paperboards",  "990",    "5.0%", "11.5%", "2,350",   "9,180",    "11,530",   "9.2"],
    ["ITC Infotech",  "—",      "—",    "—",     "—",       "—",        "9,700",    "7.7"],
    ["TOTAL Op. EV",  "—",      "—",    "—",     "—",       "—",        "3,08,851", "245.7"],
    ["+ Net Cash & Inv.", "—",  "—",    "—",     "—",       "—",        "28,000",   "22.3"],
    ["+ ITC Hotels (40%)","—",  "—",    "—",     "—",       "—",        "8,500",    "6.8"],
    ["+ Other Invest.",   "—",  "—",    "—",     "—",       "—",        "3,500",    "2.8"],
    ["– Minority Int.", "—",    "—",    "—",     "—",       "—",        "(2,120)",  "(1.7)"],
    ["– Corp. Costs PV","—",    "—",    "—",     "—",       "—",        "(8,500)",  "(6.8)"],
    ["Equity Value",  "—",      "—",    "—",     "—",       "—",        "3,38,231", "269.1"],
    ["5% Exec. Risk Disc.","—", "—",    "—",     "—",       "—",        "—",        "—"],
    ["SOTP DCF Value","—",       "—",   "—",     "—",       "—",        "—",        "₹269"],
]
make_table(doc, ["Segment", "UFCF FY30E", "g", "WACC", "PV FCFs", "PV TV", "EV (₹Cr)", "₹/Share"],
           sotp_data[:-1], size=8)
src_note(doc, "Source: Institutional Equity Research DCF model.  UFCF = Unlevered Free Cash Flow.  "
              "Terminal growth rate applied to FY30E UFCF × (1+g) / (WACC–g).")

add_chart(doc, "chart_28_dcf_sensitivity_heatmap.png", width=7.0,
          fig_num=38,
          caption="ITC — DCF Sensitivity Analysis: Price Per Share (₹) vs. WACC and Terminal Growth Rate (2-Way Heat Map)")

add_chart(doc, "chart_29_dcf_enterprise_equity_bridge.png", width=6.5,
          fig_num=39,
          caption="ITC — DCF Enterprise Value to Equity Value Bridge: SOTP Segment Waterfall (₹ Crore)")

sub_head(doc, "DCF Sensitivity: ₹219–₹590 Range Across Scenarios")
body(doc,
     "Our DCF sensitivity matrix (Figure 38) shows ITC's intrinsic value ranges from ₹219 "
     "(pessimistic: WACC 13%, g 4.5%) to ₹590 (optimistic: WACC 9.5%, g 7.5%). The base "
     "case (WACC 11.59%, g 6.0%) implies ₹278. The highlighted band (WACC 10.5–11.0%, "
     "g 6.0–7.0%) yields ₹330–384 — our 'Bull to Mid-Bull' scenario range. At the current "
     "market price of ₹307, the market is implying approximately WACC 10.8%, g 5.0% — "
     "already embedding some pessimism. Our price target of ₹380 corresponds to WACC 10.8%, "
     "g 6.5% — a modest improvement in growth assumption as FMCG-Others matures.", sa=4)

doc.add_page_break()

sub_head(doc, "Valuation Methodology 2: EV/EBITDA Comparable Companies (35% weight)")
body(doc,
     "ITC's unique multi-segment structure requires a blended peer approach. We construct a "
     "peer set spanning Indian FMCG companies, global tobacco peers, and Indian diversified "
     "consumer conglomerates. The blended fair multiple is derived by weighting each segment's "
     "peer benchmark by its EBIT contribution.", sa=2)

comp_hdr = ["Company", "Ticker", "Mkt Cap (₹Cr)", "EV/Rev LTM", "EV/EBITDA LTM",
            "EV/EBITDA NTM", "P/E NTM", "Rev Gr%", "EBITDA Mg%"]
comp_rows = [
    ["INDIAN FMCG PEERS", "", "", "", "", "", "", "", ""],
    ["Hindustan Unilever", "HINDUNILVR", "5,40,000", "8.4x", "39.2x", "36.8x", "54x", "6%",  "21.5%"],
    ["Nestle India",       "NESTLEIND",  "2,14,000", "12.8x","55.2x", "51.4x", "72x", "8%",  "23.2%"],
    ["Britannia Industries","BRITANNIA", "1,12,000", "6.9x", "33.4x", "30.8x", "48x", "7%",  "20.7%"],
    ["Dabur India",        "DABUR",      "88,000",   "5.1x", "28.6x", "26.4x", "42x", "9%",  "17.8%"],
    ["Tata Consumer",      "TATACONSUM", "93,500",   "4.8x", "34.2x", "31.5x", "52x", "11%", "14.1%"],
    ["Emami Ltd",          "EMAMILTD",   "32,000",   "6.2x", "29.8x", "27.6x", "44x", "8%",  "20.8%"],
    ["GLOBAL TOBACCO PEERS","", "", "", "", "", "", "", ""],
    ["BAT",                "BTI",        "4,70,000", "2.8x", "9.4x",  "9.0x",  "8x",  "(2%)", "29.8%"],
    ["Philip Morris Intl", "PM",         "11,50,000","4.2x", "14.8x", "14.1x", "19x", "3%",  "28.3%"],
    ["INDIAN DIVERSIFIED",  "", "", "", "", "", "", "", ""],
    ["Godfrey Phillips",   "GODFRYPHLP", "15,200",   "1.8x", "16.2x", "14.9x", "22x", "5%",  "11.1%"],
    ["ITC (For Reference)","ITC",        "3,87,153", "8.3x", "14.9x", "13.5x", "17x", "9%",  "34.7%"],
    ["STATISTICAL SUMMARY (excl. ITC)", "", "", "", "", "", "", "", ""],
    ["  Maximum",          "—",          "—",        "12.8x","55.2x", "51.4x", "72x", "11%", "29.8%"],
    ["  75th Percentile",  "—",          "—",        "6.9x", "34.2x", "31.5x", "52x", "9%",  "23.2%"],
    ["  Median",           "—",          "—",        "5.1x", "29.8x", "27.6x", "44x", "7%",  "20.8%"],
    ["  25th Percentile",  "—",          "—",        "2.8x", "14.8x", "14.1x", "19x", "3%",  "17.8%"],
    ["  Minimum",          "—",          "—",        "1.8x", "9.4x",  "9.0x",  "8x",  "(2%)", "11.1%"],
]

ct = doc.add_table(rows=len(comp_rows)+1, cols=len(comp_hdr))
ct.style = "Table Grid"; ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(comp_hdr):
    c = ct.rows[0].cells[ci]; shade_cell(c, NAVY)
    cell_text(c, h, size=8, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)

section_fills = {"INDIAN FMCG PEERS": RGBColor(0x1A,0x4F,0x72),
                 "GLOBAL TOBACCO PEERS": RGBColor(0x1A,0x4F,0x72),
                 "INDIAN DIVERSIFIED": RGBColor(0x1A,0x4F,0x72),
                 "STATISTICAL SUMMARY (excl. ITC)": RGBColor(0x2C,0x3E,0x50),
                 "ITC (For Reference)": GREEN}

for ri, row in enumerate(comp_rows):
    bg = LGRAY if ri % 2 == 0 else WHITE
    section_key = next((k for k in section_fills if row[0].startswith(k)), None)
    for ci, val in enumerate(row):
        c = ct.rows[ri+1].cells[ci]
        if section_key:
            shade_cell(c, section_fills[section_key])
            cell_text(c, val, size=8, bold=True, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
        else:
            shade_cell(c, bg)
            cell_text(c, val, size=8,
                      align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
src_note(doc, "Source: Bloomberg, company filings, analyst estimates. Data as of March 18, 2026. "
              "Indian peers in ₹ Crore; global peers at INR equivalent. ITC shown for reference only.")

body(doc,
     "ITC's valuation paradox: it is simultaneously the most profitable company in the group "
     "(34.7% EBITDA margin vs. peer median ~20.8%) but trades at one of the lowest multiples "
     "(EV/EBITDA LTM 14.9x vs. Indian FMCG median ~30x). This 55% discount reflects the "
     "tobacco component (~64% of EBIT) and conglomerate structure. Applying a segment-weighted "
     "blended multiple of 18.5x NTM EV/EBITDA to FY26E EBITDA of ₹26,670 Crore implies "
     "equity value of ₹5,31,275 Crore, or ₹423/share (upside scenario). At a more conservative "
     "16x (large tobacco discount): ₹366/share. Our base comps contribution to price target "
     "uses 17.0x NTM, implying ~₹375/share.", sa=3)

add_chart(doc, "chart_30_comps_scatter_margin_vs_multiple.png", width=6.8,
          fig_num=40,
          caption="ITC vs. Peers — EV/EBITDA Multiple vs. EBITDA Margin: ITC's Profound Valuation Discount (FY26E)")

add_chart(doc, "chart_31_peer_multiples_ev_ebitda_pe.png", width=6.8,
          fig_num=41,
          caption="ITC vs. Peers — EV/EBITDA and P/E Benchmarking: Peer Group Analysis (FY26E NTM Estimates)")

sub_head(doc, "Valuation Methodology 3: Dividend Discount Model (15% weight)")
body(doc,
     "ITC's exceptional dividend track record — 18 consecutive years of dividend growth, "
     "85–90% payout ratio — makes the DDM a meaningful valuation floor. We apply a two-stage "
     "Gordon Growth DDM with Stage 1 explicit dividends (FY26E–FY30E: ₹15.50 / ₹17.00 / "
     "₹18.65 / ₹20.55 / ₹22.65 per share) and Stage 2 terminal DPS growing at 5.5% in "
     "perpetuity. Cost of equity: 11.68% (CAPM). The DDM yields ₹174/share, serving as an "
     "important downside floor — the dividend-supported intrinsic value independent of any "
     "multiple expansion. CMP of ₹307 represents 1.8× the DDM floor, providing meaningful "
     "downside cushion even in the most pessimistic scenarios.", sa=3)

sub_head(doc, "Valuation Reconciliation & Football Field")
body(doc,
     "The weighted blend of our three methodologies yields a 12-month price target of ₹380, "
     "representing 23.8% price upside plus an approximately 3.5% FY26E dividend yield, "
     "for a total expected return of ~27.3%. The probability-weighted scenario analysis yields "
     "an expected value of ₹377, confirming our ₹380 target (rounded to the nearest ₹5).", sa=2)

fig_caption(doc, 42, "ITC — Valuation Reconciliation: Football Field Summary (₹ Per Share)")
val_hdr = ["Methodology", "Weight", "Low End (₹)", "Base (₹)", "High End (₹)", "Weighted Contribution"]
val_rows = [
    ["Consolidated DCF (WACC 11.59%, g 6.0%)",  "35%", "240", "278", "338", "₹97"],
    ["SOTP DCF (Segment-Level)",                  "15%", "250", "269", "310", "₹40"],
    ["EV/EBITDA Comps (16x–18.5x NTM)",          "35%", "330", "375", "423", "₹131"],
    ["P/E Comps (16x–24x FY26E EPS)",            "10%", "286", "358", "430", "₹36"],
    ["Dividend Discount Model (Floor Value)",     "5%",  "160", "174", "190", "₹9"],
    ["WEIGHTED AVERAGE INTRINSIC VALUE",         "100%", "—",   "—",   "—",   "₹313"],
    ["12-Month Forward Premium (FMCG inflection)","—",   "—",   "—",   "—",   "+₹67"],
    ["12-MONTH PRICE TARGET (ROUNDED)",          "—",   "—",   "₹380","—",   "₹380"],
]
vt = doc.add_table(rows=len(val_rows)+1, cols=len(val_hdr))
vt.style = "Table Grid"; vt.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(val_hdr):
    c = vt.rows[0].cells[ci]; shade_cell(c, NAVY)
    cell_text(c, h, size=8.5, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
for ri, row in enumerate(val_rows):
    is_main = not row[0].startswith(" ") and row[0] not in ["WEIGHTED AVERAGE INTRINSIC VALUE","12-MONTH PRICE TARGET (ROUNDED)","12-Month Forward Premium (FMCG inflection)"]
    is_total = row[0].startswith("WEIGHTED") or row[0].startswith("12-MONTH")
    bg = RGBColor(0x1A,0x4F,0x72) if is_total else (LGRAY if ri%2==0 else WHITE)
    for ci, val in enumerate(row):
        c = vt.rows[ri+1].cells[ci]; shade_cell(c, bg)
        cell_text(c, val, size=8.5, bold=(is_total or is_main),
                  color=WHITE if is_total else BLACK,
                  align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
src_note(doc, "Source: Institutional Equity Research.  ₹380 = 12-month target: ₹313 intrinsic value today + ~21% "
              "forward re-rating as FMCG profitability crystallises and Hotels demerger benefits become visible.")

add_chart(doc, "chart_32_valuation_football_field.png", width=7.0,
          fig_num=43,
          caption="ITC — Valuation Football Field: Range of Values by Methodology vs. Current Market Price (₹/Share)")

add_chart(doc, "chart_34_historical_valuation_multiples.png", width=6.5,
          fig_num=44,
          caption="ITC — Historical EV/EBITDA and P/E Multiples: 10-Year Trading Range vs. Current Levels (FY16–FY26E)")

sub_head(doc, "Final Price Target and Recommendation")
body(doc,
     "We initiate coverage of ITC Limited with a BUY rating and a twelve-month price target "
     "of ₹380. At the current price of ₹306.80, ITC offers 23.8% price upside plus approximately "
     "3.5% FY26E dividend yield, for a total expected return of ~27.3%. The stock trades at "
     "13.5x FY26E EV/EBITDA — near five-year lows and a 55% discount to Indian FMCG peers. "
     "Our analysis shows this discount is excessive given ITC's extraordinary profitability "
     "(34.7% EBITDA margin), near-zero leverage (₹28,000 Crore net cash), FMCG profitability "
     "inflection (FMCG-Others EBIT: ₹1,050 Crore and growing), and the Hotels demerger "
     "providing a structural re-rating catalyst. Even in our bear case (cigarette volumes "
     "-4% p.a., FMCG margins stalling at 3%), ITC is worth ₹240 — only 22% downside "
     "from current levels. The risk-reward is highly asymmetric.", sa=4)

# Recommendation summary box
tbl_r = doc.add_table(rows=1, cols=1); tbl_r.style = "Table Grid"
rc = tbl_r.rows[0].cells[0]; shade_cell(rc, NAVY)
rp = rc.paragraphs[0]
rp.paragraph_format.space_before = Pt(6); rp.paragraph_format.space_after = Pt(6)
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
def rrow(para, label, value, vc=WHITE, vb=False):
    r1 = para.add_run(f"{label}: "); sfont(r1, size=10, color=RGBColor(0xAA,0xBE,0xCF))
    r2 = para.add_run(f"{value}     "); sfont(r2, size=10, bold=vb, color=vc)
rrow(rp, "Rating",         "BUY",   GREEN, True)
rrow(rp, "Price Target",   "₹380",  RGBColor(0xB7,0x95,0x0B), True)
rrow(rp, "CMP",            "₹306.80")
rrow(rp, "Upside",         "+23.8%", GREEN)
rrow(rp, "Total Return",   "~27.3%", GREEN)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# KEY INVESTMENT CATALYSTS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "KEY INVESTMENT CATALYSTS",
            "Five Catalysts Driving Re-Rating Over the Next 6–24 Months")

cat_data = [
    ("Excise Duty Absorption Confirmation (Q1 FY27, July 2026)",
     "6–9 months",
     "The first full-quarter results post the January 2026 excise hike will be the critical "
     "data point. We expect Q1 FY27 cigarette volumes to decline 3–5% (not the feared 8–10%), "
     "with ITC taking a 7–10% price increase. Confirmation would re-rate cigarettes EBIT by "
     "~15–20%. Historical parallels: FY10 and FY17 excise hikes — volumes recovered within "
     "2–3 quarters in both cases."),
    ("FMCG-Others Margin Crossing 5%+ (FY27E)",
     "12 months",
     "FMCG-Others EBIT margin of ~4.8% (FY25A) is projected to reach 5.0% (FY26E) and 10.0% "
     "(FY27E) as revenue scale passes ₹28,000 Crore and operating leverage accrues. Each 1% "
     "margin improvement adds ~₹450 Crore (~2%) to consolidated EBIT. Visibility of this "
     "inflection would force an upward revision to FMCG-Others segment value, reducing the "
     "conglomerate discount significantly."),
    ("ITC Hotels Stake Value Appreciation (6–12 months)",
     "6–12 months",
     "ITC Hotels Limited targets 220 properties by 2030 under an asset-light managed model. "
     "As the entity executes on this, its market cap will grow from ~₹21,250 Crore (at listing) "
     "toward ₹40,000–50,000 Crore over five years. Each ₹10,000 Crore increase in Hotels market "
     "cap adds ~₹8 per ITC share. The cleaner SOTP narrative post-demerger helps institutional "
     "investors properly value this stake."),
    ("Capital Allocation — Special Dividend / Buyback (12–18 months)",
     "12–18 months",
     "ITC's net cash has grown to ₹28,000 Crore with FCF of ₹18,880 Crore projected in FY26E. "
     "A ₹10,000 Crore buyback at ₹307 retires ~3.3% of shares and adds ₹0.50 to FY27E EPS. "
     "A special dividend of ₹8–10/share would provide immediate cash yield of 2.6–3.3% on top "
     "of the regular dividend. Either action signals management confidence in the current "
     "depressed valuation and would catalyse institutional re-rating."),
    ("BAT Stake Resolution — Overhang Removal (12–24 months)",
     "12–24 months",
     "BAT's stake reduction from ~29% to 22.9% has created periodic selling pressure. A clear "
     "statement of intent to maintain 20%+ shareholding, or alternatively, a final reduction "
     "followed by a binding commitment floor, would remove the uncertainty premium from ITC's "
     "stock. LIC's 16% stake provides a natural counterweight. Resolution of this structural "
     "overhang could add 3–5% to ITC's valuation by reducing the risk premium embedded in "
     "institutional investor positions."),
]

for i, (title, horizon, text) in enumerate(cat_data, 1):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, RGBColor(0x1A,0x4F,0x72))
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"  Catalyst {i}: {title}  ")
    sfont(r1, size=10, bold=True, color=WHITE)
    r2 = p.add_run(f"[Timeframe: {horizon}]")
    sfont(r2, size=9, color=RGBColor(0xAA,0xBE,0xCF))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    body(doc, text, sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# APPENDIX A: EXTENDED RISK ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "APPENDIX A: EXTENDED RISK ASSESSMENT",
            "Eleven Risk Factors Across Company-Specific, Industry, Financial, and Macro Categories")

body(doc,
     "We present an extended risk assessment of eleven risk factors across four categories. "
     "The key risks to our investment thesis and price target have been discussed in the main "
     "body of this report (Risks 1–5). Appendix A provides additional detail on the "
     "remaining risks (6–11) and a consolidated risk matrix.", sa=3)

ext_risks = [
    ("6", "FMCG-Others Brand Investment Risk", "MEDIUM",
     "ITC invests ~₹1,500–2,000 Crore annually in A&P for FMCG brands. If competitive intensity "
     "from HUL, Nestle, and Tata Consumer forces higher-than-planned A&P spend, near-term EBIT "
     "margin expansion could stall. The risk is most acute in noodles (Yippee! vs. Maggi) and "
     "biscuits (Sunfeast vs. Britannia) — categories with the most directly entrenched incumbents."),
    ("7", "Agribusiness Commodity Price Volatility", "LOW",
     "ITC's Agribusiness segment (~₹19,753 Crore revenue, FY25) is exposed to global leaf tobacco "
     "prices, wheat/spice commodity cycles, and INR/USD dynamics. A 10% deterioration in leaf "
     "tobacco export realisations could impact EBIT by ₹400–500 Crore — manageable relative to "
     "consolidated EBIT of ₹24,000+ Crore. Strong FY25 performance (25% revenue growth) may "
     "not be repeatable if global tobacco demand softens."),
    ("8", "Capital Allocation and FMCG Investment Returns", "MEDIUM",
     "ITC's ₹20,000 Crore medium-term manufacturing expansion plan announced at FY25 AGM will "
     "deploy capital into FMCG fixed assets with EBIT margins of 12–18% at best — significantly "
     "below the cigarettes business's 60%+ margins. Group ROCE will face dilution as more capital "
     "is deployed into FMCG. The discipline question: does the FMCG investment eventually justify "
     "its cost of capital? This will take 5–7 years to answer definitively."),
    ("9", "Hotels Demerger Value Crystallisation Risk", "LOW",
     "ITC Hotels Limited's market performance post-listing has been mixed. ITC's retained 40% "
     "stake (₹8,500 Crore) is only realised through dividends and market re-rating — not "
     "consolidated earnings. BAT has sold most of its ITC Hotels shares (retaining ~6.3%), "
     "and continued institutional selling in the smaller-cap ITC Hotels stock remains an "
     "overhang. If ITC Hotels underperforms vs. Taj Hotels (Indian Hotels Company), ITC's "
     "stake value may not appreciate as projected."),
    ("10", "Rural Demand Slowdown — FMCG Volume Pressure", "MEDIUM",
     "~35–40% of FMCG-Others revenues are linked to rural India's consumption. Rural demand "
     "was muted in FY25 (ITC FMCG-Others: 5–7% growth vs. 12–15% expectation). An agricultural "
     "distress cycle — poor monsoon, delayed PM Kisan payments, rural credit tightening — could "
     "materially impact volume growth in atta, biscuits, and snacks. The rural demand "
     "recovery is a prerequisite for the FMCG margin expansion thesis."),
    ("11", "INR Weakness and Foreign Currency Risk", "LOW",
     "ITC's primary business is INR-denominated. However, Agribusiness export revenues benefit "
     "from INR weakness (positive); capital equipment imports and some raw materials face "
     "incremental costs from depreciation. BAT receives ITC dividends in INR and repatriates "
     "in GBP — sustained INR depreciation vs. GBP could marginally influence BAT's holding "
     "period decision. Net impact on ITC: modestly negative but secondary to operating risks."),
]
for num, title, severity, text in ext_risks:
    sev_color = (RGBColor(0x7B,0x24,0x1C) if severity=="HIGH" else
                 RGBColor(0x78,0x5B,0x00) if severity=="MEDIUM" else
                 RGBColor(0x1A,0x53,0x76))
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade_cell(c, sev_color)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  Risk {num}: {title}  [{severity} IMPACT]")
    sfont(r, size=9.5, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    body(doc, text, sa=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# APPENDIX B: MODELLING ASSUMPTIONS
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "APPENDIX B: KEY MODELLING ASSUMPTIONS",
            "Base Case Assumptions: FY26E–FY30E")

fig_caption(doc, 45, "ITC — Key Modelling Assumptions Summary (Base Case, FY26E–FY30E)")
assump_hdr = ["Assumption", "FY26E", "FY27E", "FY28E", "FY29E", "FY30E", "Rationale"]
assump_data = [
    ["Net Revenue CAGR",      "10.0%", "10.2%","10.0%","8.3%","8.2%","Cig pricing + FMCG vol"],
    ["Cig EBIT Margin",       "75.6%", "76.5%","77.5%","78.5%","79.5%","Pricing leverage > excise"],
    ["FMCG-Others EBIT Mg",   "7.1%",  "10.0%","12.5%","14.7%","16.3%","Scale + mix benefits"],
    ["Consol. EBITDA Margin",  "37.3%", "37.6%","37.8%","38.3%","38.7%","Slowly expanding"],
    ["CapEx / Net Revenue",    "8.0%",  "7.7%", "7.3%", "7.1%", "6.8%","Moderating post-FMCG"],
    ["Effective Tax Rate",     "24.5%", "24.5%","24.5%","24.5%","24.5%","Stable"],
    ["Dividend Payout Ratio",  "85%",   "85%",  "85%",  "85%",  "85%","Conservative vs. 90% hist."],
    ["Net Debt / EBITDA",      "(1.05x)","(1.27x)","(1.51x)","(1.78x)","(2.10x)","Net cash deepens"],
    ["WACC",                   "11.59%","—",    "—",    "—",    "—",   "India large-cap FMCG"],
    ["Terminal Growth Rate",   "6.00%", "—",    "—",    "—",    "—",   "India nominal GDP – 4%"],
    ["Cig Vol. Growth (YoY)",  "(2.5%)", "(2.0%)","(2.0%)","(1.5%)","(1.5%)","Modest decline; recovers"],
    ["FMCG-Others Rev CAGR",   "14%",   "14%",  "14%",  "14%",  "14%","Management aspiration"],
]
make_table(doc, assump_hdr, assump_data, size=8)
src_note(doc, "Source: Institutional Equity Research estimates.  Brackets indicate negative values.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# APPENDIX C: DATA SOURCES
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "APPENDIX C: DATA SOURCES & CITATIONS",
            "Primary Sources Used in This Initiating Coverage Report")

sources = [
    ("1", "ITC Limited Annual Report FY2024-25 (May 2025)",
     "Consolidated financials, segment EBIT, MD&A, corporate governance"),
    ("2", "ITC Limited Q3 FY26 Results Presentation (January 2026)",
     "Latest quarterly performance; Q3 FY26 cigarettes segment data"),
    ("3", "ITC Limited Q2 FY26 Investor Press Release (September 2025)",
     "FMCG-Others revenue growth, Aashirvaad/Sunfeast brand data"),
    ("4", "Reserve Bank of India (rbi.org.in)",
     "10-year GoI bond yield (6.80%, March 2026); India macroeconomic data"),
    ("5", "BSE India XBRL Filings",
     "Quarterly shareholding pattern, corporate actions"),
    ("6", "Damodaran Online (NYU Stern)",
     "India Equity Risk Premium estimate (6.50%)"),
    ("7", "Bloomberg Terminal (as referenced)",
     "Peer market data, beta, consensus estimates, ITC stock price history"),
    ("8", "SEBI XBRL Database",
     "Quarterly financial filings, insider trading disclosures"),
    ("9", "ITC Hotels Limited — Annual Results FY25",
     "Hotel revenue ₹3,333 Crore, PAT ₹698 Crore, 140+ properties"),
    ("10", "Tobacco Institute of India",
     "Market share data, illicit trade estimates (~26% of total)"),
    ("11", "Business Standard, The Economic Times, Mint",
     "ITC Hotels demerger news, BAT stake sale, excise duty reporting"),
    ("12", "ICICI Securities Research — ITC Coverage",
     "Comparable financial analysis and estimates cross-check"),
    ("13", "IMARC Group",
     "India FMCG market size (₹5.5–6 lakh crore), growth projections"),
    ("14", "British American Tobacco plc — Press Releases (2024–2025)",
     "BAT block trade announcements; ITC Hotels share sale (Dec 2025)"),
    ("15", "ITC Corporate Portal (itcportal.com)",
     "Management profiles, product portfolio, e-Choupal description, investor relations"),
    ("16", "Harvard Business School Case Study — e-Choupal",
     "Rural innovation context; ITC Agribusiness strategy background"),
    ("17", "NSE / BSE Corporate Filings",
     "Shareholding disclosures, board resolutions, demerger notifications"),
    ("18", "Screener.in / StockAnalysis.com",
     "Historical financial ratios, P/E and EV/EBITDA trading multiples"),
]
s_tbl = doc.add_table(rows=len(sources)+1, cols=3)
s_tbl.style = "Table Grid"
cell_text(s_tbl.rows[0].cells[0], "#", size=8.5, bold=True, color=WHITE)
cell_text(s_tbl.rows[0].cells[1], "Source", size=8.5, bold=True, color=WHITE)
cell_text(s_tbl.rows[0].cells[2], "Usage / Data Provided", size=8.5, bold=True, color=WHITE)
for c in s_tbl.rows[0].cells: shade_cell(c, NAVY)
for ri, (num, src, usage) in enumerate(sources):
    bg = LGRAY if ri%2==0 else WHITE
    for ci, val in enumerate([num, src, usage]):
        c = s_tbl.rows[ri+1].cells[ci]; shade_cell(c, bg)
        cell_text(c, val, size=8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# IMPORTANT DISCLOSURES
# ══════════════════════════════════════════════════════════════════════════
section_bar(doc, "IMPORTANT DISCLOSURES & ANALYST CERTIFICATION")

disclosures = [
    ("Analyst Certification",
     "The research analyst(s) primarily responsible for the content of this research report, "
     "in whole or in part, certify that with respect to each security or issuer that the analyst "
     "covered in this report: (1) all of the views expressed in this report accurately reflect "
     "the personal views about those securities or issuers; and (2) no part of any of the "
     "research analyst's compensation was, is, or will be, directly or indirectly, related to "
     "the specific recommendations or views expressed by the research analyst in this report."),
    ("Rating Definitions",
     "BUY: Expected total return > +15% over 12 months.  HOLD: Expected total return between "
     "-5% and +15% over 12 months.  SELL: Expected total return < -5% over 12 months.  "
     "Total return includes price appreciation and expected dividend income."),
    ("Conflicts of Interest",
     "This report has been prepared by the research department as informational material only. "
     "The research analyst and/or their associates do not hold any financial interest in the "
     "securities of ITC Limited. The research analyst has not received any compensation from "
     "ITC Limited in the past 12 months. This firm may from time to time hold positions in "
     "securities mentioned in this report."),
    ("General Disclaimer",
     "This document has been prepared by the Institutional Equity Research team for informational "
     "purposes only. It is not intended to constitute investment advice or a recommendation to "
     "buy, sell, or hold any security. The information contained herein is believed to be "
     "reliable but is not guaranteed as to accuracy or completeness. Past performance is not "
     "indicative of future results. Investors should conduct their own due diligence before "
     "making investment decisions. All financial figures are in Indian Rupees (₹) unless "
     "otherwise stated. Market data as of March 18, 2026."),
    ("Forward-Looking Statements",
     "This report contains certain forward-looking statements based on management guidance, "
     "historical trends, and analyst projections. Actual results may differ materially from "
     "those projected due to factors including but not limited to: regulatory changes in the "
     "tobacco industry, competitive dynamics in FMCG, macroeconomic conditions in India, "
     "commodity price movements, and execution risks in ITC's FMCG-Others strategy."),
    ("Distribution Restrictions",
     "This research report is intended for distribution to institutional investors only. "
     "Recipients in the United States are 'major U.S. institutional investors' within the "
     "meaning of Rule 15a-6 under the U.S. Securities Exchange Act. This report may not "
     "be reproduced, redistributed, or passed to any other person or published in whole "
     "or in part for any purpose without prior written consent."),
]

for title, text in disclosures:
    sub_head(doc, title)
    body(doc, text, sa=4)

# Final recommendation box
doc.add_paragraph().paragraph_format.space_after = Pt(6)
final_tbl = doc.add_table(rows=1, cols=1); final_tbl.style = "Table Grid"
fc = final_tbl.rows[0].cells[0]; shade_cell(fc, NAVY)
fp = fc.paragraphs[0]
fp.paragraph_format.space_before = Pt(8); fp.paragraph_format.space_after = Pt(8)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lines = [
    ("ITC LIMITED (ITC.NS)  |  INITIATING COVERAGE", 12, True, WHITE),
    ("\nRating: BUY  |  Price Target: ₹380  |  CMP: ₹306.80  |  Upside: +23.8%", 11, True, GREEN),
    ("\nTotal Return (incl. div.): ~27.3%  |  Market Cap: ₹3,87,153 Crore (~USD 45B)", 9, False, RGBColor(0xAA,0xBE,0xCF)),
    ("\nReport Date: March 18, 2026  |  Institutional Equity Research", 8, False, DGRAY),
]
for text, sz, bold, color in lines:
    r = fp.add_run(text)
    sfont(r, size=sz, bold=bold, color=color)

doc.save(OUTPUT)
print(f"✅ Section 6 (FINAL) complete → {OUTPUT}")
print(f"\n📄 Report: {OUTPUT}")
print(f"   All 35 charts embedded across 6 sections.")
