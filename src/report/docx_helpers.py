"""
Reusable python-docx helpers for building institutional equity research reports.
All styling is controlled via ReportTheme — no hardcoded colors.

These functions were extracted and DRY-fied from the ITC report build scripts.
"""

import os
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class ReportTheme:
    font_name: str = "Times New Roman"
    font_size_body: int = 9
    font_size_table: int = 8
    font_size_heading: int = 12

    # Colors as hex strings (no #)
    navy: str = "0D2B55"
    green: str = "1A753F"
    gold: str = "B7950B"
    lgray: str = "F2F2F2"
    dgray: str = "D9D9D9"
    white: str = "FFFFFF"
    black: str = "000000"
    blue_light: str = "D6E4F0"  # projection column shading


DEFAULT_THEME = ReportTheme()

# ── Alignment mapping ────────────────────────────────────────────────────────

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ── Core run / cell helpers ──────────────────────────────────────────────────

def sfont(run, size: float, bold: bool = False, italic: bool = False,
          color_hex: str = None, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Set font properties on a docx Run."""
    run.font.name = theme.font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def shade_cell(cell, fill_hex: str, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Apply background shading to a table cell via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    # Remove any existing shd element and replace
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def cell_text(cell, text: str, size: float = 8, bold: bool = False,
              italic: bool = False, color_hex: str = None,
              align: str = "left", theme: ReportTheme = DEFAULT_THEME) -> None:
    """Clear cell, add a run with styled text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    sfont(run, size, bold, italic, color_hex, theme)


# ── Document-level building blocks ──────────────────────────────────────────

def section_bar(doc, title: str, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Full-width NAVY bar with white bold text — section header."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, theme.navy, theme)
    cell_text(cell, title,
              size=float(theme.font_size_heading),
              bold=True,
              color_hex=theme.white,
              align="left",
              theme=theme)
    # Add spacing after
    doc.add_paragraph()


def sub_head(doc, text: str, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Gold underline sub-heading paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    sfont(run, float(theme.font_size_heading), bold=True,
          color_hex=theme.gold, theme=theme)
    run.font.underline = True


def body(doc, text: str, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Normal body paragraph in Times New Roman body size."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    sfont(run, float(theme.font_size_body), theme=theme)


def bullet(doc, text: str, level: int = 0,
           theme: ReportTheme = DEFAULT_THEME) -> None:
    """Bullet point paragraph using List Bullet style."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    sfont(run, float(theme.font_size_body), theme=theme)


def src_note(doc, text: str, theme: ReportTheme = DEFAULT_THEME) -> None:
    """Small gray italic source note."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    sfont(run, 7.0, italic=True, color_hex=theme.dgray, theme=theme)


def fig_caption(doc, text: str, fig_num: int = None,
                theme: ReportTheme = DEFAULT_THEME) -> None:
    """Figure caption: 'Figure X: text' in small italic."""
    label = f"Figure {fig_num}: {text}" if fig_num is not None else text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    sfont(run, 7.5, italic=True, color_hex=theme.black, theme=theme)


def add_chart(doc, image_path: str, width_inches: float = 6.8,
              caption: str = None, fig_num: int = None,
              theme: ReportTheme = DEFAULT_THEME) -> None:
    """Add an image to the document with optional caption."""
    if not image_path or not os.path.exists(image_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    if caption:
        fig_caption(doc, caption, fig_num, theme)


def make_fin_table(doc, headers: list, rows: list,
                   col_widths: list = None,
                   highlight_rows: list = None,
                   proj_start_col: int = None,
                   theme: ReportTheme = DEFAULT_THEME) -> None:
    """
    Build a styled financial table (P&L, Balance Sheet, etc.).

    headers: column header strings
    rows: list of (row_label, val1, val2, ...) tuples
          val can be float, str, or None
    col_widths: list of widths in inches. Default: first col 2.0", rest equal
    highlight_rows: row indices to bold (for subtotals/totals)
    proj_start_col: column index where projections start (light blue shading)

    Styling:
    - Header row: NAVY background, white bold text, centered
    - Odd data rows: white background
    - Even data rows: LGRAY background
    - Highlighted rows: DGRAY background, bold text
    - proj_start_col onwards: BLUE_LIGHT background for header
    - Numbers: right-aligned
    - Row labels: left-aligned
    - All text: font_size_table
    """
    if not headers or not rows:
        return

    n_cols = len(headers)
    highlight_rows = highlight_rows or []

    # Default column widths
    if col_widths is None:
        remaining = (6.5 - 2.0) / max(n_cols - 1, 1)
        col_widths = [2.0] + [remaining] * (n_cols - 1)

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Set column widths
    for i, width in enumerate(col_widths[:n_cols]):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    # Header row
    hdr_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = hdr_row.cells[col_idx]
        # Projection columns get blue_light; others get navy
        if proj_start_col is not None and col_idx >= proj_start_col:
            shade_cell(cell, theme.blue_light, theme)
            cell_text(cell, str(header),
                      size=float(theme.font_size_table),
                      bold=True, color_hex=theme.navy,
                      align="center", theme=theme)
        else:
            shade_cell(cell, theme.navy, theme)
            cell_text(cell, str(header),
                      size=float(theme.font_size_table),
                      bold=True, color_hex=theme.white,
                      align="center", theme=theme)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        table_row = table.rows[row_idx + 1]
        is_highlighted = row_idx in highlight_rows
        row_fill = theme.dgray if is_highlighted else (
            theme.white if row_idx % 2 == 0 else theme.lgray
        )

        for col_idx in range(n_cols):
            cell = table_row.cells[col_idx]
            shade_cell(cell, row_fill, theme)

            if col_idx < len(row_data):
                raw = row_data[col_idx]
            else:
                raw = None

            # Format value
            if raw is None:
                text = "\u2013"
            elif isinstance(raw, float) and col_idx > 0:
                text = f"{raw:,.1f}" if abs(raw) < 1000 and raw != int(raw) else f"{raw:,.0f}"
            else:
                text = str(raw)

            align = "left" if col_idx == 0 else "right"
            cell_text(cell, text,
                      size=float(theme.font_size_table),
                      bold=is_highlighted,
                      color_hex=theme.black,
                      align=align,
                      theme=theme)

    doc.add_paragraph()


def make_key_metrics_box(doc, metrics: dict,
                          cols: int = 3,
                          theme: ReportTheme = DEFAULT_THEME) -> None:
    """
    Compact metrics box: N items laid out in `cols` columns.
    Each cell: metric name (small, gray) + value (larger, bold, NAVY).
    Used for investment summary key stats.
    """
    if not metrics:
        return

    items = list(metrics.items())
    n_rows = (len(items) + cols - 1) // cols

    table = doc.add_table(rows=n_rows, cols=cols)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(items):
        r = i // cols
        c = i % cols
        cell = table.cell(r, c)
        shade_cell(cell, theme.lgray, theme)
        # Clear cell and build two runs in the same paragraph
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        label_run = p.add_run(str(label) + "\n")
        sfont(label_run, 7.0, italic=True, color_hex=theme.dgray, theme=theme)

        val_run = p.add_run(str(value))
        sfont(val_run, float(theme.font_size_body) + 1, bold=True,
              color_hex=theme.navy, theme=theme)

    # Fill any empty trailing cells
    total_cells = n_rows * cols
    for i in range(len(items), total_cells):
        r = i // cols
        c = i % cols
        shade_cell(table.cell(r, c), theme.lgray, theme)

    doc.add_paragraph()


def make_rating_box(doc, rating: str, price_target: float, cmp: float,
                    upside_pct: float, theme: ReportTheme = DEFAULT_THEME) -> None:
    """
    Large prominent rating box for cover page / investment summary.
    NAVY background, white text:
    - "INITIATING COVERAGE" header
    - Rating: BUY/HOLD/SELL in large bold
    - Price Target: Rs.XXX
    - CMP: Rs.XXX
    - Upside: XX%
    """
    table = doc.add_table(rows=5, cols=1)
    table.style = "Table Grid"

    labels = [
        ("INITIATING COVERAGE", 9.0, False),
        (str(rating).upper(), 16.0, True),
        (f"Price Target: \u20b9{price_target:,.0f}", 10.0, True),
        (f"CMP: \u20b9{cmp:,.0f}", 10.0, False),
        (f"Upside: {upside_pct:+.1f}%", 10.0, True),
    ]

    for row_idx, (text, size, bold) in enumerate(labels):
        cell = table.rows[row_idx].cells[0]
        shade_cell(cell, theme.navy, theme)
        cell_text(cell, text, size=size, bold=bold,
                  color_hex=theme.white, align="center", theme=theme)

    doc.add_paragraph()


def page_break(doc) -> None:
    """Add a page break."""
    doc.add_page_break()


def set_margins(doc, top: float = 1.0, bottom: float = 1.0,
                left: float = 1.0, right: float = 1.0) -> None:
    """Set page margins in inches for all sections."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
