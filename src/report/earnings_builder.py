"""
EarningsBuilder — quarterly earnings update report generator (8-12 page DOCX).

Produces a concise earnings flash note for Indian listed companies after
quarterly results. Driven by the same JSON data contract as ReportBuilder,
with the addition of a top-level "quarterly" key containing actuals,
estimates, prior-period comparisons, management commentary, and revised
estimates / price target.

Usage:
    from earnings_builder import EarningsBuilder

    builder = EarningsBuilder(output_path="reports/ITC/ITC_Q3FY25_Earnings.docx")
    builder.build(report_json)   # returns output_path
"""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

# ---------------------------------------------------------------------------
# Graceful imports — works both from within the project tree and standalone
# ---------------------------------------------------------------------------
try:
    from src.report.docx_helpers import (
        ReportTheme, DEFAULT_THEME, sfont, shade_cell, cell_text,
        section_bar, sub_head, body, bullet, src_note,
        add_chart, make_fin_table, make_key_metrics_box, make_rating_box,
        page_break, set_margins,
    )
except ImportError:
    try:
        from docx_helpers import (  # type: ignore
            ReportTheme, DEFAULT_THEME, sfont, shade_cell, cell_text,
            section_bar, sub_head, body, bullet, src_note,
            add_chart, make_fin_table, make_key_metrics_box, make_rating_box,
            page_break, set_margins,
        )
    except ImportError as exc:
        raise ImportError(
            "Cannot import docx_helpers. Run from the IndiaStockResearch "
            "project root or install the package."
        ) from exc

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Color constants (project-wide palette, same as chart_style.py)
# ---------------------------------------------------------------------------
NAVY  = "0D2B55"
GREEN = "1A753F"
GOLD  = "B7950B"
RED   = "8B0000"   # dark red for "miss" cells
LGRAY = "F2F2F2"
DGRAY = "D9D9D9"
WHITE = "FFFFFF"
BLACK = "000000"
BEAT_FILL  = "E8F5E9"   # light green for beat rows
MISS_FILL  = "FFEBEE"   # light red for miss rows


# ---------------------------------------------------------------------------
# EarningsBuilder
# ---------------------------------------------------------------------------

class EarningsBuilder:
    """
    Builds a quarterly earnings update DOCX report (8-12 pages).

    Parameters
    ----------
    output_path : str
        Destination file path for the generated DOCX.
    theme : ReportTheme, optional
        Styling theme; defaults to DEFAULT_THEME (Times New Roman, NAVY/GREEN/GOLD).
    """

    def __init__(self, output_path: str, theme: ReportTheme = None):
        self.output_path = output_path
        self.theme = theme or DEFAULT_THEME
        self.doc = Document()

    # =========================================================================
    # Public entry point
    # =========================================================================

    def build(self, report_json: dict) -> str:
        """
        Assemble the complete earnings update DOCX from report_json.

        Sections
        --------
        1.  Earnings Flash Header          (Page 1)
        2.  Quarterly P&L Comparison Table (Page 2)
        3.  QoQ Trend Charts               (Page 2-3)
        4.  Key Segment Performance        (Page 3-4)
        5.  Updated Estimates Table        (Page 4-5)
        6.  Updated Valuation              (Page 5-6)
        7.  Financial Summary Table        (Page 6-8)
        8.  Rating and Disclosure          (Page 8)

        Returns
        -------
        str
            Absolute path of the saved DOCX file.
        """
        meta       = report_json.get("meta", {})
        market_data = report_json.get("market_data", {})
        financials = report_json.get("financials", {})
        ratios     = report_json.get("ratios", {})
        projections = report_json.get("projections", {})
        valuation  = report_json.get("valuation", {})
        charts     = report_json.get("charts", {})
        narrative  = report_json.get("narrative", {})
        quarterly  = report_json.get("quarterly")   # may be None

        set_margins(self.doc, top=0.9, bottom=0.9, left=1.0, right=1.0)

        # Section 1 — Earnings Flash Header
        self._build_flash_header(meta, market_data, valuation, quarterly)

        # Section 2 — Quarterly P&L Comparison Table
        self._build_quarterly_pl_table(quarterly)

        # Section 3 — QoQ Trend Charts
        self._build_trend_charts(charts)

        # Section 4 — Key Segment Performance
        self._build_segment_performance(financials, narrative)

        # Section 5 — Updated Estimates Table
        self._build_estimates_revision(quarterly, report_json)

        # Section 6 — Updated Valuation
        self._build_updated_valuation(quarterly, valuation, narrative, market_data)

        # Section 7 — Financial Summary Table
        self._build_financial_summary(financials, ratios, projections)

        # Section 8 — Rating and Disclosure
        self._build_rating_disclosure(meta, valuation, quarterly, narrative)

        # Save
        out_dir = os.path.dirname(os.path.abspath(self.output_path))
        os.makedirs(out_dir, exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path

    # =========================================================================
    # Section 1 — Earnings Flash Header
    # =========================================================================

    def _build_flash_header(self, meta: dict, market_data: dict,
                             valuation: dict, quarterly) -> None:
        """
        Page 1: cover-flash combining company identity, rating box,
        beat/miss summary and management commentary highlights.
        """
        doc, theme = self.doc, self.theme
        q = quarterly or {}

        company  = meta.get("company_name", "Company")
        ticker   = meta.get("ticker", "TICK")
        sector   = meta.get("sector", "Sector")
        qlabel   = q.get("quarter_label", "Q?FY??")
        rdate    = meta.get("report_date", "")

        # ── Company / quarter title ──────────────────────────────────────────
        p_company = doc.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_company = p_company.add_run(company)
        sfont(run_company, 20.0, bold=True, color_hex=NAVY, theme=theme)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(
            f"{qlabel} Results Review  |  {sector}  |  NSE: {ticker}"
        )
        sfont(run_sub, 10.0, italic=True, color_hex=DGRAY, theme=theme)

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_date = p_date.add_run(f"Report Date: {rdate}")
        sfont(run_date, 8.0, italic=True, color_hex=BLACK, theme=theme)

        doc.add_paragraph()

        # ── Rating box — uses revised values if available ────────────────────
        rating  = q.get("revised_rating", valuation.get("rating", "BUY"))
        pt      = q.get("revised_price_target",
                        valuation.get("price_target", 0.0))
        cmp_val = market_data.get("cmp", valuation.get("cmp", 0.0))
        upside  = self._pct_change(pt, cmp_val) if cmp_val else 0.0
        make_rating_box(doc, rating, pt, cmp_val, upside, theme)

        # ── Key metrics box: Revenue, EBITDA, PAT — actual vs estimate ───────
        section_bar(doc, f"{qlabel} — KEY RESULTS SNAPSHOT", theme)

        actuals   = q.get("actuals", {})
        estimates = q.get("estimates", {})

        metrics = {}
        for label, a_key, e_key in [
            ("Revenue",     "revenue",  "revenue"),
            ("EBITDA",      "ebitda",   "ebitda"),
            ("EBITDA Mgn",  "ebitda_margin", "ebitda_margin"),
            ("PAT",         "pat",      "pat"),
            ("EPS (Rs.)",   "eps",      "eps"),
        ]:
            a_val = actuals.get(a_key)
            e_val = estimates.get(e_key)
            if a_key in ("ebitda_margin",):
                a_str = f"{a_val:.1f}%" if a_val is not None else "N/A"
                e_str = f"{e_val:.1f}%" if e_val is not None else "N/A"
            elif a_key == "eps":
                a_str = f"Rs. {a_val:.2f}" if a_val is not None else "N/A"
                e_str = f"Rs. {e_val:.2f}" if e_val is not None else "N/A"
            else:
                a_str = self._format_cr(a_val)
                e_str = self._format_cr(e_val)
            metrics[label] = f"A: {a_str}  |  E: {e_str}"

        make_key_metrics_box(doc, metrics, cols=3, theme=theme)

        # ── Beat / Miss 3-column summary ─────────────────────────────────────
        sub_head(doc, "Beat / Miss Summary", theme)
        self._build_beat_miss_summary(q)

        # ── What happened — management commentary bullets ─────────────────────
        section_bar(doc, "WHAT HAPPENED — MANAGEMENT COMMENTARY HIGHLIGHTS", theme)
        commentary = q.get("management_commentary", [])
        if commentary:
            for point in commentary[:6]:
                bullet(doc, str(point), level=0, theme=theme)
        else:
            bullet(doc, "Management commentary not available.", theme=theme)

        page_break(doc)

    def _build_beat_miss_summary(self, q: dict) -> None:
        """3-column table: Revenue beat/miss | EBITDA beat/miss | PAT beat/miss."""
        doc, theme = self.doc, self.theme

        actuals   = q.get("actuals", {})
        estimates = q.get("estimates", {})

        rev_pct   = q.get("beat_miss_revenue_pct",
                          self._pct_change(actuals.get("revenue"),
                                           estimates.get("revenue")))
        ebitda_pct = self._pct_change(actuals.get("ebitda"),
                                      estimates.get("ebitda"))
        pat_pct   = q.get("beat_miss_pat_pct",
                          self._pct_change(actuals.get("pat"),
                                           estimates.get("pat")))

        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"

        headers = ["Revenue vs. Estimate", "EBITDA vs. Estimate", "PAT vs. Estimate"]
        values  = [
            self._beat_miss_str(rev_pct),
            self._beat_miss_str(ebitda_pct),
            self._beat_miss_str(pat_pct),
        ]
        pcts    = [rev_pct, ebitda_pct, pat_pct]

        for col_idx, (hdr, val, pct) in enumerate(zip(headers, values, pcts)):
            h_cell = table.rows[0].cells[col_idx]
            shade_cell(h_cell, NAVY, theme)
            cell_text(h_cell, hdr, size=float(theme.font_size_table),
                      bold=True, color_hex=WHITE, align="center", theme=theme)

            v_cell = table.rows[1].cells[col_idx]
            fill = BEAT_FILL if (pct is not None and pct >= 0) else MISS_FILL
            shade_cell(v_cell, fill, theme)
            txt_color = GREEN if (pct is not None and pct >= 0) else RED
            cell_text(v_cell, val, size=float(theme.font_size_body) + 2,
                      bold=True, color_hex=txt_color, align="center", theme=theme)

        doc.add_paragraph()

    # =========================================================================
    # Section 2 — Quarterly P&L Comparison Table
    # =========================================================================

    def _build_quarterly_pl_table(self, quarterly) -> None:
        """
        5-column quarterly P&L comparison table.
        Columns: Metric | Q?FY?? A | Q?FY?? E | vs Est | Q?FY??-1 A | YoY%
        Beat rows shaded green, miss rows shaded red.
        """
        doc, theme = self.doc, self.theme
        q = quarterly or {}

        qlabel = q.get("quarter_label", "Q?FY??")

        # Derive prior year quarter label (e.g. Q3FY25 -> Q3FY24)
        pyq_label = self._prior_year_quarter_label(qlabel)

        section_bar(doc,
                    f"QUARTERLY P&L COMPARISON — {qlabel}",
                    theme)

        if not q:
            body(doc, "Quarterly data not available.", theme)
            page_break(doc)
            return

        self._quarterly_comparison_table(doc, q)
        page_break(doc)

    def _quarterly_comparison_table(self, doc, quarterly_data: dict) -> None:
        """
        Build the 6-column beat/miss-colored quarterly comparison table.

        Columns
        -------
        Metric | Actual | Estimate | vs Est% | Prior Quarter | YoY%
        """
        theme = self.theme
        q     = quarterly_data or {}

        qlabel    = q.get("quarter_label", "Q?FY??")
        pyq_label = self._prior_year_quarter_label(qlabel)
        pqq_label = self._prior_quarter_label(qlabel)

        actuals   = q.get("actuals", {})
        estimates = q.get("estimates", {})
        pqq       = q.get("prior_quarter", {})
        pyq       = q.get("prior_year_quarter", {})

        headers = [
            "Metric",
            f"{qlabel} Actual",
            f"{qlabel} Estimate",
            "vs. Est",
            f"{pqq_label} Actual",
            f"QoQ%",
            f"{pyq_label} Actual",
            "YoY%",
        ]
        col_widths = [1.7, 0.8, 0.8, 0.6, 0.8, 0.55, 0.8, 0.55]

        # Each row: (label, actual_key, est_key, pqq_key, pyq_key, is_margin)
        row_specs = [
            ("Net Revenue (Rs. Cr)",   "revenue",      "revenue",      "revenue",      "revenue",      False),
            ("EBITDA (Rs. Cr)",        "ebitda",       "ebitda",       "ebitda",       "ebitda",       False),
            ("EBITDA Margin (%)",      "ebitda_margin","ebitda_margin","ebitda_margin","ebitda_margin", True),
            ("Depreciation (Rs. Cr)",  "depreciation", "depreciation", "depreciation", "depreciation", False),
            ("EBIT (Rs. Cr)",          "ebit",         "ebit",         "ebit",         "ebit",         False),
            ("Other Income (Rs. Cr)",  "other_income", "other_income", "other_income", "other_income", False),
            ("PBT (Rs. Cr)",           "pbt",          "pbt",          "pbt",          "pbt",          False),
            ("Tax (Rs. Cr)",           "tax",          "tax",          "tax",          "tax",          False),
            ("PAT (Rs. Cr)",           "pat",          "pat",          "pat",          "pat",          False),
            ("EPS (Rs.)",              "eps",          "eps",          "eps",          "eps",          False),
        ]

        # Bold / highlight rows (0-indexed among data rows)
        highlight_idx = {0, 1, 8, 9}  # Revenue, EBITDA, PAT, EPS

        # Build table
        n_data_rows = len(row_specs)
        table = doc.add_table(rows=1 + n_data_rows, cols=len(headers))
        table.style = "Table Grid"

        # Set column widths
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

        # Header row
        for ci, hdr in enumerate(headers):
            cell = table.rows[0].cells[ci]
            shade_cell(cell, NAVY, theme)
            cell_text(cell, hdr, size=float(theme.font_size_table),
                      bold=True, color_hex=WHITE, align="center", theme=theme)

        # Data rows
        for ri, (label, a_key, e_key, pqq_key, pyq_key, is_margin) in enumerate(row_specs):
            a_val   = actuals.get(a_key)
            e_val   = estimates.get(e_key)
            pqq_val = pqq.get(pqq_key)
            pyq_val = pyq.get(pyq_key)

            vs_est_pct = self._pct_change(a_val, e_val)
            qoq_pct    = self._pct_change(a_val, pqq_val)
            yoy_pct    = self._pct_change(a_val, pyq_val)

            # Base row fill
            is_bold = ri in highlight_idx
            if vs_est_pct is not None:
                row_fill = BEAT_FILL if vs_est_pct >= 0 else MISS_FILL
            else:
                row_fill = WHITE if ri % 2 == 0 else LGRAY

            table_row = table.rows[ri + 1]

            def _fmt(val, margin=False):
                if val is None:
                    return "\u2013"
                if margin:
                    return f"{val:.1f}%"
                if a_key == "eps":
                    return f"{val:.2f}"
                return f"{val:,.0f}"

            def _pct_str(pct):
                if pct is None:
                    return "\u2013"
                sign = "+" if pct >= 0 else ""
                return f"{sign}{pct:.1f}%"

            row_vals = [
                label,
                _fmt(a_val, is_margin),
                _fmt(e_val, is_margin),
                _pct_str(vs_est_pct),
                _fmt(pqq_val, is_margin),
                _pct_str(qoq_pct),
                _fmt(pyq_val, is_margin),
                _pct_str(yoy_pct),
            ]

            for ci, text in enumerate(row_vals):
                cell = table_row.cells[ci]
                shade_cell(cell, row_fill, theme)

                # Color-code the vs. est column
                if ci == 3 and vs_est_pct is not None:
                    txt_color = GREEN if vs_est_pct >= 0 else RED
                else:
                    txt_color = BLACK

                align = "left" if ci == 0 else "right"
                cell_text(cell, text,
                          size=float(theme.font_size_table),
                          bold=is_bold,
                          color_hex=txt_color,
                          align=align,
                          theme=theme)

        doc.add_paragraph()
        src_note(doc,
                 f"A = Actual; E = Consensus estimate; "
                 f"QoQ vs. {self._prior_quarter_label(q.get('quarter_label',''))};  "
                 f"YoY vs. {self._prior_year_quarter_label(q.get('quarter_label',''))}",
                 theme)
        doc.add_paragraph()

    # =========================================================================
    # Section 3 — QoQ Trend Charts
    # =========================================================================

    def _build_trend_charts(self, charts: dict) -> None:
        """Embed revenue/PAT trend and margin trend charts if present."""
        doc, theme = self.doc, self.theme

        has_chart = False
        for chart_key, caption in [
            ("chart_02", "Revenue & PAT Trend (Quarterly)"),
            ("chart_10", "EBITDA Margin Trend (Quarterly)"),
        ]:
            path = (charts or {}).get(chart_key, "")
            if path and os.path.exists(str(path)):
                if not has_chart:
                    section_bar(doc, "QoQ TREND CHARTS", theme)
                    has_chart = True
                add_chart(doc, str(path), width_inches=6.5,
                          caption=caption, theme=theme)

        if has_chart:
            page_break(doc)

    # =========================================================================
    # Section 4 — Key Segment Performance
    # =========================================================================

    def _build_segment_performance(self, financials: dict,
                                    narrative: dict) -> None:
        """
        Segment revenue, EBIT and margin table (where segment data exists),
        plus a brief commentary excerpt from narrative.financial_analysis.
        """
        doc, theme = self.doc, self.theme

        section_bar(doc, "KEY SEGMENT PERFORMANCE", theme)

        segments = (financials or {}).get("segments", {})
        if segments:
            sub_head(doc, "Segment Revenue & Profitability", theme)
            self._build_segment_table(segments)
        else:
            body(doc, "Segment-level quarterly data not available in this report.", theme)

        # Commentary excerpt
        fin_analysis = (narrative or {}).get("financial_analysis", "")
        if fin_analysis:
            sub_head(doc, "Financial Analysis Commentary", theme)
            excerpt = str(fin_analysis)[:700].rstrip()
            if len(fin_analysis) > 700:
                excerpt += "..."
            body(doc, excerpt, theme)

        page_break(doc)

    def _build_segment_table(self, segments: dict) -> None:
        """
        Table: Segment | Revenue | QoQ% | YoY% | EBIT Margin
        segments expected to be a dict of dicts, e.g.:
        {"FMCG-Cigarettes": {"revenue": 9000, "ebit": 3800, "margin": 42.2}, ...}
        """
        doc, theme = self.doc, self.theme

        if not segments:
            return

        headers = ["Segment", "Revenue (Rs. Cr)", "EBIT (Rs. Cr)", "EBIT Margin (%)"]
        col_widths = [2.2, 1.3, 1.3, 1.3]
        rows = []

        for seg_name, seg_data in segments.items():
            if not isinstance(seg_data, dict):
                continue
            rev  = seg_data.get("revenue")
            ebit = seg_data.get("ebit")
            mgn  = seg_data.get("margin") or seg_data.get("ebit_margin")

            rows.append((
                str(seg_name),
                f"{rev:,.0f}"   if rev  is not None else "\u2013",
                f"{ebit:,.0f}"  if ebit is not None else "\u2013",
                f"{mgn:.1f}%"   if mgn  is not None else "\u2013",
            ))

        if rows:
            make_fin_table(doc, headers, rows,
                           col_widths=col_widths,
                           theme=theme)

    # =========================================================================
    # Section 5 — Updated Estimates Table
    # =========================================================================

    def _build_estimates_revision(self, quarterly, report_json: dict) -> None:
        """
        'Revising Estimates' section with a table comparing old vs new
        revenue and PAT estimates for forward FY periods.
        """
        doc, theme = self.doc, self.theme

        section_bar(doc, "REVISING ESTIMATES", theme)
        self._estimates_revision_table(doc, quarterly, report_json)
        page_break(doc)

    def _estimates_revision_table(self, doc, quarterly_data,
                                   report_json: dict) -> None:
        """
        Table: FY | Old Revenue | New Revenue | Rev Chg% | Old PAT | New PAT | PAT Chg%
        Falls back to "No Change" if revised_estimates is absent.
        """
        theme = self.theme
        q     = quarterly_data or {}

        projections       = report_json.get("projections", {})
        revised_estimates = q.get("revised_estimates", {})
        proj_fy_labels    = projections.get("fy_labels", [])

        sub_head(doc, "Estimate Revisions", theme)

        if not proj_fy_labels:
            body(doc, "Forward projection data not available.", theme)
            return

        # Old estimates come from projections (pre-earnings)
        old_rev_series = projections.get("net_revenue", [])
        old_pat_series = projections.get("pat", [])

        headers    = ["FY Period", "Old Revenue", "New Revenue", "Rev Chg%",
                      "Old PAT", "New PAT", "PAT Chg%"]
        col_widths = [1.1, 1.05, 1.05, 0.75, 1.05, 1.05, 0.75]
        rows       = []
        highlight_rows = []

        for idx, fy in enumerate(proj_fy_labels):
            old_rev = old_rev_series[idx] if idx < len(old_rev_series) else None
            old_pat = old_pat_series[idx]  if idx < len(old_pat_series) else None

            # Revised values: look for keys like "fy26e_pat", "fy26e_revenue"
            fy_key_base = fy.lower().replace(" ", "")   # "fy26e"
            new_rev = revised_estimates.get(f"{fy_key_base}_revenue") or \
                      revised_estimates.get(f"{fy_key_base}_rev")
            new_pat = revised_estimates.get(f"{fy_key_base}_pat")

            # If no revision found, treat as unchanged
            if new_rev is None:
                new_rev = old_rev
            if new_pat is None:
                new_pat = old_pat

            rev_chg = self._pct_change(new_rev, old_rev)
            pat_chg = self._pct_change(new_pat, old_pat)

            def _fmt_cr(v):
                return self._format_cr(v)

            def _chg_str(pct):
                if pct is None:
                    return "No Change"
                sign = "+" if pct >= 0 else ""
                return f"{sign}{pct:.1f}%"

            rows.append((
                fy,
                _fmt_cr(old_rev),
                _fmt_cr(new_rev),
                _chg_str(rev_chg),
                _fmt_cr(old_pat),
                _fmt_cr(new_pat),
                _chg_str(pat_chg),
            ))
            # Highlight rows where estimates were actually revised
            if rev_chg or pat_chg:
                highlight_rows.append(len(rows) - 1)

        if rows:
            make_fin_table(doc, headers, rows,
                           col_widths=col_widths,
                           highlight_rows=highlight_rows,
                           theme=theme)
        else:
            body(doc, "No estimate revisions to display.", theme)

        # Revised price target note
        revised_pt = q.get("revised_price_target")
        if revised_pt is not None:
            old_pt = report_json.get("valuation", {}).get("price_target")
            pt_note = (
                f"Price target revised to \u20b9{revised_pt:,.0f}"
                + (f" from \u20b9{old_pt:,.0f}" if old_pt else "")
                + " on revised earnings."
            )
            sub_head(doc, "Price Target Revision", theme)
            body(doc, pt_note, theme)

    # =========================================================================
    # Section 6 — Updated Valuation
    # =========================================================================

    def _build_updated_valuation(self, quarterly, valuation: dict,
                                  narrative: dict, market_data: dict) -> None:
        """
        Revised price target, valuation multiples at CMP, 1-paragraph rationale.
        """
        doc, theme = self.doc, self.theme
        q   = quarterly or {}

        section_bar(doc, "UPDATED VALUATION", theme)

        # Effective rating / PT (use revised if available)
        rating  = q.get("revised_rating", valuation.get("rating", "BUY"))
        pt      = q.get("revised_price_target", valuation.get("price_target", 0.0))
        cmp_val = market_data.get("cmp", valuation.get("cmp", 0.0))
        upside  = self._pct_change(pt, cmp_val) if cmp_val else 0.0

        make_rating_box(doc, rating, pt, cmp_val, upside, theme)

        # Methodology note
        sub_head(doc, "Valuation Methodology", theme)
        val_narrative = (narrative or {}).get("valuation", "")
        if val_narrative:
            body(doc, str(val_narrative)[:500], theme)
        else:
            method = valuation.get("methodology", "DCF / EV-EBITDA blended approach")
            body(doc, f"Valuation based on {method}.", theme)

        # Key multiples at CMP
        sub_head(doc, "Key Valuation Multiples at CMP", theme)

        pe_series       = valuation.get("pe_series", [])
        ev_series       = valuation.get("ev_ebitda_series", [])
        proj_fy_labels  = valuation.get("proj_fy_labels", [])

        # Build a 3-column metrics dict
        multiples = {}

        if cmp_val:
            multiples["CMP"] = f"\u20b9{cmp_val:,.0f}"
        if pt:
            multiples["Revised Price Target"] = f"\u20b9{pt:,.0f}"
        if upside is not None:
            sign = "+" if upside >= 0 else ""
            multiples["Upside / (Downside)"] = f"{sign}{upside:.1f}%"

        for idx, fy in enumerate(proj_fy_labels[:3]):
            if idx < len(pe_series) and pe_series[idx] is not None:
                multiples[f"P/E ({fy})"] = f"{pe_series[idx]:.1f}x"
            if idx < len(ev_series) and ev_series[idx] is not None:
                multiples[f"EV/EBITDA ({fy})"] = f"{ev_series[idx]:.1f}x"

        # Fallback: derive from projections if valuation series absent
        if len(multiples) <= 3:
            revised_estimates = q.get("revised_estimates", {})
            if revised_estimates:
                new_pat = revised_estimates.get("fy26e_pat") or \
                          revised_estimates.get("fy27e_pat")
                mktcap  = market_data.get("market_cap_cr")
                if new_pat and mktcap:
                    implied_pe = mktcap / new_pat
                    fy_key     = "FY26E" if revised_estimates.get("fy26e_pat") else "FY27E"
                    multiples[f"Implied P/E ({fy_key})"] = f"{implied_pe:.1f}x"

        if multiples:
            make_key_metrics_box(doc, multiples, cols=3, theme=theme)

        # Valuation rationale paragraph
        sub_head(doc, "Investment Rationale — Post Results", theme)
        rec_text = (narrative or {}).get("recommendation", "")
        if rec_text:
            body(doc, str(rec_text)[:600], theme)
        else:
            body(doc,
                 f"We maintain our {rating} rating on {q.get('quarter_label', '')} "
                 f"results with a revised price target of "
                 f"\u20b9{pt:,.0f}. The results were broadly in line with / above "
                 f"expectations and our investment thesis remains intact.",
                 theme)

        page_break(doc)

    # =========================================================================
    # Section 7 — Financial Summary Table
    # =========================================================================

    def _build_financial_summary(self, financials: dict, ratios: dict,
                                  projections: dict) -> None:
        """
        Last 4 quarters actuals (if available) + last 2 annual FYs + forward FYs.
        If quarterly series is not in financials, fall back to last 4 annual FYs.
        """
        doc, theme = self.doc, self.theme

        section_bar(doc, "FINANCIAL SUMMARY", theme)

        fy_labels       = (financials or {}).get("fy_labels", [])
        actuals_end_idx = (financials or {}).get("actuals_end_idx", len(fy_labels))
        actuals_end_fy  = fy_labels[actuals_end_idx - 1] if (actuals_end_idx and fy_labels) else ""

        inc = (financials or {}).get("income_statement", {})

        # Use last 5 annual FYs + up to 3 projection FYs for a compact table
        proj_fy    = (projections or {}).get("fy_labels", [])
        n_actuals  = min(5, len(fy_labels))
        show_fy    = list(fy_labels[-n_actuals:]) + list(proj_fy[:3])

        if not show_fy:
            body(doc, "Financial data not available.", theme)
            page_break(doc)
            return

        sub_head(doc, "Income Statement Summary", theme)
        if inc and show_fy:
            # Slice series to match show_fy
            def _slice(series, all_labels, show_labels):
                idx_map = {lb: i for i, lb in enumerate(all_labels)}
                proj_data = (projections or {})
                proj_map  = {lb: i for i, lb in enumerate(proj_data.get("fy_labels", []))}
                out = []
                for lb in show_labels:
                    if lb in idx_map:
                        val = series[idx_map[lb]] if idx_map[lb] < len(series) else None
                        out.append(val)
                    elif lb in proj_map:
                        # Try to find the key in projections
                        out.append(None)  # placeholder; caller handles per-key
                    else:
                        out.append(None)
                return out

            # Build merged labels and sliced series
            all_labels = list(fy_labels) + list(proj_fy)
            all_series = {}

            # Annual income statement keys
            for k in ["net_revenue", "ebitda", "ebitda_margin", "pat", "eps"]:
                annual = inc.get(k, [])
                proj   = (projections or {}).get(k, [])
                combined = list(annual) + list(proj)
                label_pool = list(fy_labels) + list(proj_fy)
                out = []
                for lb in show_fy:
                    if lb in label_pool:
                        i = label_pool.index(lb)
                        out.append(combined[i] if i < len(combined) else None)
                    else:
                        out.append(None)
                all_series[k] = out

            row_specs = [
                ("Net Revenue (Rs. Cr)",  "net_revenue",   "{:,.0f}",   True),
                ("EBITDA (Rs. Cr)",       "ebitda",        "{:,.0f}",   True),
                ("EBITDA Margin (%)",     "ebitda_margin", "{:.1f}%",   False),
                ("PAT (Rs. Cr)",          "pat",           "{:,.0f}",   True),
                ("EPS (Rs.)",             "eps",           "{:.1f}",    True),
            ]

            headers = ["Metric"] + show_fy
            rows = []
            highlight_idx = []

            for row_label, key, fmt, is_highlight in row_specs:
                series = all_series.get(key, [None] * len(show_fy))
                vals = []
                for v in series:
                    if v is None:
                        vals.append("\u2013")
                    else:
                        try:
                            vals.append(fmt.format(v))
                        except (ValueError, TypeError):
                            vals.append(str(v))
                rows.append(tuple([row_label] + vals))
                if is_highlight:
                    highlight_idx.append(len(rows) - 1)

            # Determine projection start column for shading
            proj_start_col = None
            if actuals_end_fy and actuals_end_fy in show_fy:
                proj_start_col = show_fy.index(actuals_end_fy) + 2

            make_fin_table(doc, headers, rows,
                           highlight_rows=highlight_idx,
                           proj_start_col=proj_start_col,
                           theme=theme)

        # Key ratios sub-section
        sub_head(doc, "Key Ratios", theme)
        ratio_labels_all = (ratios or {}).get("fy_labels", fy_labels)
        ratio_data       = ratios or {}

        ratio_specs = [
            ("ROE (%)",          "roe",           "{:.1f}%"),
            ("ROCE (%)",         "roce",          "{:.1f}%"),
            ("Debt / Equity",    "debt_equity",   "{:.2f}x"),
            ("P/E (x)",          "pe",            "{:.1f}"),
            ("EV/EBITDA (x)",    "ev_ebitda",     "{:.1f}"),
        ]

        ratio_show_fy = [lb for lb in show_fy if lb in ratio_labels_all]
        if ratio_show_fy:
            ratio_headers = ["Ratio"] + ratio_show_fy
            ratio_rows    = []
            for rlabel, rkey, rfmt in ratio_specs:
                rseries = ratio_data.get(rkey, [])
                rvals   = []
                for lb in ratio_show_fy:
                    if lb in ratio_labels_all:
                        idx = ratio_labels_all.index(lb)
                        v   = rseries[idx] if idx < len(rseries) else None
                    else:
                        v = None
                    if v is None:
                        rvals.append("\u2013")
                    else:
                        try:
                            rvals.append(rfmt.format(v))
                        except (ValueError, TypeError):
                            rvals.append(str(v))
                ratio_rows.append(tuple([rlabel] + rvals))

            if ratio_rows:
                make_fin_table(doc, ratio_headers, ratio_rows, theme=theme)

        page_break(doc)

    # =========================================================================
    # Section 8 — Rating and Disclosure
    # =========================================================================

    def _build_rating_disclosure(self, meta: dict, valuation: dict,
                                  quarterly, narrative: dict) -> None:
        """
        Investment thesis update paragraph and SEBI disclosure boilerplate.
        """
        doc, theme = self.doc, self.theme
        q = quarterly or {}

        section_bar(doc, "INVESTMENT THESIS UPDATE & DISCLOSURES", theme)

        # Thesis update
        sub_head(doc, "Investment Thesis — Post Results Update", theme)
        rating  = q.get("revised_rating", valuation.get("rating", "BUY"))
        pt      = q.get("revised_price_target", valuation.get("price_target", 0.0))
        company = meta.get("company_name", "the company")
        qlabel  = q.get("quarter_label", "")

        thesis_update = (narrative or {}).get("investment_summary", "")
        if thesis_update:
            summary_para = (thesis_update if isinstance(thesis_update, str)
                            else " ".join(thesis_update))
            body(doc, summary_para[:600], theme)
        else:
            body(doc,
                 f"Following {qlabel} results, we maintain our {rating} rating on "
                 f"{company} with a price target of \u20b9{pt:,.0f}. The results "
                 f"reaffirm our core investment thesis of sustained earnings growth "
                 f"and strong cash generation.",
                 theme)

        doc.add_paragraph()

        # SEBI disclosures
        section_bar(doc, "ANALYST CERTIFICATION & SEBI DISCLOSURES", theme)

        analyst = meta.get("analyst_name", "The Research Analyst")
        ticker  = meta.get("ticker", "")

        cert = (
            f"{analyst} certifies that the views expressed in this earnings update "
            f"accurately reflect the analyst's personal views about {company} ({ticker}) "
            f"and its securities. Compensation is not directly or indirectly related to "
            f"the specific recommendations or views expressed in this report."
        )
        body(doc, cert, theme)
        doc.add_paragraph()

        disclosures = [
            "This document is an earnings update note prepared for information purposes only "
            "and does not constitute an offer to buy or sell any securities.",
            "The information herein is based on publicly available data including the company's "
            "quarterly results, exchange filings, and management commentary.",
            "Past performance is not indicative of future results. Investors should seek "
            "independent financial advice before acting on this report.",
            f"The research analyst(s) may have financial interest in {company}. Full disclosure "
            "is available in the firm's disclosure document.",
            "SEBI Registration: Research Analyst — refer to firm disclosure documents for full "
            "SEBI registration details and conflict-of-interest policy.",
            "This report is for the exclusive use of the intended recipient. Redistribution or "
            "reproduction in whole or part requires prior written consent.",
        ]

        sub_head(doc, "Important Disclosures", theme)
        for disc in disclosures:
            bullet(doc, disc, level=0, theme=theme)

        doc.add_paragraph()
        src_note(doc,
                 f"Report Date: {meta.get('report_date', '')}  |  "
                 f"Analyst: {meta.get('analyst_name', '')}  |  "
                 f"Firm: {meta.get('firm_name', '')}  |  "
                 f"Quarter: {q.get('quarter_label', '')}",
                 theme)

    # =========================================================================
    # Helper methods
    # =========================================================================

    @staticmethod
    def _pct_change(actual, base) -> float | None:
        """
        Return (actual / base - 1) * 100 as a float, or None if inputs invalid.

        Parameters
        ----------
        actual : numeric or None
        base   : numeric or None  (the denominator / comparison value)

        Returns
        -------
        float | None
        """
        try:
            if actual is None or base is None:
                return None
            base_f = float(base)
            if base_f == 0.0:
                return None
            return (float(actual) / base_f - 1.0) * 100.0
        except (TypeError, ValueError, ZeroDivisionError):
            return None

    @staticmethod
    def _beat_miss_str(pct) -> str:
        """
        Convert a percentage change to a human-readable beat/miss string.

        Examples
        --------
        3.7  -> "+3.7% Beat"
        -2.1 -> "-2.1% Miss"
        None -> "N/A"
        """
        if pct is None:
            return "N/A"
        sign = "+" if pct >= 0 else ""
        label = "Beat" if pct >= 0 else "Miss"
        return f"{sign}{pct:.1f}% {label}"

    @staticmethod
    def _format_cr(val) -> str:
        """
        Format a value in Crores with the Rupee sign.

        Examples
        --------
        19500  -> "\u20b919,500 Cr"
        None   -> "N/A"
        """
        if val is None:
            return "N/A"
        try:
            return f"\u20b9{float(val):,.0f} Cr"
        except (TypeError, ValueError):
            return "N/A"

    @staticmethod
    def _prior_year_quarter_label(qlabel: str) -> str:
        """
        Derive the prior-year quarter label.
        "Q3FY25" -> "Q3FY24"
        Returns qlabel unchanged if parsing fails.
        """
        if not qlabel:
            return ""
        import re
        m = re.match(r"(Q[1-4]FY)(\d+)", qlabel, re.IGNORECASE)
        if m:
            fy_year = int(m.group(2))
            return f"{m.group(1)}{fy_year - 1:02d}"
        return qlabel

    @staticmethod
    def _prior_quarter_label(qlabel: str) -> str:
        """
        Derive the immediately prior quarter label.
        "Q3FY25" -> "Q2FY25"
        "Q1FY25" -> "Q4FY24"
        Returns qlabel unchanged if parsing fails.
        """
        if not qlabel:
            return ""
        import re
        m = re.match(r"Q([1-4])FY(\d+)", qlabel, re.IGNORECASE)
        if m:
            q_num  = int(m.group(1))
            fy_num = int(m.group(2))
            if q_num == 1:
                return f"Q4FY{fy_num - 1:02d}"
            return f"Q{q_num - 1}FY{fy_num:02d}"
        return qlabel


# ---------------------------------------------------------------------------
# Quick smoke-test — run directly as a script
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    import tempfile

    _SAMPLE_JSON = {
        "meta": {
            "ticker": "ITC",
            "company_name": "ITC Limited",
            "sector": "FMCG",
            "model_type": "fmcg",
            "report_date": "2026-03-18",
            "analyst_name": "Equity Research Team",
            "firm_name": "IndiaStockResearch",
        },
        "market_data": {
            "cmp": 307.0,
            "market_cap_cr": 387153,
            "52w_high": 444.0,
            "52w_low": 272.0,
        },
        "financials": {
            "fy_labels": ["FY21", "FY22", "FY23", "FY24", "FY25"],
            "actuals_end_fy": "FY25",
            "actuals_end_idx": 5,
            "income_statement": {
                "net_revenue": [49327, 59186, 69446, 73928, 80002],
                "ebitda":      [17400, 20100, 22300, 24500, 26800],
                "ebitda_margin": [35.3, 34.0, 32.1, 33.1, 33.5],
                "pat":         [13032, 15159, 18753, 20457, 21800],
                "eps":         [10.4,  12.0,  14.9,  16.2,  17.3],
            },
            "balance_sheet": {},
            "cash_flow": {},
        },
        "ratios": {
            "fy_labels": ["FY21", "FY22", "FY23", "FY24", "FY25"],
            "ebitda_margin": [35.3, 34.0, 32.1, 33.1, 33.5],
            "pat_margin":    [26.4, 25.6, 27.0, 27.7, 27.2],
            "roce":          [32.1, 34.0, 35.5, 36.2, 35.9],
            "pe":            [22.0, 20.5, 19.0, 18.2, 17.5],
            "ev_ebitda":     [16.0, 15.0, 13.5, 12.8, 12.2],
            "roe":           [24.5, 25.1, 26.3, 27.0, 26.8],
            "debt_equity":   [0.01, 0.01, 0.01, 0.01, 0.01],
        },
        "projections": {
            "fy_labels":   ["FY26E", "FY27E", "FY28E"],
            "net_revenue": [87000, 95000, 103000],
            "pat":         [23500, 26000, 28500],
            "ebitda":      [29000, 32000, 35000],
            "ebitda_margin": [33.3, 33.7, 34.0],
            "eps":         [18.6, 20.6, 22.6],
            "fcf":         [20500, 23000, 25500],
        },
        "valuation": {
            "rating": "BUY",
            "price_target": 380.0,
            "cmp": 307.0,
            "upside_pct": 23.8,
            "methodology": "25x FY27E EPS",
            "proj_fy_labels": ["FY26E", "FY27E", "FY28E"],
            "pe_series":       [16.5, 14.9, 13.6],
            "ev_ebitda_series": [12.1, 10.9,  9.9],
        },
        "charts": {},
        "narrative": {
            "investment_summary": (
                "Market share gains across cigarettes and FMCG categories. "
                "Agri-business recovery adding to earnings visibility. "
                "Hotels segment on a multi-year upcycle."
            ),
            "financial_analysis": (
                "ITC delivered a strong Q3FY25 with revenue growth of 14% YoY, "
                "driven by cigarette volume growth of 8% and robust FMCG sales. "
                "EBITDA margins expanded 40bps sequentially aided by benign leaf tobacco "
                "prices and operating leverage in the FMCG segment."
            ),
        },
        "quarterly": {
            "quarter_label": "Q3FY25",
            "actuals": {
                "revenue": 19500,
                "ebitda": 5200,
                "ebitda_margin": 26.7,
                "pat": 4900,
                "eps": 3.90,
                "depreciation": 350,
                "ebit": 4850,
                "other_income": 420,
                "pbt": 5270,
                "tax": 370,
            },
            "estimates": {
                "revenue": 18800,
                "ebitda": 4950,
                "ebitda_margin": 26.3,
                "pat": 4700,
                "eps": 3.75,
                "depreciation": 345,
                "ebit": 4605,
                "other_income": 400,
                "pbt": 5005,
                "tax": 305,
            },
            "prior_quarter": {
                "revenue": 18200,
                "ebitda": 4800,
                "ebitda_margin": 26.4,
                "pat": 4500,
                "eps": 3.60,
            },
            "prior_year_quarter": {
                "revenue": 17100,
                "ebitda": 4500,
                "ebitda_margin": 26.3,
                "pat": 4200,
                "eps": 3.36,
            },
            "beat_miss_revenue_pct": 3.7,
            "beat_miss_pat_pct": 4.3,
            "management_commentary": [
                "Cigarette volume growth strong at 8% YoY — sustained pricing power.",
                "FMCG (ex-cigarettes) revenue grew 12% YoY with margin improvement.",
                "Rural recovery clearly visible — distributor fill-rates back to pre-COVID levels.",
                "Input cost tailwinds from lower leaf tobacco prices expected to persist through Q4.",
                "Hotels segment posted best-ever quarterly revenue on strong RevPAR growth.",
                "Board approved Rs. 3,500 Cr buyback at Rs. 350/share.",
            ],
            "revised_estimates": {
                "fy26e_pat": 24200,
                "fy27e_pat": 26500,
                "fy26e_revenue": 88500,
                "fy27e_revenue": 96800,
            },
            "revised_price_target": 395,
            "revised_rating": "BUY",
        },
    }

    out_path = "/tmp/phase3/ITC_Q3FY25_Earnings_Test.docx"
    builder  = EarningsBuilder(output_path=out_path)
    result   = builder.build(_SAMPLE_JSON)
    print(f"Earnings report saved to: {result}")
