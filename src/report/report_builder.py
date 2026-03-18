"""
Generic DOCX report builder for Indian equity research reports.
Driven by a JSON data contract — no company-specific logic.

Usage:
    from src.report.report_builder import ReportBuilder

    builder = ReportBuilder(output_path="reports/ITC/ITC_Report_2026.docx")
    builder.build(report_json)  # report_json is the full JSON data contract
    # -> saves DOCX to output_path
"""

import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

# Graceful imports — fall back to stubs when running outside the full project tree
try:
    from src.report.docx_helpers import (
        ReportTheme, DEFAULT_THEME, sfont, shade_cell, cell_text,
        section_bar, sub_head, body, bullet, src_note, fig_caption,
        add_chart, make_fin_table, make_key_metrics_box, make_rating_box,
        page_break, set_margins,
    )
except ImportError:
    from docx_helpers import (  # type: ignore
        ReportTheme, DEFAULT_THEME, sfont, shade_cell, cell_text,
        section_bar, sub_head, body, bullet, src_note, fig_caption,
        add_chart, make_fin_table, make_key_metrics_box, make_rating_box,
        page_break, set_margins,
    )

try:
    from src.utils.fy_utils import label_with_suffix, fy_end_year  # type: ignore
except ImportError:
    def label_with_suffix(label, actuals_end_fy=""):  # type: ignore
        return label

    def fy_end_year(fy_label):  # type: ignore
        return fy_label

try:
    from config.settings import FONT_NAME, REPORTS_DIR  # type: ignore
except ImportError:
    FONT_NAME = "Times New Roman"
    REPORTS_DIR = "reports"

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from enum import Enum


# ── Section enum ─────────────────────────────────────────────────────────────

class ReportSection(Enum):
    COVER = "cover"
    INVESTMENT_SUMMARY = "investment_summary"
    INVESTMENT_THESIS = "investment_thesis"
    KEY_RISKS = "key_risks"
    COMPANY_101 = "company_101"
    PRODUCTS_INDUSTRY = "products_industry"
    FINANCIAL_ANALYSIS = "financial_analysis"
    VALUATION = "valuation"
    CATALYSTS = "catalysts"
    APPENDIX = "appendix"
    DISCLOSURES = "disclosures"


# ── ReportBuilder ─────────────────────────────────────────────────────────────

class ReportBuilder:
    """Builds a complete institutional equity research DOCX from a JSON contract."""

    def __init__(self, output_path: str, theme: ReportTheme = None):
        self.output_path = output_path
        self.theme = theme or DEFAULT_THEME
        self.doc = Document()

    # ── Main entry point ─────────────────────────────────────────────────────

    def build(self, report_json: dict) -> str:
        """
        Main entry point. Builds complete DOCX from report_json.
        Calls each section method in order.
        Returns output_path.
        """
        # Extract top-level sections with safe defaults
        meta = report_json.get("meta", {})
        market_data = report_json.get("market_data", {})
        financials = report_json.get("financials", {})
        ratios = report_json.get("ratios", {})
        shareholding = report_json.get("shareholding", {})
        projections = report_json.get("projections", {})
        valuation = report_json.get("valuation", {})
        charts = report_json.get("charts", {})
        narrative = report_json.get("narrative", {})

        # Apply margins
        set_margins(self.doc, top=1.0, bottom=1.0, left=1.0, right=1.0)

        # Build each section in order
        self._build_cover(meta, market_data, valuation)
        self._build_investment_summary(meta, market_data, financials, valuation,
                                       narrative, charts)
        self._build_investment_thesis(narrative, charts)
        self._build_key_risks(narrative, charts)
        self._build_company_101(meta, narrative, charts)
        self._build_products_industry(narrative, charts)
        self._build_financial_analysis(financials, ratios, projections,
                                        narrative, charts)
        self._build_valuation(valuation, narrative, charts)
        self._build_catalysts(narrative)
        self._build_appendix(narrative)
        self._build_disclosures(meta)

        # Save document
        os.makedirs(os.path.dirname(os.path.abspath(self.output_path)), exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path

    # ── Section builders ─────────────────────────────────────────────────────

    def _build_cover(self, meta: dict, market_data: dict, valuation: dict) -> None:
        """
        Cover page:
        - Company name (large, NAVY, centered)
        - "Initiating Coverage | [Sector] | [Exchange]: [Ticker]"
        - Rating box (BUY/HOLD/SELL, price target, CMP, upside)
        - Key stats row: Market Cap, Revenue, EBITDA, P/E, EV/EBITDA
        - Report date and analyst name
        - Page break
        """
        doc, theme = self.doc, self.theme

        # Company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta.get("company_name", "Company Name"))
        sfont(run, 20.0, bold=True, color_hex=theme.navy, theme=theme)

        # Sub-title line
        ticker = meta.get("ticker", "TICK")
        sector = meta.get("sector", "Sector")
        exchange = meta.get("exchange", "NSE")
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(
            f"Initiating Coverage  |  {sector}  |  {exchange}: {ticker}"
        )
        sfont(sub_run, 10.0, italic=True, color_hex=theme.dgray, theme=theme)

        doc.add_paragraph()

        # Rating box
        rating = valuation.get("rating", "BUY")
        pt = valuation.get("price_target", 0.0)
        cmp = market_data.get("cmp", 0.0)
        upside = ((pt - cmp) / cmp * 100) if cmp else 0.0
        upside = valuation.get("upside_pct", upside)
        make_rating_box(doc, rating, pt, cmp, upside, theme)

        # Key stats row
        stats = {
            "Market Cap": f"\u20b9{market_data.get('market_cap_cr', 0):,.0f} Cr",
            "Revenue (TTM)": f"\u20b9{market_data.get('revenue_cr', 0):,.0f} Cr",
            "EBITDA (TTM)": f"\u20b9{market_data.get('ebitda_cr', 0):,.0f} Cr",
            "P/E (TTM)": f"{market_data.get('pe_ttm', 0):.1f}x",
            "EV/EBITDA": f"{market_data.get('ev_ebitda', 0):.1f}x",
        }
        make_key_metrics_box(doc, stats, cols=5, theme=theme)

        # Report date and analyst
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_str = meta.get("report_date", "")
        analyst = meta.get("analyst_name", "")
        run2 = p2.add_run(f"Date: {date_str}   |   Analyst: {analyst}")
        sfont(run2, 8.0, italic=True, color_hex=theme.black, theme=theme)

        page_break(doc)

    def _build_investment_summary(self, meta: dict, market_data: dict,
                                   financials: dict, valuation: dict,
                                   narrative: dict, charts: dict) -> None:
        """
        Investment Summary (Page 1):
        - Rating box
        - 4-5 investment thesis bullet points (from narrative.investment_summary)
        - Financial summary table: last 3 actual FY + 2 projected FY
        - chart_02 (revenue growth trajectory) if available
        - Key metrics box: Market Cap, CMP, 52w H/L, Dividend Yield, Beta
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "INVESTMENT SUMMARY", theme)

        # Rating box (compact repeat)
        rating = valuation.get("rating", "BUY")
        pt = valuation.get("price_target", 0.0)
        cmp = market_data.get("cmp", 0.0)
        upside = valuation.get("upside_pct",
                               ((pt - cmp) / cmp * 100) if cmp else 0.0)
        make_rating_box(doc, rating, pt, cmp, upside, theme)

        # Investment thesis bullets
        sub_head(doc, "Investment Thesis", theme)
        summary_text = narrative.get("investment_summary", "")
        if summary_text:
            # Support both a plain string (split on newlines) and a list
            if isinstance(summary_text, list):
                points = summary_text
            else:
                points = [ln.strip(" -\u2022") for ln in summary_text.splitlines()
                          if ln.strip()]
            for pt_text in points[:5]:
                bullet(doc, pt_text, level=0, theme=theme)
        else:
            bullet(doc, "Key investment point to be added.", theme=theme)

        # Financial summary table
        sub_head(doc, "Financial Snapshot", theme)
        fy_labels = financials.get("fy_labels", [])
        actuals_end_idx = financials.get("actuals_end_idx", len(fy_labels))
        inc = financials.get("income_statement", {})

        summary_keys = ["net_revenue", "ebitda", "ebitda_margin", "pat", "eps"]
        summary_labels = {
            "net_revenue": "Net Revenue (Rs. Cr)",
            "ebitda": "EBITDA (Rs. Cr)",
            "ebitda_margin": "EBITDA Margin (%)",
            "pat": "PAT (Rs. Cr)",
            "eps": "EPS (Rs.)",
        }
        fmt_map = {
            "ebitda_margin": "{:.1f}%",
            "eps": "{:.1f}",
        }

        if fy_labels and inc:
            self._add_financial_table_from_json(
                fin_dict=inc,
                row_keys=summary_keys,
                row_labels=summary_labels,
                fy_labels=fy_labels,
                actuals_end_fy=fy_labels[actuals_end_idx - 1] if actuals_end_idx else "",
                highlight_keys=["net_revenue", "pat"],
                fmt_map=fmt_map,
            )

        # Valuation multiples rows
        val_rows = []
        pe_list = valuation.get("pe_series", [])
        ev_list = valuation.get("ev_ebitda_series", [])
        if pe_list:
            val_rows.append(("P/E (x)",) + tuple(
                f"{v:.1f}" if v is not None else "\u2013" for v in pe_list
            ))
        if ev_list:
            val_rows.append(("EV/EBITDA (x)",) + tuple(
                f"{v:.1f}" if v is not None else "\u2013" for v in ev_list
            ))
        if val_rows and fy_labels:
            make_fin_table(doc, ["Metric"] + fy_labels, val_rows, theme=theme)

        # Chart 02 — revenue growth trajectory
        self._embed_chart(charts, "chart_02", width=6.5,
                          caption="Revenue Growth Trajectory")

        # Key metrics box
        sub_head(doc, "Key Market Data", theme)
        h52 = market_data.get("52w_high", 0)
        l52 = market_data.get("52w_low", 0)
        metrics = {
            "Market Cap": f"\u20b9{market_data.get('market_cap_cr', 0):,.0f} Cr",
            "CMP": f"\u20b9{market_data.get('cmp', 0):,.0f}",
            "52W High / Low": f"\u20b9{h52:,.0f} / \u20b9{l52:,.0f}",
            "Dividend Yield": f"{market_data.get('dividend_yield', 0):.2f}%",
            "Beta": f"{market_data.get('beta', 0):.2f}",
        }
        make_key_metrics_box(doc, metrics, cols=3, theme=theme)
        page_break(doc)

    def _build_investment_thesis(self, narrative: dict, charts: dict) -> None:
        """
        5 investment pillars with supporting text.
        Each pillar: sub_head + 2-3 body paragraphs + relevant chart if available.
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "INVESTMENT THESIS", theme)

        pillars = narrative.get("investment_pillars", [])
        if not pillars:
            # Fallback to generic text sections
            pillars = [
                {
                    "title": "Pillar 1: Strong Revenue Visibility",
                    "text": narrative.get("thesis_pillar_1", ""),
                    "chart": "chart_03",
                },
            ]

        for i, pillar in enumerate(pillars, start=1):
            title = pillar.get("title", f"Pillar {i}")
            text = pillar.get("text", "")
            chart_key = pillar.get("chart", f"chart_0{i + 2}")

            sub_head(doc, title, theme)
            if text:
                if isinstance(text, list):
                    for para in text:
                        body(doc, para, theme)
                else:
                    for para in text.splitlines():
                        if para.strip():
                            body(doc, para.strip(), theme)
            else:
                body(doc, "Supporting analysis to be added.", theme)

            self._embed_chart(charts, chart_key, width=6.5)

        page_break(doc)

    def _build_key_risks(self, narrative: dict, charts: dict) -> None:
        """
        Risk section: 5-8 risks with probability/impact table.
        Table: Risk | Category | Probability | Impact | Mitigation
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "KEY RISKS", theme)

        body(doc, narrative.get("risks_intro", ""), theme)

        risks = narrative.get("risks", [])
        if risks:
            headers = ["Risk Factor", "Category", "Probability", "Impact", "Mitigation"]
            rows = []
            highlight_rows = []
            for risk in risks:
                rows.append((
                    risk.get("title", ""),
                    risk.get("category", ""),
                    risk.get("probability", ""),
                    risk.get("impact", ""),
                    risk.get("mitigation", ""),
                ))
            col_widths = [1.8, 1.0, 0.9, 0.9, 1.9]
            make_fin_table(doc, headers, rows,
                           col_widths=col_widths,
                           highlight_rows=highlight_rows,
                           theme=theme)
        else:
            body(doc, "Risk details to be populated.", theme)

        self._embed_chart(charts, "chart_risks", width=6.5)
        page_break(doc)

    def _build_company_101(self, meta: dict, narrative: dict, charts: dict) -> None:
        """
        Company overview, history, management bios, governance.
        Uses narrative.company_overview, narrative.management_bios, etc.
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, f"COMPANY OVERVIEW: {meta.get('company_name', '')}", theme)

        sub_head(doc, "Business Overview", theme)
        overview = narrative.get("company_overview", narrative.get("business_overview", ""))
        if overview:
            for para in (overview if isinstance(overview, list) else [overview]):
                body(doc, para, theme)

        self._embed_chart(charts, "chart_01", width=6.5,
                          caption="Business Segments Overview")

        sub_head(doc, "Company History & Milestones", theme)
        history = narrative.get("company_history", "")
        if history:
            body(doc, history, theme)

        sub_head(doc, "Management Team", theme)
        mgmt_bios = narrative.get("management_bios", [])
        if mgmt_bios:
            for bio in mgmt_bios:
                if isinstance(bio, dict):
                    name = bio.get("name", "")
                    title = bio.get("title", "")
                    desc = bio.get("description", "")
                    p = doc.add_paragraph()
                    run = p.add_run(f"{name} — {title}")
                    sfont(run, float(theme.font_size_body), bold=True,
                          color_hex=theme.navy, theme=theme)
                    if desc:
                        body(doc, desc, theme)
                else:
                    body(doc, str(bio), theme)
        else:
            body(doc, "Management profiles to be added.", theme)

        sub_head(doc, "Corporate Governance", theme)
        governance = narrative.get("governance", "")
        if governance:
            body(doc, governance, theme)

        page_break(doc)

    def _build_products_industry(self, narrative: dict, charts: dict) -> None:
        """
        Products/services, customer channels, industry overview,
        competitive landscape, TAM.
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "PRODUCTS, INDUSTRY & COMPETITIVE LANDSCAPE", theme)

        sub_head(doc, "Products & Services", theme)
        products = narrative.get("products_services", "")
        if isinstance(products, list):
            for item in products:
                bullet(doc, item, theme=theme)
        elif products:
            body(doc, products, theme)

        self._embed_chart(charts, "chart_products", width=6.5,
                          caption="Revenue Mix by Segment")

        sub_head(doc, "Customer Channels & Distribution", theme)
        channels = narrative.get("customer_channels", "")
        if channels:
            body(doc, channels, theme)

        sub_head(doc, "Industry Overview & TAM", theme)
        industry = narrative.get("industry_overview", "")
        if industry:
            body(doc, industry, theme)

        self._embed_chart(charts, "chart_industry", width=6.5,
                          caption="Industry Growth Trend")

        sub_head(doc, "Competitive Landscape", theme)
        competition = narrative.get("competitive_landscape", "")
        if competition:
            body(doc, competition, theme)

        peers = narrative.get("competitor_table", [])
        if peers:
            headers = ["Company", "Market Cap (Cr)", "Revenue (Cr)", "EBITDA Margin", "P/E"]
            rows = [
                (
                    p.get("name", ""),
                    p.get("market_cap", ""),
                    p.get("revenue", ""),
                    p.get("ebitda_margin", ""),
                    p.get("pe", ""),
                )
                for p in peers
            ]
            make_fin_table(doc, headers, rows, theme=theme)

        page_break(doc)

    def _build_financial_analysis(self, financials: dict, ratios: dict,
                                   projections: dict, narrative: dict,
                                   charts: dict) -> None:
        """
        Historical income statement table (5-10yr).
        Balance sheet table.
        Cash flow table.
        Key ratios table.
        Projection table (5yr forward).
        Scenario comparison table.
        Charts embedded throughout.
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "FINANCIAL ANALYSIS", theme)

        fy_labels = financials.get("fy_labels", [])
        actuals_end_idx = financials.get("actuals_end_idx", len(fy_labels))
        actuals_end_fy = fy_labels[actuals_end_idx - 1] if actuals_end_idx and fy_labels else ""

        # Income statement
        sub_head(doc, "Income Statement", theme)
        inc = financials.get("income_statement", {})
        if inc and fy_labels:
            inc_keys = [
                "net_revenue", "gross_profit", "ebitda", "ebitda_margin",
                "depreciation", "ebit", "interest", "pbt", "tax", "pat",
            ]
            inc_labels = {
                "net_revenue": "Net Revenue (Rs. Cr)",
                "gross_profit": "Gross Profit (Rs. Cr)",
                "ebitda": "EBITDA (Rs. Cr)",
                "ebitda_margin": "EBITDA Margin (%)",
                "depreciation": "Depreciation (Rs. Cr)",
                "ebit": "EBIT (Rs. Cr)",
                "interest": "Interest Expense (Rs. Cr)",
                "pbt": "PBT (Rs. Cr)",
                "tax": "Tax (Rs. Cr)",
                "pat": "PAT (Rs. Cr)",
            }
            self._add_financial_table_from_json(
                inc, inc_keys, inc_labels, fy_labels, actuals_end_fy,
                highlight_keys=["net_revenue", "ebitda", "pat"],
                fmt_map={"ebitda_margin": "{:.1f}%"},
            )

        self._embed_chart(charts, "chart_income", width=6.5,
                          caption="Revenue and EBITDA Trend")

        # Balance sheet
        sub_head(doc, "Balance Sheet", theme)
        bs = financials.get("balance_sheet", {})
        if bs and fy_labels:
            bs_keys = [
                "total_assets", "fixed_assets", "current_assets", "cash",
                "total_equity", "total_debt", "current_liabilities",
            ]
            bs_labels = {
                "total_assets": "Total Assets (Rs. Cr)",
                "fixed_assets": "Net Fixed Assets (Rs. Cr)",
                "current_assets": "Current Assets (Rs. Cr)",
                "cash": "Cash & Equivalents (Rs. Cr)",
                "total_equity": "Total Equity (Rs. Cr)",
                "total_debt": "Total Debt (Rs. Cr)",
                "current_liabilities": "Current Liabilities (Rs. Cr)",
            }
            self._add_financial_table_from_json(
                bs, bs_keys, bs_labels, fy_labels, actuals_end_fy,
                highlight_keys=["total_assets", "total_equity"],
            )

        # Cash flow
        sub_head(doc, "Cash Flow Statement", theme)
        cf = financials.get("cash_flow", {})
        if cf and fy_labels:
            cf_keys = [
                "cfo", "capex", "fcf", "cfi", "cff", "net_change_cash",
            ]
            cf_labels = {
                "cfo": "Cash from Operations (Rs. Cr)",
                "capex": "Capex (Rs. Cr)",
                "fcf": "Free Cash Flow (Rs. Cr)",
                "cfi": "Cash from Investing (Rs. Cr)",
                "cff": "Cash from Financing (Rs. Cr)",
                "net_change_cash": "Net Change in Cash (Rs. Cr)",
            }
            self._add_financial_table_from_json(
                cf, cf_keys, cf_labels, fy_labels, actuals_end_fy,
                highlight_keys=["cfo", "fcf"],
            )

        # Key ratios
        sub_head(doc, "Key Financial Ratios", theme)
        ratio_labels_all = fy_labels
        if ratios:
            ratio_keys = [
                "gross_margin", "ebitda_margin", "pat_margin",
                "roe", "roce", "roa",
                "debt_equity", "current_ratio", "asset_turnover",
                "eps", "bvps", "pe", "ev_ebitda",
            ]
            ratio_labels = {
                "gross_margin": "Gross Margin (%)",
                "ebitda_margin": "EBITDA Margin (%)",
                "pat_margin": "PAT Margin (%)",
                "roe": "ROE (%)",
                "roce": "ROCE (%)",
                "roa": "ROA (%)",
                "debt_equity": "Debt / Equity (x)",
                "current_ratio": "Current Ratio (x)",
                "asset_turnover": "Asset Turnover (x)",
                "eps": "EPS (Rs.)",
                "bvps": "Book Value / Share (Rs.)",
                "pe": "P/E (x)",
                "ev_ebitda": "EV/EBITDA (x)",
            }
            ratio_fy = ratios.get("fy_labels", ratio_labels_all)
            self._add_financial_table_from_json(
                ratios, ratio_keys, ratio_labels, ratio_fy, actuals_end_fy,
                fmt_map={
                    k: "{:.1f}%" for k in
                    ["gross_margin", "ebitda_margin", "pat_margin", "roe", "roce", "roa"]
                },
            )

        self._embed_chart(charts, "chart_ratios", width=6.5,
                          caption="Key Profitability Ratios")

        # Projections table
        sub_head(doc, "Financial Projections", theme)
        if projections:
            proj_fy = projections.get("fy_labels", [])
            proj_keys = [
                "net_revenue", "ebitda", "ebitda_margin", "pat", "eps",
                "pe", "ev_ebitda",
            ]
            proj_labels = {
                "net_revenue": "Net Revenue (Rs. Cr)",
                "ebitda": "EBITDA (Rs. Cr)",
                "ebitda_margin": "EBITDA Margin (%)",
                "pat": "PAT (Rs. Cr)",
                "eps": "EPS (Rs.)",
                "pe": "P/E (x)",
                "ev_ebitda": "EV/EBITDA (x)",
            }
            if proj_fy:
                self._add_financial_table_from_json(
                    projections, proj_keys, proj_labels,
                    proj_fy, actuals_end_fy="",
                    highlight_keys=["net_revenue", "pat"],
                    fmt_map={"ebitda_margin": "{:.1f}%", "eps": "{:.1f}"},
                )

        # Scenario comparison
        sub_head(doc, "Scenario Analysis", theme)
        narrative_text = narrative.get("financial_analysis", "")
        if narrative_text:
            body(doc, narrative_text, theme)

        page_break(doc)

    def _build_valuation(self, valuation: dict, narrative: dict, charts: dict) -> None:
        """
        DCF analysis table (SOTP or consolidated).
        DCF sensitivity table.
        Comparable companies table with statistical summary.
        Valuation football field chart.
        Final recommendation box.
        """
        doc, theme = self.doc, self.theme
        section_bar(doc, "VALUATION", theme)

        sub_head(doc, "Valuation Methodology", theme)
        val_narrative = narrative.get("valuation", "")
        if val_narrative:
            body(doc, val_narrative, theme)

        # DCF / SOTP table
        sub_head(doc, "DCF / SOTP Analysis", theme)
        dcf = valuation.get("dcf", {})
        if dcf:
            sotp_rows = dcf.get("sotp_rows", [])
            if sotp_rows:
                headers = ["Segment", "EBITDA (Cr)", "EV/EBITDA (x)",
                           "Enterprise Value (Cr)", "% of Total"]
                rows = [
                    (
                        r.get("segment", ""),
                        self._format_value(r.get("ebitda")),
                        self._format_value(r.get("multiple"), "{:.1f}"),
                        self._format_value(r.get("ev")),
                        self._format_value(r.get("pct"), "{:.1f}%"),
                    )
                    for r in sotp_rows
                ]
                make_fin_table(doc, headers, rows, theme=theme)
            else:
                # Consolidated DCF summary
                dcf_display = [
                    ("WACC (%)", self._format_value(dcf.get("wacc"), "{:.1f}%")),
                    ("Terminal Growth Rate (%)", self._format_value(dcf.get("tgr"), "{:.1f}%")),
                    ("Intrinsic Value (Rs.)", self._format_value(dcf.get("intrinsic_value"))),
                    ("Equity Value (Rs. Cr)", self._format_value(dcf.get("equity_value"))),
                ]
                headers = ["Parameter", "Value"]
                make_fin_table(doc, headers, dcf_display, theme=theme)

        # Sensitivity table
        sub_head(doc, "DCF Sensitivity Analysis", theme)
        sensitivity = valuation.get("dcf", {}).get("sensitivity", {})
        if sensitivity:
            wacc_range = sensitivity.get("wacc_range", [])
            tgr_range = sensitivity.get("tgr_range", [])
            matrix = sensitivity.get("matrix", [])
            if wacc_range and tgr_range and matrix:
                headers = ["WACC \\ TGR"] + [f"{t:.1f}%" for t in tgr_range]
                rows = []
                for i, wacc_val in enumerate(wacc_range):
                    row = [f"{wacc_val:.1f}%"] + [
                        self._format_value(matrix[i][j]) if i < len(matrix) and j < len(matrix[i]) else "\u2013"
                        for j in range(len(tgr_range))
                    ]
                    rows.append(tuple(row))
                make_fin_table(doc, headers, rows, theme=theme)

        # Comparable companies table
        sub_head(doc, "Comparable Companies", theme)
        peers = valuation.get("peers", [])
        if peers:
            headers = ["Company", "Mkt Cap (Cr)", "P/E (x)", "EV/EBITDA (x)",
                       "Revenue CAGR (%)", "EBITDA Margin (%)"]
            rows = [
                (
                    p.get("name", ""),
                    self._format_value(p.get("market_cap")),
                    self._format_value(p.get("pe"), "{:.1f}"),
                    self._format_value(p.get("ev_ebitda"), "{:.1f}"),
                    self._format_value(p.get("revenue_cagr"), "{:.1f}%"),
                    self._format_value(p.get("ebitda_margin"), "{:.1f}%"),
                )
                for p in peers
            ]
            # Add statistical summary rows
            pe_vals = [p.get("pe") for p in peers if p.get("pe") is not None]
            evebitda_vals = [p.get("ev_ebitda") for p in peers if p.get("ev_ebitda") is not None]
            if pe_vals and evebitda_vals:
                rows.append(("Median",
                             "\u2013",
                             self._format_value(sorted(pe_vals)[len(pe_vals) // 2], "{:.1f}"),
                             self._format_value(sorted(evebitda_vals)[len(evebitda_vals) // 2], "{:.1f}"),
                             "\u2013", "\u2013"))
            make_fin_table(doc, headers, rows,
                           highlight_rows=[len(rows) - 1] if pe_vals and evebitda_vals else [],
                           theme=theme)

        self._embed_chart(charts, "chart_football_field", width=6.5,
                          caption="Valuation Football Field")

        # Final recommendation
        sub_head(doc, "Investment Recommendation", theme)
        rating = valuation.get("rating", "BUY")
        pt = valuation.get("price_target", 0.0)
        cmp_val = valuation.get("cmp", 0.0)
        upside = valuation.get("upside_pct", 0.0)
        make_rating_box(doc, rating, pt, cmp_val, upside, theme)

        rec_text = narrative.get("recommendation", "")
        if rec_text:
            body(doc, rec_text, theme)

        page_break(doc)

    def _build_catalysts(self, narrative: dict) -> None:
        """5 key catalysts with timeframe and impact description."""
        doc, theme = self.doc, self.theme
        section_bar(doc, "KEY CATALYSTS", theme)

        catalysts = narrative.get("catalysts", [])
        if catalysts:
            headers = ["Catalyst", "Timeframe", "Potential Impact"]
            rows = []
            for cat in catalysts[:5]:
                if isinstance(cat, dict):
                    rows.append((
                        cat.get("title", ""),
                        cat.get("timeframe", ""),
                        cat.get("impact", ""),
                    ))
                else:
                    rows.append((str(cat), "", ""))
            col_widths = [2.5, 1.2, 2.8]
            make_fin_table(doc, headers, rows, col_widths=col_widths, theme=theme)
        else:
            catalyst_text = narrative.get("catalysts_text", "")
            if catalyst_text:
                body(doc, catalyst_text, theme)
            else:
                body(doc, "Key catalysts to be identified.", theme)

        page_break(doc)

    def _build_appendix(self, narrative: dict) -> None:
        """Extended risks, modelling assumptions table, data sources."""
        doc, theme = self.doc, self.theme
        section_bar(doc, "APPENDIX", theme)

        sub_head(doc, "Extended Risk Analysis", theme)
        ext_risks = narrative.get("extended_risks", "")
        if ext_risks:
            body(doc, ext_risks, theme)

        sub_head(doc, "Modelling Assumptions", theme)
        assumptions = narrative.get("model_assumptions", [])
        if assumptions:
            if isinstance(assumptions, list):
                headers = ["Assumption", "Value / Range", "Rationale"]
                rows = [
                    (
                        a.get("name", ""),
                        a.get("value", ""),
                        a.get("rationale", ""),
                    )
                    for a in assumptions if isinstance(a, dict)
                ]
                if rows:
                    make_fin_table(doc, headers, rows, theme=theme)
                else:
                    for a in assumptions:
                        bullet(doc, str(a), theme=theme)
            else:
                body(doc, str(assumptions), theme)

        sub_head(doc, "Data Sources", theme)
        sources = narrative.get("data_sources", [])
        if sources:
            for src in sources:
                src_note(doc, str(src), theme)
        else:
            src_note(doc, "Sources: Company filings, NSE/BSE, Bloomberg, CMIE Prowess.", theme)

    def _build_disclosures(self, meta: dict) -> None:
        """Standard analyst certification and important disclosures."""
        doc, theme = self.doc, self.theme
        page_break(doc)
        section_bar(doc, "ANALYST CERTIFICATION & IMPORTANT DISCLOSURES", theme)

        analyst = meta.get("analyst_name", "The Research Analyst")
        company = meta.get("company_name", "the subject company")
        ticker = meta.get("ticker", "")

        cert_text = (
            f"{analyst} hereby certifies that the views expressed in this research report "
            f"accurately reflect the analyst's personal views about the subject company ({company} "
            f"[{ticker}]) and its securities. The analyst's compensation is not directly or indirectly "
            f"related to the specific recommendations or views expressed in this report."
        )
        body(doc, cert_text, theme)
        doc.add_paragraph()

        disclosures = [
            "This document has been prepared by the research team for information purposes only.",
            "This report does not constitute an offer to sell or a solicitation of an offer to buy "
            "any securities mentioned herein.",
            "The information contained herein has been obtained from sources believed to be reliable "
            "but is not necessarily complete and its accuracy cannot be guaranteed.",
            "Past performance is not indicative of future results. Investors should obtain independent "
            "financial advice before acting on any information herein.",
            f"The research analyst(s) and/or their associates may have financial interest in {company}.",
            "SEBI Registration: Research Analyst — refer to firm disclosure documents for full details.",
        ]
        sub_head(doc, "Important Disclosures", theme)
        for disc in disclosures:
            bullet(doc, disc, level=0, theme=theme)

        doc.add_paragraph()
        src_note(doc,
                 f"Report Date: {meta.get('report_date', '')}  |  "
                 f"Analyst: {meta.get('analyst_name', '')}  |  "
                 f"Firm: {meta.get('firm_name', '')}",
                 theme)

    # ── Generic table helpers ────────────────────────────────────────────────

    def _add_financial_table_from_json(self, fin_dict: dict,
                                        row_keys: list,
                                        row_labels: dict,
                                        fy_labels: list,
                                        actuals_end_fy: str,
                                        highlight_keys: list = None,
                                        fmt_map: dict = None) -> None:
        """
        Build a financial table from the JSON data contract.

        row_keys: keys to include from fin_dict
        row_labels: {key: display_label} mapping
        fy_labels: column headers
        actuals_end_fy: FY label where actuals end (for projection shading)
        highlight_keys: row keys to bold (subtotals)
        fmt_map: {key: format_string} e.g. {'ebitda_margin': '{:.1f}%'}
        """
        if not fin_dict or not fy_labels:
            return

        highlight_keys = highlight_keys or []
        fmt_map = fmt_map or {}

        # Determine projection start column (1-indexed because col 0 is row label)
        proj_start_col = None
        if actuals_end_fy and actuals_end_fy in fy_labels:
            proj_start_col = fy_labels.index(actuals_end_fy) + 2  # +1 for label col, +1 for next

        headers = [""] + list(fy_labels)
        rows = []
        highlight_row_indices = []

        for key in row_keys:
            if key not in fin_dict:
                continue
            series = fin_dict[key]
            if not isinstance(series, list):
                continue
            label = row_labels.get(key, key)
            fmt = fmt_map.get(key, "{:,.0f}")
            row_vals = [label]
            for val in series:
                row_vals.append(self._format_value(val, fmt))
            # Pad if series shorter than fy_labels
            while len(row_vals) < len(headers):
                row_vals.append("\u2013")
            rows.append(tuple(row_vals[:len(headers)]))
            if key in highlight_keys:
                highlight_row_indices.append(len(rows) - 1)

        if rows:
            make_fin_table(
                self.doc,
                headers,
                rows,
                highlight_rows=highlight_row_indices,
                proj_start_col=proj_start_col,
                theme=self.theme,
            )

    def _format_value(self, value, fmt: str = "{:,.0f}") -> str:
        """Format a financial value for display. None -> em dash."""
        if value is None:
            return "\u2013"
        try:
            return fmt.format(value)
        except (ValueError, TypeError):
            return str(value)

    def _embed_chart(self, charts: dict, chart_key: str,
                     width: float = 6.5, caption: str = None) -> None:
        """Safely embed a chart if the path exists. Skip silently if not found."""
        if not charts:
            return
        path = charts.get(chart_key, "")
        if path and os.path.exists(str(path)):
            add_chart(self.doc, str(path), width_inches=width,
                      caption=caption, theme=self.theme)
